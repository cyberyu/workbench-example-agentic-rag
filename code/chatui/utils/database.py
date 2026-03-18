# SPDX-FileCopyrightText: Copyright (c) 2024 NVIDIA CORPORATION & AFFILIATES. All rights reserved.
# SPDX-License-Identifier: Apache-2.0
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
# http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    WebBaseLoader,
    UnstructuredPDFLoader,
    UnstructuredWordDocumentLoader,
    TextLoader,
    CSVLoader 
)
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings

from typing import Any, Dict, List, Tuple, Union
from urllib.parse import urlparse
import os
import shutil
import mimetypes



# Handling which API to use based on public vs NVIDIA internal
# Check if the INTERNAL_API environment variable is set to yes. Don't set if not NVIDIA employee, as you can't access them.
INTERNAL_API = os.getenv('INTERNAL_API', 'no')

# Local embedding model (runs via sentence-transformers, no API key needed)
EMBEDDINGS_MODEL = 'all-MiniLM-L6-v2'

# Get project root dynamically - works both inside and outside AI Workbench
PROJECT_ROOT = os.environ.get('PROJECT_ROOT', os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..')))
DATA_DIR = os.path.join(PROJECT_ROOT, 'data')
# Ensure data directory exists
os.makedirs(DATA_DIR, exist_ok=True)

# Set the chunk size and overlap for the text splitter. Uses defaults but allows them to be set as environment variables.
DEFAULT_CHUNK_SIZE = 250
DEFAULT_CHUNK_OVERLAP = 0

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", DEFAULT_CHUNK_SIZE))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", DEFAULT_CHUNK_OVERLAP))

# Style annotation for CALS <entry> elements (bold / indent).
# STYLE_ANNOTATE_METHOD: "none" (default) | "pdfplumber" | "vlm"
#   pdfplumber — reads font name + x0 position directly from the LibreOffice PDF (fast, accurate)
#   vlm        — queries Qwen2.5-VL via LM Studio (slow, requires local model server)
VLM_LMSTUDIO_URL = os.getenv("VLM_LMSTUDIO_URL", "http://localhost:1234/v1")
STYLE_ANNOTATE_METHOD = os.getenv("STYLE_ANNOTATE_METHOD", "none").lower()
# Legacy alias: VLM_STYLE_ANNOTATE=yes is treated as STYLE_ANNOTATE_METHOD=vlm
if os.getenv("VLM_STYLE_ANNOTATE", "no").lower() == "yes" and STYLE_ANNOTATE_METHOD == "none":
    STYLE_ANNOTATE_METHOD = "vlm"

# Module-level caches so LM Studio is not queried for the model name on every
# table, and page images are not re-encoded to base64 on every VLM call.
# Keys: API base_url -> model_id string
_vlm_model_id_cache: dict = {}
# Keys: (img_path, max_img_px) -> data URI string
_vlm_img_b64_cache: dict = {}

print(f"[config] Using local embedding model: {EMBEDDINGS_MODEL}")


# Adding nltk data
import nltk

def download_nltk_if_missing():
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')

    try:
        nltk.data.find('taggers/averaged_perceptron_tagger')
    except LookupError:
        nltk.download('averaged_perceptron_tagger')

download_nltk_if_missing()

    
# nltk.download("punkt")
# nltk.download("averaged_perceptron_tagger")

# Functions for dealing with URLs
def is_valid_url(url: str) -> bool:
    """ This is a helper function for checking if the URL is valid. It isn't fail proof, but it will catch most common errors. """
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False

def safe_load(url):
    """ This is a helper function for loading the URL. It protects against false negatives from is_value_url and 
        filters for actual web pages. Returns None if it fails.
    """
    try:
        return WebBaseLoader(url).load()
    except Exception as e:
        print(f"[upload] Skipping {url}: {e}")
        return None


def upload(urls: List[str]):
    """ This is a helper function for parsing the user inputted URLs and uploading them into the vector store. """

    urls = [url for url in urls if is_valid_url(url)]

    docs = []
    for url in urls:
        result = safe_load(url)
        if result is not None:
            docs.append(result)

    docs_list = [item for sublist in docs for item in sublist]

    if not docs_list:
        # If no documents were loaded, return None
        print("[upload] No URLs provided.")
        return None
    
    try:
        text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=CHUNK_SIZE, chunk_overlap=0
        )
        doc_splits = text_splitter.split_documents(docs_list)

        vectorstore = Chroma.from_documents(
            documents=doc_splits,
            collection_name="rag-chroma",
            embedding=HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL),
            persist_directory=DATA_DIR,
        )
        return vectorstore

    except Exception as e:
        print(f"[upload] Vectorstore creation failed: {e}")
        return None


# Functions for dealing with file uploads/embeddings

## Helper functions

def _indent_xml(elem, level: int = 0) -> None:
    """Add pretty-print indentation to an ElementTree element in-place."""
    pad = "\n" + "  " * level
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = pad + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = pad
        last = None
        for child in elem:
            _indent_xml(child, level + 1)
            last = child
        if last is not None and (not last.tail or not last.tail.strip()):
            last.tail = pad
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = pad


def _extract_tc_text(tc_el, qn_fn) -> str:
    """Robustly extract all text from a <w:tc> lxml element.

    Handles the full range of OOXML structures found in real-world tables:

    * Multi-paragraph cells   — joined with ``\\n`` (preserves line breaks).
    * `<w:del>` (tracked deletions) — **excluded**; only content visible in the
      final accepted document is returned.
    * `<w:ins>` / `<w:hyperlink>` / `<w:sdt>` — runs inside these wrappers are
      included because they represent permanent visible text.
    * `<w:tab>`               — replaced with a single space (tab characters in
      table cells are decorative indentation, not meaningful delimiters).
    * `<w:sym>`               — skipped (symbol-font glyphs have no Unicode text).
    * Nested `<w:tbl>`        — recursed; each nested cell's text is appended as
      an additional paragraph so the content is not lost.
    """
    qn = qn_fn

    _SKIP_TAGS = {qn("w:del")}          # excluded entirely
    _WRAP_TAGS = {qn("w:ins"), qn("w:hyperlink"), qn("w:sdt"),
                  qn("w:sdtContent"), qn("w:smartTag")}   # transparent wrappers

    def _run_text(r_el) -> str:
        """Extract text from a single <w:r> element, handling tab/t children."""
        buf = []
        for child in r_el:
            if child.tag == qn("w:t"):
                buf.append(child.text or "")
            elif child.tag == qn("w:tab"):
                buf.append(" ")
            # w:br → newline within a run; treat as space
            elif child.tag == qn("w:br"):
                buf.append(" ")
        return "".join(buf)

    def _para_text(p_el) -> str:
        """Extract text from a <w:p>, skipping deleted content."""
        buf = []
        def _walk(el):
            for child in el:
                if child.tag in _SKIP_TAGS:
                    continue          # deleted → skip whole subtree
                elif child.tag == qn("w:r"):
                    buf.append(_run_text(child))
                elif child.tag in _WRAP_TAGS:
                    _walk(child)      # transparent — descend
                # w:tbl inside w:p is unusual but skip (handled at tc level)
        _walk(p_el)
        return "".join(buf).strip()

    def _tc_recurse(tc_el) -> list:
        """Return list of non-empty paragraph text strings from tc_el."""
        parts = []
        for child in tc_el:
            if child.tag == qn("w:p"):
                t = _para_text(child)
                if t:
                    parts.append(t)
            elif child.tag == qn("w:tbl"):
                # Nested table — recurse into every (non-continuation) cell
                for tr in child.findall(qn("w:tr")):
                    for nested_tc in tr.findall(qn("w:tc")):
                        # Skip vMerge-continuation cells
                        vm = nested_tc.find(f"{qn('w:tcPr')}/{qn('w:vMerge')}")
                        if vm is not None and vm.get(qn("w:val")) != "restart":
                            continue
                        parts.extend(_tc_recurse(nested_tc))
        return parts

    return "\n".join(_tc_recurse(tc_el))


def _build_cals_xml(table, title: str = "", table_id: str = "") -> str:
    """Build an OASIS CALS table XML string from a python-docx Table object.

    Maps Word XML properties to CALS attributes:
      - w:gridSpan   → namest / nameend  (horizontal column span)
      - w:vMerge     → morerows="N"     (vertical row span count)
      - paragraph jc → align            (right / center / justify; left is default)
      - w:tblGrid    → proportional colwidth values
      - first row    → <thead>  (remaining rows go into <tbody>)

    Returns a well-formed XML fragment (no XML declaration) ready for embedding
    in an SGML/XML XPP source document.
    """
    import xml.etree.ElementTree as ET
    from docx.oxml.ns import qn

    tbl = table._tbl  # lxml element representing <w:tbl>

    # Convenience: extract all text from a <w:tc> lxml element
    def _tc_text(tc_el) -> str:
        return _extract_tc_text(tc_el, qn)

    # ------------------------------------------------------------------
    # 1. Column widths from <w:tblGrid>
    # ------------------------------------------------------------------
    tbl_grid = tbl.find(qn("w:tblGrid"))
    col_widths: list = []
    if tbl_grid is not None:
        for gc in tbl_grid.findall(qn("w:gridCol")):
            w_val = gc.get(qn("w:w"))
            try:
                col_widths.append(int(w_val) if w_val else 1440)
            except ValueError:
                col_widths.append(1440)

    if not col_widths:
        # Fallback: count unique cells in first row
        try:
            seen: set = set()
            for c in table.rows[0].cells:
                seen.add(id(c._tc))
            col_widths = [1440] * len(seen)
        except Exception:
            col_widths = [1440]

    num_cols = len(col_widths)
    total_w = sum(col_widths) or 1

    # ------------------------------------------------------------------
    # 2. Build per-row cell info list using raw XML (not row.cells)
    #
    # row.cells uses python-docx's virtual grid and can return cells from
    # OTHER rows to fill empty grid positions — which breaks span detection.
    # Iterating _tr.findall(qn('w:tc')) gives only actual XML cells.
    #
    # vMerge values:
    #   None       – not part of any vertical merge
    #   "restart"  – first (content) row of a vertical span
    #   "continue" – subsequent (covered) rows of a vertical span
    # ------------------------------------------------------------------
    rows = table.rows
    nrows = len(rows)
    row_data: list = []

    for row in rows:
        row_entry: list = []
        col_cursor = 0  # tracks grid-column index as we walk left→right

        for tc in row._tr.findall(qn("w:tc")):
            # Horizontal span
            gs_el = tc.find(f"{qn('w:tcPr')}/{qn('w:gridSpan')}")
            gridspan = 1
            if gs_el is not None:
                try:
                    gridspan = int(gs_el.get(qn("w:val")) or 1)
                except ValueError:
                    pass

            # Vertical span marker from Word XML
            vm_el = tc.find(f"{qn('w:tcPr')}/{qn('w:vMerge')}")
            vmerge = None
            if vm_el is not None:
                vm_val = vm_el.get(qn("w:val"))
                # "restart" = first row of span (has content)
                # absent val or any other val = continuation (no content)
                vmerge = "restart" if vm_val == "restart" else "continue"

            # Paragraph alignment from the first <w:p>
            align = "left"
            p_el = tc.find(qn("w:p"))
            if p_el is not None:
                jc_el = p_el.find(f"{qn('w:pPr')}/{qn('w:jc')}")
                if jc_el is not None:
                    jc_val = jc_el.get(qn("w:val")) or ""
                    _amap = {
                        "right": "right", "center": "center",
                        "both": "justify", "distribute": "justify",
                    }
                    align = _amap.get(jc_val, "left")

            # Cell text — collect text runs from all paragraphs
            text = _tc_text(tc)

            row_entry.append({
                "col_start": col_cursor,
                "col_end": col_cursor + gridspan,
                "text": text,
                "gridSpan": gridspan,
                "vMerge": vmerge,
                "align": align,
            })
            col_cursor += gridspan

        row_data.append(row_entry)

    # ------------------------------------------------------------------
    # 3. Map each cell to its starting grid-column (already in cell_dict)
    #    and pre-compute morerows for "restart" vMerge cells.
    # ------------------------------------------------------------------
    # Build a lookup: (row_idx, col_start) → vMerge type, for morerows counting
    vm_lookup: dict = {}
    for r_idx, row_entry in enumerate(row_data):
        for cd in row_entry:
            vm_lookup[(r_idx, cd["col_start"])] = cd["vMerge"]

    def _morerows(start_r: int, col_start: int) -> int:
        """Count consecutive vMerge-continuation cells below start_r at col_start."""
        count = 0
        for r in range(start_r + 1, nrows):
            if vm_lookup.get((r, col_start)) == "continue":
                count += 1
            else:
                break
        return count

    # cell_grid kept as flat alias for row_data (col_start/col_end already in dicts)
    cell_grid = row_data

    # ------------------------------------------------------------------
    # 4. Assemble CALS XML using ElementTree
    # ------------------------------------------------------------------
    root = ET.Element("table")
    if table_id:
        root.set("id", table_id)
    if title:
        t_el = ET.SubElement(root, "title")
        t_el.text = title

    tgroup = ET.SubElement(root, "tgroup")
    tgroup.set("cols", str(num_cols))

    for i, w in enumerate(col_widths):
        cs = ET.SubElement(tgroup, "colspec")
        cs.set("colnum", str(i + 1))
        cs.set("colname", f"c{i + 1}")
        # Proportional width: widths sum to num_cols * (so columns are comparable)
        cs.set("colwidth", f"{w / total_w * num_cols:.3f}*")

    # First row → <thead> when there are multiple rows
    head_rows = 1 if nrows > 1 else 0
    thead_el = ET.SubElement(tgroup, "thead") if head_rows else None
    tbody_el = ET.SubElement(tgroup, "tbody")

    for r_idx, row_positions in enumerate(cell_grid):
        parent = thead_el if (r_idx < head_rows and thead_el is not None) else tbody_el
        row_el = ET.SubElement(parent, "row")

        for cell_dict in row_positions:
            if cell_dict["vMerge"] == "continue":
                continue  # covered by a row-spanning entry above

            col_start = cell_dict["col_start"]
            col_end = cell_dict["col_end"]

            entry_el = ET.SubElement(row_el, "entry")

            if cell_dict["gridSpan"] > 1:
                entry_el.set("namest", f"c{col_start + 1}")
                entry_el.set("nameend", f"c{col_end}")
            else:
                entry_el.set("colname", f"c{col_start + 1}")

            mr = _morerows(r_idx, col_start) if cell_dict["vMerge"] == "restart" else 0
            if mr:
                entry_el.set("morerows", str(mr))

            if cell_dict["align"] != "left":
                entry_el.set("align", cell_dict["align"])

            entry_el.text = cell_dict["text"] if cell_dict["text"] else None

    _indent_xml(root)
    return ET.tostring(root, encoding="unicode")


def _annotate_entry_styles_from_ooxml(cals_xml: str, table) -> str:
    """Annotate CALS <entry> elements with bold/indent attributes sourced
    directly from the OOXML run and paragraph properties of *table*.

    This is more reliable than the pdfplumber-based annotator because it reads
    Word's own structural markup instead of inferring attributes from glyph
    metrics in the LibreOffice-rendered PDF.  Text that is split across runs,
    wrapped, or uses non-breaking spaces is handled correctly.

    Attributes written:
      bold="true"   — any run in the cell's first paragraph has <w:b>
                      (and the value is not explicitly "false" or "0")
      indent="N"    — paragraph left-indent level 1–3; the smallest non-zero
                      indent across all non-header cells defines level 1.

    The function traverses <w:tr>/<w:tc> in the same order as _build_cals_xml
    (skipping vMerge-continuation cells) and applies attributes to the
    positionally corresponding <entry> elements in the CALS XML.

    Returns the annotated XML string, or *cals_xml* unchanged on any error.
    """
    import xml.etree.ElementTree as ET
    from docx.oxml.ns import qn

    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError:
        return cals_xml

    tbl = table._tbl

    # ── Pass 1: collect (bold, indent_twips) per non-continuation tc ─────────
    tc_props: list = []   # [(bold: bool, indent_twips: int), ...]
    for tr in tbl.findall(qn("w:tr")):
        for tc in tr.findall(qn("w:tc")):
            # Skip vMerge continuation cells — they are omitted from CALS output
            vm_el = tc.find(f"{qn('w:tcPr')}/{qn('w:vMerge')}")
            if vm_el is not None and vm_el.get(qn("w:val")) != "restart":
                continue

            # Bold: any run in any paragraph has <w:b> not explicitly disabled
            bold = False
            for p_el in tc.findall(qn("w:p")):
                for r_el in p_el.findall(qn("w:r")):
                    rpr = r_el.find(qn("w:rPr"))
                    if rpr is not None:
                        b_el = rpr.find(qn("w:b"))
                        if b_el is not None:
                            val = b_el.get(qn("w:val"), "true")
                            if val not in ("false", "0"):
                                bold = True
                                break
                if bold:
                    break

            # Indent: minimum left-indent across all paragraphs in the cell
            cell_indent = 0
            for p_el in tc.findall(qn("w:p")):
                ppr = p_el.find(qn("w:pPr"))
                if ppr is not None:
                    ind = ppr.find(qn("w:ind"))
                    if ind is not None:
                        try:
                            twips = int(ind.get(qn("w:left"), "0") or "0")
                        except ValueError:
                            twips = 0
                        if twips > 0 and (cell_indent == 0 or twips < cell_indent):
                            cell_indent = twips
                    # Also check <w:numPr> — numbered list items get a hanging
                    # indent that is NOT reflected in <w:ind w:left> on older docs.
                    num_el = ppr.find(qn("w:numPr"))
                    if num_el is not None and cell_indent == 0:
                        cell_indent = 360   # treat as one implicit level

            tc_props.append((bold, cell_indent))

    # ── Pass 2: compute normalised indent levels ──────────────────────────────
    # The smallest non-zero indent value seen across ALL cells becomes level 1.
    nonzero = sorted({t for _, t in tc_props if t > 0})
    indent_step = nonzero[0] if nonzero else 360   # 360 twips ≈ 0.25 in

    def _level(twips: int) -> int:
        return min(3, max(1, round(twips / indent_step))) if twips > 0 else 0

    # ── Pass 3: apply to CALS <entry> elements in document order ─────────────
    entries = list(root.iter("entry"))
    for i, entry_el in enumerate(entries):
        if i >= len(tc_props):
            break
        bold, indent_twips = tc_props[i]
        if bold:
            entry_el.set("bold", "true")
        lvl = _level(indent_twips)
        if lvl > 0:
            entry_el.set("indent", str(lvl))

    _indent_xml(root)
    return ET.tostring(root, encoding="unicode")


def _cals_to_html(cals_xml: str) -> str:
    """Convert a CALS XML string (produced by _build_cals_xml) to a styled HTML table.

    Colour coding matches the verification annotation:
      - green background  : verify="ok"
      - red background    : verify="unconfirmed"
      - default           : no verify attribute (e.g. un-inspected tables)
    """
    import xml.etree.ElementTree as ET

    if not cals_xml:
        return "<p style='color:#888; font-family:sans-serif;'>No CALS XML available</p>"

    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError as exc:
        return f"<p style='color:red; font-family:sans-serif;'>XML parse error: {exc}</p>"

    title_el = root.find("title")
    title_html = ""
    if title_el is not None and title_el.text:
        title_html = (
            f"<caption style='font-weight:bold;text-align:left;"
            f"padding:4px 0;font-family:sans-serif;'>{title_el.text}</caption>"
        )

    TABLE_STYLE = (
        "border-collapse:collapse;width:100%;font-family:monospace;"
        "font-size:11px;margin-top:4px;color:#111;background:#fff;"
    )
    TH_STYLE = "border:1px solid #888;padding:4px 8px;background:#d8e8f8;color:#111;font-weight:bold;"
    TD_STYLE = "border:1px solid #ccc;padding:4px 8px;color:#111;background:#fff;"
    OK_STYLE  = "border:1px solid #5a5;padding:4px 8px;background:#d4f8d4;color:#111;"
    BAD_STYLE = "border:1px solid #a55;padding:4px 8px;background:#f8d4d4;color:#111;"

    def _cell_markup(entry_el, is_head=False):
        tag = "th" if is_head else "td"

        # colspan from namest/nameend (column names are "c1", "c2", ...)
        namest  = entry_el.get("namest")
        nameend = entry_el.get("nameend")
        colspan = 1
        if namest and nameend:
            try:
                colspan = int(nameend.lstrip("c")) - int(namest.lstrip("c")) + 1
            except (ValueError, AttributeError):
                pass

        # rowspan from morerows attribute
        morerows = entry_el.get("morerows")
        rowspan = 1
        if morerows:
            try:
                rowspan = int(morerows) + 1
            except ValueError:
                pass

        verify = entry_el.get("verify")
        if is_head:
            style = TH_STYLE
        elif verify == "ok":
            style = OK_STYLE
        elif verify == "unconfirmed":
            style = BAD_STYLE
        else:
            style = TD_STYLE

        text = (entry_el.text or "").replace("<", "&lt;").replace(">", "&gt;")

        attrs = f'style="{style}"'
        if colspan > 1:
            attrs += f' colspan="{colspan}"'
        if rowspan > 1:
            attrs += f' rowspan="{rowspan}"'
        return f"<{tag} {attrs}>{text}</{tag}>"

    parts = [f'<div style="overflow-x:auto;color:#111;background:#fff;padding:4px;"><table style="{TABLE_STYLE}">{title_html}']

    for tgroup in root.findall("tgroup"):
        thead = tgroup.find("thead")
        if thead is not None:
            parts.append("<thead>")
            for row in thead.findall("row"):
                parts.append("<tr>")
                for entry in row.findall("entry"):
                    parts.append(_cell_markup(entry, is_head=True))
                parts.append("</tr>")
            parts.append("</thead>")

        tbody = tgroup.find("tbody")
        if tbody is not None:
            parts.append("<tbody>")
            for row in tbody.findall("row"):
                parts.append("<tr>")
                for entry in row.findall("entry"):
                    parts.append(_cell_markup(entry, is_head=False))
                parts.append("</tr>")
            parts.append("</tbody>")

    parts.append("</table></div>")
    return "\n".join(parts)


# ── Table rendering themes ────────────────────────────────────────────────────
_CALS_HTML_THEMES = {
    "verify": {
        "wrap":    "overflow-x:auto;color:#111;background:#fff;padding:4px;",
        "table":   "border-collapse:collapse;width:100%;font-family:monospace;font-size:11px;margin-top:4px;color:#111;background:#fff;",
        "caption": "font-weight:bold;text-align:left;padding:4px 0;font-family:sans-serif;color:#111;",
        "th":      "border:1px solid #888;padding:4px 8px;background:#d8e8f8;color:#111;font-weight:bold;cursor:pointer;",
        "td":      "border:1px solid #ccc;padding:4px 8px;color:#111;background:#fff;cursor:pointer;",
        "ok":      "border:1px solid #5a5;padding:4px 8px;background:#d4f8d4;color:#111;cursor:pointer;",
        "bad":     "border:1px solid #a55;padding:4px 8px;background:#f8d4d4;color:#111;cursor:pointer;",
        "use_verify": True,
    },
    "finance": {
        "wrap":    "overflow-x:auto;color:#1a1a1a;background:#fafbfc;padding:10px 4px;",
        "table":   "border-collapse:collapse;width:100%;font-family:Arial,Helvetica,sans-serif;font-size:11px;margin-top:4px;color:#1a1a1a;",
        "caption": "font-weight:bold;text-align:left;padding:6px 0;font-family:Arial,Helvetica,sans-serif;font-size:12px;color:#0d2a6e;text-transform:uppercase;letter-spacing:0.05em;",
        "th":      "border-bottom:2px solid #0d2a6e;border-top:1px solid #0d2a6e;padding:6px 10px;background:#fafbfc;color:#0d2a6e;font-weight:bold;cursor:pointer;",
        "td":      "border-bottom:1px solid #e0e4ea;padding:5px 10px;color:#1a1a1a;background:#fafbfc;cursor:pointer;",
        "ok":      "border-bottom:1px solid #e0e4ea;padding:5px 10px;color:#1a1a1a;background:#fafbfc;cursor:pointer;",
        "bad":     "border-bottom:1px solid #e0e4ea;padding:5px 10px;color:#777;background:#fafbfc;cursor:pointer;font-style:italic;",
        "use_verify": False,
    },
    "dark": {
        "wrap":    "overflow-x:auto;color:#cdd6f4;background:#1e1e2e;padding:8px;border-radius:6px;",
        "table":   "border-collapse:collapse;width:100%;font-family:monospace;font-size:11px;margin-top:4px;color:#cdd6f4;background:#1e1e2e;",
        "caption": "font-weight:bold;text-align:left;padding:4px 0;font-family:monospace;color:#89b4fa;font-size:12px;",
        "th":      "border:1px solid #3a4a6b;padding:5px 10px;background:#263050;color:#89b4fa;font-weight:bold;cursor:pointer;",
        "td":      "border:1px solid #2a2a3e;padding:4px 10px;color:#cdd6f4;background:#1e1e2e;cursor:pointer;",
        "ok":      "border:1px solid #2a2a3e;padding:4px 10px;color:#a6e3a1;background:#1e2b20;cursor:pointer;",
        "bad":     "border:1px solid #2a2a3e;padding:4px 10px;color:#f38ba8;background:#2e1e20;cursor:pointer;",
        "use_verify": True,
    },
    "minimal": {
        "wrap":    "overflow-x:auto;color:#222;background:#fff;padding:8px 4px;",
        "table":   "border-collapse:collapse;width:100%;font-family:Georgia,serif;font-size:11.5px;margin-top:4px;color:#222;background:#fff;",
        "caption": "font-weight:bold;text-align:left;padding:6px 0;font-family:Georgia,serif;font-size:13px;color:#222;text-transform:uppercase;letter-spacing:0.07em;",
        "th":      "border-bottom:2px solid #222;padding:6px 12px;background:#fff;color:#222;font-weight:bold;cursor:pointer;",
        "td":      "border-bottom:1px solid #e0e0e0;padding:5px 12px;color:#222;background:#fff;cursor:pointer;",
        "ok":      "border-bottom:1px solid #e0e0e0;padding:5px 12px;color:#222;background:#fff;cursor:pointer;",
        "bad":     "border-bottom:1px solid #e0e0e0;padding:5px 12px;color:#aaa;background:#fff;cursor:pointer;",
        "use_verify": False,
    },
    "striped": {
        "wrap":    "overflow-x:auto;color:#111;background:#fff;padding:4px;",
        "table":   "border-collapse:collapse;width:100%;font-family:sans-serif;font-size:11px;margin-top:4px;color:#111;background:#fff;",
        "caption": "font-weight:bold;text-align:left;padding:4px 0;font-family:sans-serif;color:#333;",
        "th":      "padding:6px 10px;background:#404060;color:#fff;font-weight:bold;cursor:pointer;",
        "td":      None,  # computed per-row
        "td_even": "padding:5px 10px;color:#111;background:#eef0f8;cursor:pointer;",
        "td_odd":  "padding:5px 10px;color:#111;background:#fff;cursor:pointer;",
        "ok":      None,  # computed per-row (same as td, no verify highlight)
        "bad":     None,  # computed per-row (same as td, no verify highlight)
        "use_verify": False,
    },
    # Based on the iXBRL financial-statements CSS published by XII at
    # https://gitlab.xbrl.org/xii/ixbrl-samples (xii-2020 branch)
    "ixbrl": {
        "wrap":    "overflow-x:auto;color:#231f20;background:#fff;padding:10px 4px;",
        "table":   "border-collapse:collapse;width:100%;font-family:'Gotham',Arial,Helvetica,sans-serif;font-size:9pt;line-height:13pt;color:#231f20;background:#fff;font-weight:300;",
        "caption": "font-weight:normal;text-align:left;padding:6px 0;font-family:'Gotham',Arial,Helvetica,sans-serif;font-size:10pt;color:#46a2df;",
        "th":      "border-bottom:3px solid #bdbdbc;padding:4px 8px;color:#46a2df;font-weight:normal;box-sizing:border-box;height:25px;background:#fff;cursor:pointer;",
        "td":      "padding:4px 8px;color:#231f20;background:#fff;box-sizing:border-box;height:25px;cursor:pointer;",
        "ok":      "padding:4px 8px;color:#231f20;background:#fff;box-sizing:border-box;height:25px;cursor:pointer;",
        "bad":     "padding:4px 8px;color:#999;background:#fff;box-sizing:border-box;height:25px;cursor:pointer;font-style:italic;",
        "use_verify": False,
    },
}


def _cals_to_interactive_html(cals_xml: str, theme: str = "verify", clean_xml: str = None):
    """Return (table_html, xml_panel_html) for the interactive Table Browser.

    cals_xml       — annotated CALS XML (may contain verify= attributes for colour coding)
    clean_xml      — optional clean CALS XML (no verify= attrs); when supplied,
                     each annotation span gets a data-clean-line attribute containing
                     the 0-based line number of the corresponding <entry> in clean_xml,
                     enabling precise jump-to-line in the Edit Clean XML CodeMirror panel.
    table_html     — rendered HTML table; each cell has an onclick handler.
    xml_panel_html — dark-themed <pre> block with the pretty-printed annotation XML;
                     each <entry> is wrapped in <span id="entry-rN-cK"> so the
                     onclick handler can locate it via document.getElementById.

    Because Gradio injects gr.HTML content via innerHTML (which does not execute
    <script> tags), all JavaScript is written as inline onclick= attributes.
    """
    import xml.etree.ElementTree as ET
    import html as _html
    import re

    if not cals_xml:
        empty = "<p style='color:#888;font-family:sans-serif;'>No CALS XML available</p>"
        return empty, empty

    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError as exc:
        err = f"<p style='color:red;font-family:sans-serif;'>XML parse error: {exc}</p>"
        return err, err

    # CSS injected into both components so the highlight class is always defined.
    HI_CSS = (
        "<style>.cals-hi{background:#ffe080!important;color:#111!important;"
        "outline:2px solid #d4a000;border-radius:2px;}</style>"
    )

    # Inline onclick JS — no named function needed (innerHTML doesn't run <script>).
    # Single quotes inside JS are safe inside a double-quoted HTML attribute.
    # Bridge strategy for Edit Clean XML sync:
    #   1. Highlight the annotation XML span (existing, reliable).
    #   2. Read its textContent to extract the cell value.
    #   3. Scroll CM6 to the approximate line, then search visible .cm-line
    #      elements by colname AND cell value — colname alone is not unique
    #      (every row has a col c2), but colname+value is almost always unique.
    # IMPORTANT: no literal " may appear inside onclick="..." HTML attributes.
    # col_key is a plain identifier (e.g. "c2") safe to interpolate directly.
    def _onclick(eid):
        return (
            # Step 1: switch to View Annotation XML tab so the span is visible
            "var _tabs=document.querySelectorAll('#tb-right-tabs button[role=tab]');"
            "for(var _ti=0;_ti<_tabs.length;_ti++){"
              "if(_tabs[_ti].textContent.indexOf('View Annotation XML')>=0)"
              "{_tabs[_ti].click();break;}"
            "}"
            # Step 2: clear old highlight
            "var _a=document.querySelectorAll('.cals-hi');"
            "for(var _i=0;_i<_a.length;_i++)_a[_i].classList.remove('cals-hi');"
            # Step 3: after tab reveal, highlight + scroll span;
            #         store data-clean-line for optional Edit Clean XML jump
            f"var _eid='{eid}';"
            "setTimeout(function(){"
              "var _el=document.getElementById(_eid);"
              "if(_el){"
                "_el.classList.add('cals-hi');"
                "_el.scrollIntoView({behavior:'smooth',block:'center'});"
                "window._calsJumpLine=parseInt(_el.getAttribute('data-clean-line')||-1,10);"
              "}"
            "},150);"
            # Step 4: open inline popup editor
            f"if(window._calsEditPopup)window._calsEditPopup('{eid}');"
        )

    # Assign stable DOM IDs:  entry-r{row_idx}-{col_key}
    def _assign_ids(xml_root):
        eids = {}
        for tg in xml_root.findall("tgroup"):
            row_idx = 0
            for section in (tg.find("thead"), tg.find("tbody")):
                if section is None:
                    continue
                for row in section.findall("row"):
                    for entry in row.findall("entry"):
                        col_key = (
                            entry.get("colname")
                            or entry.get("namest")
                            or f"x{row_idx}"
                        )
                        eids[id(entry)] = f"entry-r{row_idx}-{col_key}"
                    row_idx += 1
        return eids

    entry_ids = _assign_ids(root)

    # ── Rendered table HTML (onclick on every cell) ────────────────────────
    _th         = _CALS_HTML_THEMES.get(theme, _CALS_HTML_THEMES["verify"])
    TABLE_STYLE = _th["table"]
    WRAP        = _th["wrap"]
    TH          = _th["th"]
    TD          = _th["td"] or _th.get("td_even", "padding:4px 8px;cursor:pointer;")
    OK          = _th["ok"]  or TD
    BAD         = _th["bad"] or TD
    USE_VERIFY  = _th.get("use_verify", True)
    IS_STRIPED  = (theme == "striped")

    title_el = root.find("title")
    caption  = ""
    if title_el is not None and title_el.text:
        caption = f"<caption style='{_th['caption']}'>{_html.escape(title_el.text)}</caption>"

    def _norm_display(v):
        """Return (raw, normalized) strings for mismatch explanation."""
        raw = str(v or "").strip()
        n = re.sub(r"[\$,\(\)\s\u2014\u2013]", "", raw).replace("-", "").strip()
        if not n:
            return raw, "(empty after normalisation)"
        try:
            return raw, str(float(n))
        except ValueError:
            return raw, n.lower()

    def _cell(entry_el, is_head=False, row_idx=0):
        tag  = "th" if is_head else "td"
        eid  = entry_ids.get(id(entry_el), "")
        namest, nameend = entry_el.get("namest"), entry_el.get("nameend")
        colspan = 1
        if namest and nameend:
            try:
                colspan = int(nameend.lstrip("c")) - int(namest.lstrip("c")) + 1
            except (ValueError, AttributeError):
                pass
        morerows = entry_el.get("morerows")
        rowspan  = (int(morerows) + 1) if morerows else 1
        verify   = entry_el.get("verify")
        if is_head:
            css = TH
        elif IS_STRIPED:
            css = _th["td_even"] if row_idx % 2 == 0 else _th["td_odd"]
        elif USE_VERIFY:
            css = OK if verify == "ok" else (BAD if verify == "unconfirmed" else TD)
        else:
            css = TD
        # Respect the CALS align attribute (left/right/center/justify)
        _align = entry_el.get("align")
        if _align in ("left", "right", "center", "justify"):
            css = css + f"text-align:{_align};"
        # VLM-annotated bold and indent
        if entry_el.get("bold") == "true":
            css = css + "font-weight:bold;"
        _indent = entry_el.get("indent")
        if _indent and _indent.isdigit() and int(_indent) > 0:
            css = css + f"padding-left:{int(_indent) * 1.5:.1f}em;"
        # Use itertext() so text inside <para> children is included
        cell_full_text = " ".join("".join(entry_el.itertext()).split())
        text     = _html.escape(cell_full_text)

        # Build onclick: always highlight XML panel entry; for mismatched cells
        # also populate the info banner via data-mismatch attribute.
        base_oc = _onclick(eid) if eid else ""
        if not is_head and verify == "unconfirmed":
            raw_txt, norm_txt = _norm_display(cell_full_text)
            verify_reason = entry_el.get("verify-reason", "")
            reason_line = _html.escape(verify_reason) if verify_reason else "Not found in pdfplumber cross-check."
            # &#10; encodes newline inside an HTML attribute; textContent renders it
            # as a real newline with white-space:pre-wrap on the container.
            detail = (
                f"CALS cell text : {_html.escape(raw_txt)}&#10;"
                f"Normalised to  : {_html.escape(norm_txt)}&#10;"
                f"&#10;"
                f"Reason         : {reason_line}&#10;"
                f"&#10;"
                f"Common causes  : leading/trailing spaces, line-breaks inside&#10;"
                f"                 a cell, merged-cell text split across rows,&#10;"
                f"                 or special characters dropped by pdfplumber."
            )
            show_oc = (
                "var _b=document.getElementById('cals-mismatch-info');"
                "_b.style.display='block';"
                "document.getElementById('cals-mismatch-body').textContent="
                "this.getAttribute('data-mismatch');"
            )
            extra = (
                f' onclick="{base_oc}{show_oc}"'
                f' data-mismatch="{detail}"'
                f' title="Mismatch — click for details"'
            )
        elif not is_head and verify == "ok":
            hide_oc = (
                "var _b=document.getElementById('cals-mismatch-info');"
                "_b.style.display='none';"
            )
            extra = f' onclick="{base_oc}{hide_oc}" title="Confirmed by pdfplumber"'
        else:
            extra = f' onclick="{base_oc}" title="Click to locate in XML"' if base_oc else ""

        if colspan > 1:
            extra += f' colspan="{colspan}"'
        if rowspan > 1:
            extra += f' rowspan="{rowspan}"'
        return f'<{tag} style="{css}"{extra}>{text}</{tag}>'

    # Info banner — shown/hidden by the onclick handlers above.
    INFO_BOX = (
        '<div id="cals-mismatch-info" style="display:none;background:#fff3cd;'
        'border:1px solid #c9a700;padding:8px 14px;margin-bottom:10px;'
        'border-radius:4px;font-family:monospace;font-size:11px;color:#111;'
        'white-space:pre-wrap;">'
        '<b style="color:#7a5200;">&#9888; Mismatch detail</b>'
        '<span style="float:right;cursor:pointer;font-size:14px;color:#7a5200;" '
        'onclick="this.parentElement.style.display=\'none\'">&#x2715;</span>'
        '<div id="cals-mismatch-body" style="margin-top:6px;"></div>'
        '</div>'
    )

    tbl = [
        HI_CSS,
        f'<div style="{WRAP}">',
        INFO_BOX,
        f'<table style="{TABLE_STYLE}">{caption}',
    ]
    for tg in root.findall("tgroup"):
        thead = tg.find("thead")
        if thead is not None:
            tbl.append("<thead>")
            for row in thead.findall("row"):
                tbl.append("<tr>" + "".join(_cell(e, True)  for e in row.findall("entry")) + "</tr>")
            tbl.append("</thead>")
        tbody = tg.find("tbody")
        if tbody is not None:
            tbl.append("<tbody>")
            for _ri, row in enumerate(tbody.findall("row")):
                tbl.append("<tr>" + "".join(_cell(e, False, _ri) for e in row.findall("entry")) + "</tr>")
            tbl.append("</tbody>")
    tbl.append("</table></div>")
    table_html = "\n".join(tbl)

    # ── XML panel HTML (dark code view with span IDs per entry) ───────────
    # Clone the tree, inject _eid attributes, pretty-print, then build HTML
    # with each <entry> line wrapped in a <span id="..."> for the onclick hook.

    # Pre-compute eid → line-in-clean-xml mapping when clean_xml is available.
    # Both annotation and clean XMLs have the same CALS structure so _assign_ids
    # yields the same eid for each structurally equivalent entry.
    eid_to_clean_line: dict = {}
    if clean_xml:
        try:
            root_cl = ET.fromstring(clean_xml)
            eids_cl = _assign_ids(root_cl)
            for tg_cl in root_cl.findall("tgroup"):
                for sec_cl in (tg_cl.find("thead"), tg_cl.find("tbody")):
                    if sec_cl is None:
                        continue
                    for row_cl in sec_cl.findall("row"):
                        for entry_cl in row_cl.findall("entry"):
                            eid_cl = eids_cl.get(id(entry_cl), "")
                            if eid_cl:
                                entry_cl.set("_eid", eid_cl)
            _indent_xml(root_cl)
            cl_raw = ET.tostring(root_cl, encoding="unicode")
            for cl_no, cl_line in enumerate(cl_raw.splitlines()):
                mc = re.search(r'_eid="([^"]+)"', cl_line)
                if mc:
                    eid_to_clean_line[mc.group(1)] = cl_no
        except Exception:
            pass  # clean_xml parse failure — falls back to no data-clean-line

    root2 = ET.fromstring(cals_xml)
    eids2 = _assign_ids(root2)
    for tg in root2.findall("tgroup"):
        for section in (tg.find("thead"), tg.find("tbody")):
            if section is None:
                continue
            for row in section.findall("row"):
                for entry in row.findall("entry"):
                    eid = eids2.get(id(entry), "")
                    if eid:
                        entry.set("_eid", eid)
    _indent_xml(root2)
    xml_raw = ET.tostring(root2, encoding="unicode")

    # Build per-entry metadata dict for popup data attrs (cell text, verify status, reason).
    eid_to_entry_data: dict = {}  # eid -> (cell_text, verify, reason)
    for _tg2 in root2.findall("tgroup"):
        for _sec2 in (_tg2.find("thead"), _tg2.find("tbody")):
            if _sec2 is None:
                continue
            for _row2 in _sec2.findall("row"):
                for _e2 in _row2.findall("entry"):
                    _eid2 = eids2.get(id(_e2), "")
                    if _eid2:
                        _ct2 = " ".join("".join(_e2.itertext()).split())
                        eid_to_entry_data[_eid2] = (
                            _ct2,
                            _e2.get("verify", ""),
                            _e2.get("verify-reason", ""),
                        )

    xml_lines = []
    in_multi  = False
    cur_eid   = ""
    eid_to_line: dict = {}   # eid -> 0-based line index (first line of the entry)
    for line_no, line in enumerate(xml_raw.splitlines()):
        m = re.search(r'_eid="([^"]+)"', line)
        if m:
            eid   = m.group(1)
            cur_eid = eid
            eid_to_line[eid] = line_no
            _cl = eid_to_clean_line.get(eid, -1)
            clean = re.sub(r'\s*_eid="[^"]*"', "", line)
            esc   = _html.escape(clean)
            _ct, _vfy, _rsn = eid_to_entry_data.get(eid, ("", "", ""))
            _da = (
                f' data-verify="{_html.escape(_vfy)}"'
                f' data-reason="{_html.escape(_rsn)}"'
                f' data-celltext="{_html.escape(_ct)}"'
            )
            if "</entry>" in line:
                xml_lines.append(
                    f'<span id="{eid}" data-line="{line_no}" data-clean-line="{_cl}"{_da} class="cals-entry" style="display:block;">{esc}</span>'
                )
                cur_eid = ""
            else:
                xml_lines.append(
                    f'<span id="{eid}" data-line="{line_no}" data-clean-line="{_cl}"{_da} class="cals-entry" style="display:block;">{esc}'
                )
                in_multi = True
        elif in_multi and "</entry>" in line:
            xml_lines.append(f'{_html.escape(line)}</span>')
            in_multi = False
            cur_eid = ""
        else:
            xml_lines.append(_html.escape(line))

    total_lines = len(xml_lines)
    # Gutter: line numbers, right-aligned; width scales with digit count
    _gutter_w = f"{max(len(str(total_lines)) * 0.65 + 0.6, 2.4):.1f}em"
    GUTTER = (
        "font-family:'Courier New',monospace;font-size:11px;line-height:1.6;"
        "background:#2a2a2a;color:#606060;padding:12px 6px 12px 8px;margin:0;"
        f"white-space:pre;text-align:right;user-select:none;min-width:{_gutter_w};"
        "border-right:1px solid #444;"
    )
    CODE = (
        "font-family:'Courier New',monospace;font-size:11px;line-height:1.6;"
        "background:#1e1e1e;color:#d4d4d4;padding:12px;margin:0;"
        "white-space:pre;overflow-x:auto;flex:1;"
    )
    line_nums = "\n".join(str(i + 1) for i in range(total_lines))
    xml_panel_html = (
        HI_CSS
        + '<div style="min-height:600px;overflow-y:auto;background:#1e1e1e;'
        'border-radius:4px;display:flex;align-items:stretch;">'
        + f'<pre style="{GUTTER}">{line_nums}</pre>'
        + f'<pre style="{CODE}">'
        + "\n".join(xml_lines)
        + "</pre></div>"
    )

    return table_html, xml_panel_html


# XSL-FO theme properties — mirrors _CALS_HTML_THEMES but for Apache FOP output.
# Each key is substituted into _CALS_TO_FO_XSLT via simple str.replace().
_CALS_FO_THEMES = {
    "verify": {
        "table_border":     "0.5pt solid #888888",
        "head_bg":          "#d8e8f8",
        "head_cell_border": "0.5pt solid #888888",
        "font_family":      "monospace",
        "font_size":        "8pt",
        "head_font_weight": "bold",
        "head_color":       "#111111",
        "body_cell_border": "0.5pt solid #cccccc",
        "body_color":       "#111111",
        "bg_ok":            "#d4f8d4",
        "bg_bad":           "#f8d4d4",
        "bg_default":       "#ffffff",
        "caption_font":     "sans-serif",
        "caption_size":     "11pt",
    },
    "finance": {
        "table_border":     "0.5pt solid #cccccc",
        "head_bg":          "#0d2a6e",
        "head_cell_border": "2pt solid #0d2a6e",
        "font_family":      "Arial, Helvetica, sans-serif",
        "font_size":        "8pt",
        "head_font_weight": "bold",
        "head_color":       "#ffffff",
        "body_cell_border": "0.5pt solid #cccccc",
        "body_color":       "#1a1a1a",
        "bg_ok":            "#ffffff",
        "bg_bad":           "#ffffff",
        "bg_default":       "#ffffff",
        "caption_font":     "Arial, Helvetica, sans-serif",
        "caption_size":     "11pt",
    },
    "dark": {
        "table_border":     "0.5pt solid #45475a",
        "head_bg":          "#313244",
        "head_cell_border": "0.5pt solid #45475a",
        "font_family":      "monospace",
        "font_size":        "8pt",
        "head_font_weight": "bold",
        "head_color":       "#89b4fa",
        "body_cell_border": "0.5pt solid #45475a",
        "body_color":       "#cdd6f4",
        "bg_ok":            "#1e4a2e",
        "bg_bad":           "#4a1e1e",
        "bg_default":       "#1e1e2e",
        "caption_font":     "monospace",
        "caption_size":     "11pt",
    },
    "minimal": {
        "table_border":     "0.5pt solid #cccccc",
        "head_bg":          "#ffffff",
        "head_cell_border": "1pt solid #222222",
        "font_family":      "Georgia, Times New Roman, serif",
        "font_size":        "8pt",
        "head_font_weight": "bold",
        "head_color":       "#222222",
        "body_cell_border": "0.5pt solid #e0e0e0",
        "body_color":       "#222222",
        "bg_ok":            "#ffffff",
        "bg_bad":           "#ffffff",
        "bg_default":       "#ffffff",
        "caption_font":     "Georgia, Times New Roman, serif",
        "caption_size":     "11pt",
    },
    "striped": {
        "table_border":     "0.5pt solid #888888",
        "head_bg":          "#404060",
        "head_cell_border": "0.5pt solid #888888",
        "font_family":      "Arial, Helvetica, sans-serif",
        "font_size":        "8pt",
        "head_font_weight": "bold",
        "head_color":       "#ffffff",
        "body_cell_border": "0.5pt solid #cccccc",
        "body_color":       "#111111",
        "bg_ok":            "#eef0f8",
        "bg_bad":           "#eef0f8",
        "bg_default":       "#ffffff",
        "caption_font":     "Arial, Helvetica, sans-serif",
        "caption_size":     "11pt",
    },
    "ixbrl": {
        "table_border":     "0.5pt solid #bdbdbc",
        "head_bg":          "#ffffff",
        "head_cell_border": "3pt solid #bdbdbc",
        "font_family":      "Arial, Helvetica, sans-serif",
        "font_size":        "9pt",
        "head_font_weight": "normal",
        "head_color":       "#46a2df",
        "body_cell_border": "0.5pt solid #bdbdbc",
        "body_color":       "#231f20",
        "bg_ok":            "#ffffff",
        "bg_bad":           "#ffffff",
        "bg_default":       "#ffffff",
        "caption_font":     "Arial, Helvetica, sans-serif",
        "caption_size":     "11pt",
    },
}

# XSLT that converts a CALS <table> fragment into a self-contained XSL-FO document.
# Supports: colspec widths, thead/tbody, colspan (namest/nameend), rowspan (morerows),
#           align, valign, verify="ok"/"unconfirmed" colouring, <para> children.
# Style tokens (__FO_*__) are replaced at render time from _CALS_FO_THEMES.
_CALS_TO_FO_XSLT = """\
<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
  xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
  xmlns:fo="http://www.w3.org/1999/XSL/Format">

  <xsl:output method="xml" indent="yes"/>

  <!-- ═══ page wrapper ═══════════════════════════════════════════════════ -->
  <xsl:template match="/table">
    <fo:root>
      <fo:layout-master-set>
        <fo:simple-page-master master-name="p"
            page-width="330mm" page-height="420mm"
            margin-top="8mm" margin-bottom="8mm"
            margin-left="8mm" margin-right="8mm">
          <fo:region-body/>
        </fo:simple-page-master>
      </fo:layout-master-set>
      <fo:page-sequence master-reference="p">
        <fo:flow flow-name="xsl-region-body">
          <!-- optional caption from <title> -->
          <xsl:if test="title">
            <fo:block font-family="__FO_CAPTION_FONT__" font-size="__FO_CAPTION_SIZE__"
                      font-weight="bold" space-after="4pt">
              <xsl:value-of select="title"/>
            </fo:block>
          </xsl:if>
          <xsl:apply-templates select="tgroup"/>
        </fo:flow>
      </fo:page-sequence>
    </fo:root>
  </xsl:template>

  <!-- ═══ tgroup ══════════════════════════════════════════════════════════ -->
  <xsl:template match="tgroup">
    <fo:table table-layout="fixed" width="100%"
              border-collapse="collapse"
              border="__FO_TABLE_BORDER__">
      <!-- colspec → fo:table-column with proportional widths -->
      <xsl:for-each select="colspec">
        <xsl:variable name="w">
          <xsl:choose>
            <!-- CALS proportional: N* → proportional-column-width(N) -->
            <xsl:when test="@colwidth and substring(@colwidth, string-length(@colwidth)) = '*'">
              <xsl:text>proportional-column-width(</xsl:text>
              <xsl:value-of select="substring(@colwidth, 1, string-length(@colwidth) - 1)"/>
              <xsl:text>)</xsl:text>
            </xsl:when>
            <!-- absolute dimension (mm, cm, pt, in …) — pass through -->
            <xsl:when test="@colwidth"><xsl:value-of select="@colwidth"/></xsl:when>
            <xsl:otherwise>proportional-column-width(1)</xsl:otherwise>
          </xsl:choose>
        </xsl:variable>
        <fo:table-column column-width="{$w}"/>
      </xsl:for-each>
      <!-- fallback: equal columns if no colspec -->
      <xsl:if test="not(colspec)">
        <xsl:call-template name="equal-cols">
          <xsl:with-param name="n" select="@cols"/>
        </xsl:call-template>
      </xsl:if>
      <xsl:if test="thead">
        <fo:table-header>
          <xsl:apply-templates select="thead/row" mode="head"/>
        </fo:table-header>
      </xsl:if>
      <fo:table-body>
        <xsl:apply-templates select="tbody/row" mode="body"/>
      </fo:table-body>
    </fo:table>
  </xsl:template>

  <!-- equal-columns fallback -->
  <xsl:template name="equal-cols">
    <xsl:param name="n" select="1"/>
    <xsl:param name="i" select="1"/>
    <xsl:if test="$i &lt;= $n">
      <fo:table-column column-width="proportional-column-width(1)"/>
      <xsl:call-template name="equal-cols">
        <xsl:with-param name="n" select="$n"/>
        <xsl:with-param name="i" select="$i + 1"/>
      </xsl:call-template>
    </xsl:if>
  </xsl:template>

  <!-- ═══ rows ════════════════════════════════════════════════════════════ -->
  <xsl:template match="row" mode="head">
    <fo:table-row>
      <xsl:apply-templates select="entry" mode="head"/>
    </fo:table-row>
  </xsl:template>

  <xsl:template match="row" mode="body">
    <fo:table-row>
      <xsl:apply-templates select="entry" mode="body"/>
    </fo:table-row>
  </xsl:template>

  <!-- ═══ head cell ════════════════════════════════════════════════════════ -->
  <xsl:template match="entry" mode="head">
    <fo:table-cell border="__FO_HEAD_CELL_BORDER__"
                   background-color="__FO_HEAD_BG__"
                   padding="3pt">
      <xsl:call-template name="cell-spans"/>
      <fo:block font-family="__FO_FONT_FAMILY__" font-size="__FO_FONT_SIZE__" font-weight="__FO_HEAD_FONT_WEIGHT__"
                color="__FO_HEAD_COLOR__">
        <xsl:call-template name="cell-align"/>
        <xsl:call-template name="cell-content"/>
      </fo:block>
    </fo:table-cell>
  </xsl:template>

  <!-- ═══ body cell ════════════════════════════════════════════════════════ -->
  <xsl:template match="entry" mode="body">
    <xsl:variable name="bg">
      <xsl:choose>
        <xsl:when test="@verify='ok'">__FO_BG_OK__</xsl:when>
        <xsl:when test="@verify='unconfirmed'">__FO_BG_BAD__</xsl:when>
        <xsl:otherwise>__FO_BG_DEFAULT__</xsl:otherwise>
      </xsl:choose>
    </xsl:variable>
    <fo:table-cell border="__FO_BODY_CELL_BORDER__"
                   background-color="{$bg}"
                   padding="3pt">
      <xsl:call-template name="cell-spans"/>
      <xsl:call-template name="cell-valign"/>
      <fo:block font-family="__FO_FONT_FAMILY__" font-size="__FO_FONT_SIZE__" color="__FO_BODY_COLOR__">
        <xsl:call-template name="cell-align"/>
        <xsl:call-template name="cell-bold"/>
        <xsl:call-template name="cell-indent"/>
        <xsl:call-template name="cell-content"/>
      </fo:block>
    </fo:table-cell>
  </xsl:template>

  <!-- ═══ helpers ══════════════════════════════════════════════════════════ -->
  <xsl:template name="cell-spans">
    <xsl:if test="@namest and @nameend">
      <!-- derive number-columns-spanned from colname numbers -->
      <xsl:variable name="s" select="substring-after(@namest, 'c')"/>
      <xsl:variable name="e" select="substring-after(@nameend, 'c')"/>
      <xsl:if test="string(number($s)) != 'NaN' and string(number($e)) != 'NaN'">
        <xsl:attribute name="number-columns-spanned">
          <xsl:value-of select="$e - $s + 1"/>
        </xsl:attribute>
      </xsl:if>
    </xsl:if>
    <xsl:if test="@morerows">
      <xsl:attribute name="number-rows-spanned">
        <xsl:value-of select="@morerows + 1"/>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <xsl:template name="cell-align">
    <xsl:if test="@align">
      <xsl:attribute name="text-align">
        <xsl:value-of select="@align"/>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <xsl:template name="cell-valign">
    <xsl:if test="@valign">
      <xsl:attribute name="display-align">
        <xsl:choose>
          <xsl:when test="@valign='top'">before</xsl:when>
          <xsl:when test="@valign='bottom'">after</xsl:when>
          <xsl:otherwise>center</xsl:otherwise>
        </xsl:choose>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <!-- render <para> children as separate fo:block lines; fall back to text() -->
  <xsl:template name="cell-content">
    <xsl:choose>
      <xsl:when test="para">
        <xsl:for-each select="para">
          <xsl:value-of select="."/>
          <xsl:if test="position() != last()">
            <fo:block/>
          </xsl:if>
        </xsl:for-each>
      </xsl:when>
      <xsl:otherwise>
        <xsl:value-of select="."/>
      </xsl:otherwise>
    </xsl:choose>
  </xsl:template>

  <!-- bold from VLM annotation -->
  <xsl:template name="cell-bold">
    <xsl:attribute name="font-weight">
      <xsl:choose>
        <xsl:when test="@bold='true'">bold</xsl:when>
        <xsl:otherwise>normal</xsl:otherwise>
      </xsl:choose>
    </xsl:attribute>
  </xsl:template>

  <!-- indent from VLM annotation (each level = 6 mm) -->
  <xsl:template name="cell-indent">
    <xsl:if test="@indent and @indent != '0'">
      <xsl:attribute name="start-indent">
        <xsl:choose>
          <xsl:when test="@indent='1'">6mm</xsl:when>
          <xsl:when test="@indent='2'">12mm</xsl:when>
          <xsl:when test="@indent='3'">18mm</xsl:when>
          <xsl:otherwise>0mm</xsl:otherwise>
        </xsl:choose>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

</xsl:stylesheet>
"""


# XSLT that converts a full <document> (as produced by reconstruct_document("xml"))
# into a self-contained XSL-FO document for A4 rendering via Apache FOP.
# <para> elements become fo:block paragraphs; <table> elements are rendered as
# CALS tables using the same tgroup/thead/tbody/entry templates as _CALS_TO_FO_XSLT.
# Style tokens (__FO_*__) are replaced at render time from _CALS_FO_THEMES.
_DOCUMENT_TO_FO_XSLT = """\
<?xml version="1.0" encoding="UTF-8"?>
<xsl:stylesheet version="1.0"
  xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
  xmlns:fo="http://www.w3.org/1999/XSL/Format">

  <xsl:output method="xml" indent="yes"/>

  <!-- ═══ page wrapper ═══════════════════════════════════════════════════ -->
  <xsl:template match="/document">
    <fo:root>
      <fo:layout-master-set>
        <fo:simple-page-master master-name="a4"
            page-width="210mm" page-height="297mm"
            margin-top="20mm" margin-bottom="20mm"
            margin-left="20mm" margin-right="20mm">
          <fo:region-body/>
        </fo:simple-page-master>
      </fo:layout-master-set>
      <fo:page-sequence master-reference="a4">
        <fo:flow flow-name="xsl-region-body">
          <xsl:apply-templates/>
        </fo:flow>
      </fo:page-sequence>
    </fo:root>
  </xsl:template>

  <!-- ═══ paragraph ═══════════════════════════════════════════════════════ -->
  <xsl:template match="para">
    <fo:block font-family="__FO_FONT_FAMILY__" font-size="__FO_FONT_SIZE__"
              color="__FO_BODY_COLOR__" space-after="6pt" line-height="1.4">
      <xsl:value-of select="."/>
    </fo:block>
  </xsl:template>

  <!-- ═══ table (CALS root) ════════════════════════════════════════════════ -->
  <xsl:template match="table">
    <fo:block space-before="10pt">
      <xsl:if test="title">
        <fo:block font-family="__FO_CAPTION_FONT__" font-size="__FO_CAPTION_SIZE__"
                  font-weight="bold" space-after="4pt">
          <xsl:value-of select="title"/>
        </fo:block>
      </xsl:if>
      <xsl:apply-templates select="tgroup"/>
    </fo:block>
    <fo:block space-after="12pt"/>
  </xsl:template>

  <!-- ═══ tgroup ══════════════════════════════════════════════════════════ -->
  <xsl:template match="tgroup">
    <fo:table table-layout="fixed" width="100%"
              border-collapse="collapse"
              border="__FO_TABLE_BORDER__">
      <!-- colspec → fo:table-column with proportional widths -->
      <xsl:for-each select="colspec">
        <xsl:variable name="w">
          <xsl:choose>
            <!-- CALS proportional: N* → proportional-column-width(N) -->
            <xsl:when test="@colwidth and substring(@colwidth, string-length(@colwidth)) = '*'">
              <xsl:text>proportional-column-width(</xsl:text>
              <xsl:value-of select="substring(@colwidth, 1, string-length(@colwidth) - 1)"/>
              <xsl:text>)</xsl:text>
            </xsl:when>
            <!-- absolute dimension (mm, cm, pt, in …) — pass through -->
            <xsl:when test="@colwidth"><xsl:value-of select="@colwidth"/></xsl:when>
            <xsl:otherwise>proportional-column-width(1)</xsl:otherwise>
          </xsl:choose>
        </xsl:variable>
        <fo:table-column column-width="{$w}"/>
      </xsl:for-each>
      <!-- fallback: equal columns if no colspec -->
      <xsl:if test="not(colspec)">
        <xsl:call-template name="equal-cols">
          <xsl:with-param name="n" select="@cols"/>
        </xsl:call-template>
      </xsl:if>
      <xsl:if test="thead">
        <fo:table-header>
          <xsl:apply-templates select="thead/row" mode="head"/>
        </fo:table-header>
      </xsl:if>
      <fo:table-body>
        <xsl:apply-templates select="tbody/row" mode="body"/>
      </fo:table-body>
    </fo:table>
  </xsl:template>

  <!-- equal-columns fallback -->
  <xsl:template name="equal-cols">
    <xsl:param name="n" select="1"/>
    <xsl:param name="i" select="1"/>
    <xsl:if test="$i &lt;= $n">
      <fo:table-column column-width="proportional-column-width(1)"/>
      <xsl:call-template name="equal-cols">
        <xsl:with-param name="n" select="$n"/>
        <xsl:with-param name="i" select="$i + 1"/>
      </xsl:call-template>
    </xsl:if>
  </xsl:template>

  <!-- ═══ rows ════════════════════════════════════════════════════════════ -->
  <xsl:template match="row" mode="head">
    <fo:table-row>
      <xsl:apply-templates select="entry" mode="head"/>
    </fo:table-row>
  </xsl:template>

  <xsl:template match="row" mode="body">
    <fo:table-row>
      <xsl:apply-templates select="entry" mode="body"/>
    </fo:table-row>
  </xsl:template>

  <!-- ═══ head cell ════════════════════════════════════════════════════════ -->
  <xsl:template match="entry" mode="head">
    <fo:table-cell border="__FO_HEAD_CELL_BORDER__"
                   background-color="__FO_HEAD_BG__"
                   padding="3pt">
      <xsl:call-template name="cell-spans"/>
      <fo:block font-family="__FO_FONT_FAMILY__" font-size="__FO_FONT_SIZE__" font-weight="__FO_HEAD_FONT_WEIGHT__"
                color="__FO_HEAD_COLOR__">
        <xsl:call-template name="cell-align"/>
        <xsl:call-template name="cell-content"/>
      </fo:block>
    </fo:table-cell>
  </xsl:template>

  <!-- ═══ body cell ════════════════════════════════════════════════════════ -->
  <xsl:template match="entry" mode="body">
    <fo:table-cell border="__FO_BODY_CELL_BORDER__"
                   background-color="__FO_BG_DEFAULT__"
                   padding="3pt">
      <xsl:call-template name="cell-spans"/>
      <xsl:call-template name="cell-valign"/>
      <fo:block font-family="__FO_FONT_FAMILY__" font-size="__FO_FONT_SIZE__" color="__FO_BODY_COLOR__">
        <xsl:call-template name="cell-align"/>
        <xsl:call-template name="cell-bold"/>
        <xsl:call-template name="cell-indent"/>
        <xsl:call-template name="cell-content"/>
      </fo:block>
    </fo:table-cell>
  </xsl:template>

  <!-- ═══ helpers ══════════════════════════════════════════════════════════ -->
  <xsl:template name="cell-spans">
    <xsl:if test="@namest and @nameend">
      <xsl:variable name="s" select="substring-after(@namest, 'c')"/>
      <xsl:variable name="e" select="substring-after(@nameend, 'c')"/>
      <xsl:if test="string(number($s)) != 'NaN' and string(number($e)) != 'NaN'">
        <xsl:attribute name="number-columns-spanned">
          <xsl:value-of select="$e - $s + 1"/>
        </xsl:attribute>
      </xsl:if>
    </xsl:if>
    <xsl:if test="@morerows">
      <xsl:attribute name="number-rows-spanned">
        <xsl:value-of select="@morerows + 1"/>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <xsl:template name="cell-align">
    <xsl:if test="@align">
      <xsl:attribute name="text-align">
        <xsl:value-of select="@align"/>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <xsl:template name="cell-valign">
    <xsl:if test="@valign">
      <xsl:attribute name="display-align">
        <xsl:choose>
          <xsl:when test="@valign='top'">before</xsl:when>
          <xsl:when test="@valign='bottom'">after</xsl:when>
          <xsl:otherwise>center</xsl:otherwise>
        </xsl:choose>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

  <xsl:template name="cell-content">
    <xsl:choose>
      <xsl:when test="para">
        <xsl:for-each select="para">
          <xsl:value-of select="."/>
          <xsl:if test="position() != last()">
            <fo:block/>
          </xsl:if>
        </xsl:for-each>
      </xsl:when>
      <xsl:otherwise>
        <xsl:value-of select="."/>
      </xsl:otherwise>
    </xsl:choose>
  </xsl:template>

  <xsl:template name="cell-bold">
    <xsl:attribute name="font-weight">
      <xsl:choose>
        <xsl:when test="@bold='true'">bold</xsl:when>
        <xsl:otherwise>normal</xsl:otherwise>
      </xsl:choose>
    </xsl:attribute>
  </xsl:template>

  <xsl:template name="cell-indent">
    <xsl:if test="@indent and @indent != '0'">
      <xsl:attribute name="start-indent">
        <xsl:choose>
          <xsl:when test="@indent='1'">6mm</xsl:when>
          <xsl:when test="@indent='2'">12mm</xsl:when>
          <xsl:when test="@indent='3'">18mm</xsl:when>
          <xsl:otherwise>0mm</xsl:otherwise>
        </xsl:choose>
      </xsl:attribute>
    </xsl:if>
  </xsl:template>

</xsl:stylesheet>
"""


def _sanitize_rowspans(root) -> None:
    """Clamp entry/@morerows so no cell claims to span beyond its section's row count.

    FOP raises ValidationException when number-rows-spanned exceeds the number of
    rows remaining in the enclosing fo:table-header or fo:table-body.  CALS allows
    morerows values that browsers silently truncate, so we fix them in-place on the
    ElementTree before writing to disk.

    Operates on any root element that contains <tgroup> descendants (e.g. <table>
    or <document>).
    """
    import xml.etree.ElementTree as _ET
    for tgroup in root.iter("tgroup"):
        for section_tag in ("thead", "tbody"):
            section = tgroup.find(section_tag)
            if section is None:
                continue
            rows = section.findall("row")
            total_rows = len(rows)
            for row_idx, row in enumerate(rows):
                for entry in row.findall("entry"):
                    mr = entry.get("morerows")
                    if mr is None:
                        continue
                    try:
                        mr_int = int(mr)
                    except ValueError:
                        entry.attrib.pop("morerows", None)
                        continue
                    # Maximum additional rows this cell can span = rows remaining after this one
                    max_mr = total_rows - 1 - row_idx
                    if mr_int > max_mr:
                        if max_mr <= 0:
                            entry.attrib.pop("morerows", None)
                        else:
                            entry.set("morerows", str(max_mr))


def _cals_to_fop_pdf(cals_xml: str, page_width_mm: float = None, page_height_mm: float = None,
                     theme: str = "verify") -> tuple:
    """Convert a CALS XML table string to a PDF via Apache FOP.

    Returns (pdf_bytes, error_message).  On success, error_message is ''.
    On failure, pdf_bytes is None and error_message describes the problem.

    page_width_mm / page_height_mm: override the default 330×420mm page size
    to match the source Word PDF dimensions.
    theme: one of the keys in _CALS_FO_THEMES; controls XSL-FO typography and colours.

    Requires:
      - fop  on PATH  (apt install fop)
      - Java 11+  (JAVA_HOME pointing at java-11-openjdk if default is Java 8)
    """
    import subprocess, tempfile, base64

    # FOP 2.8 requires Java 9+ due to CharBuffer API change.
    # Detect Java 11 and override JAVA_HOME if the default is Java 8.
    java11 = "/usr/lib/jvm/java-11-openjdk-arm64"
    java11_x86 = "/usr/lib/jvm/java-11-openjdk-amd64"
    import os as _os
    env = _os.environ.copy()
    if _os.path.isdir(java11):
        env["JAVA_HOME"] = java11
    elif _os.path.isdir(java11_x86):
        env["JAVA_HOME"] = java11_x86

    with tempfile.TemporaryDirectory() as tmp:
        xslt_path = _os.path.join(tmp, "cals2fo.xsl")
        xml_path  = _os.path.join(tmp, "input.xml")
        fo_path   = _os.path.join(tmp, "output.fo")
        pdf_path  = _os.path.join(tmp, "output.pdf")

        # Write XSLT — substitute page dimensions and theme tokens
        xslt_str = _CALS_TO_FO_XSLT
        if page_width_mm is not None and page_height_mm is not None:
            w = f"{page_width_mm:.1f}mm"
            h = f"{page_height_mm:.1f}mm"
            xslt_str = xslt_str.replace('page-width="330mm"', f'page-width="{w}"')
            xslt_str = xslt_str.replace('page-height="420mm"', f'page-height="{h}"')
        fo = _CALS_FO_THEMES.get(theme, _CALS_FO_THEMES["verify"])
        xslt_str = (xslt_str
            .replace("__FO_TABLE_BORDER__",     fo["table_border"])
            .replace("__FO_HEAD_CELL_BORDER__",  fo["head_cell_border"])
            .replace("__FO_HEAD_BG__",           fo["head_bg"])
            .replace("__FO_FONT_FAMILY__",       fo["font_family"])
            .replace("__FO_FONT_SIZE__",         fo["font_size"])
            .replace("__FO_HEAD_FONT_WEIGHT__",  fo["head_font_weight"])
            .replace("__FO_HEAD_COLOR__",        fo["head_color"])
            .replace("__FO_BODY_CELL_BORDER__",  fo["body_cell_border"])
            .replace("__FO_BODY_COLOR__",        fo["body_color"])
            .replace("__FO_BG_OK__",             fo["bg_ok"])
            .replace("__FO_BG_BAD__",            fo["bg_bad"])
            .replace("__FO_BG_DEFAULT__",        fo["bg_default"])
            .replace("__FO_CAPTION_FONT__",      fo["caption_font"])
            .replace("__FO_CAPTION_SIZE__",      fo["caption_size"])
        )
        with open(xslt_path, "w", encoding="utf-8") as f:
            f.write(xslt_str)

        # Normalise: ensure root element is <table>; sanitize rowspans for FOP
        xml_str = cals_xml.strip()
        if not xml_str.startswith("<table"):
            xml_str = f"<table>{xml_str}</table>"
        try:
            import xml.etree.ElementTree as _ET2
            _root_t = _ET2.fromstring(xml_str)
            _sanitize_rowspans(_root_t)
            xml_str = _ET2.tostring(_root_t, encoding="unicode")
        except Exception:
            pass
        with open(xml_path, "w", encoding="utf-8") as f:
            f.write('<?xml version="1.0" encoding="UTF-8"?>\n' + xml_str)

        # Step 1: XSLT transform (CALS → XSL-FO) using Java's built-in XSLT
        xslt_cmd = [
            env.get("JAVA_HOME", "/usr") + "/bin/java",
            "-cp", "/usr/share/java/xalan2.jar:/usr/share/java/serializer.jar"
                   ":/usr/share/java/xercesImpl.jar:/usr/share/java/xml-apis.jar",
            "org.apache.xalan.xslt.Process",
            "-IN", xml_path, "-XSL", xslt_path, "-OUT", fo_path,
        ]
        # Fall back to Saxon or the JDK's built-in Xalan if xalan2.jar absent
        xalan_jar = "/usr/share/java/xalan2.jar"
        if not _os.path.exists(xalan_jar):
            # Use JDK built-in XSLT (javax.xml.transform) via a tiny wrapper isn't
            # easily invocable from CLI — instead pass XSL directly to fop -xsl
            fo_path = None  # signal to use fop -xsl path below

        if fo_path is not None:
            r = subprocess.run(xslt_cmd, capture_output=True, text=True, env=env, timeout=30)
            if r.returncode != 0 or not _os.path.exists(fo_path):
                # Fall through to fop -xsl
                fo_path = None

        # Step 2: FOP render
        if fo_path and _os.path.exists(fo_path):
            fop_cmd = ["fop", fo_path, pdf_path]
        else:
            # Let FOP do the XSLT itself (fop -xsl ... -xml ... -pdf ...)
            fop_cmd = ["fop", "-xsl", xslt_path, "-xml", xml_path, "-pdf", pdf_path]

        r2 = subprocess.run(fop_cmd, capture_output=True, text=True, env=env, timeout=60)
        stderr = r2.stderr or ""
        # Filter noisy warnings that aren't errors
        errors = [l for l in stderr.splitlines()
                  if "ERROR" in l or "Exception" in l or "FATAL" in l]

        if not _os.path.exists(pdf_path) or _os.path.getsize(pdf_path) < 100:
            return None, "FOP failed:\n" + "\n".join(errors) or stderr[:500]

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    b64 = base64.b64encode(pdf_bytes).decode()
    return b64, ""


def _fop_vs_word_diff_image(fop_b64: str, word_img_path: str, out_path: str = None) -> str:
    """Create a side-by-side PNG comparing the FOP-rendered page and the Word→PDF page.

    Parameters
    ----------
    fop_b64       : base64-encoded PDF bytes from _cals_to_fop_pdf
    word_img_path : path to the existing Word→PDF page PNG
    out_path      : where to save the composite PNG (auto-generated if None)

    Returns the path to the composite PNG, or None on failure.
    """
    import base64 as _b64, tempfile, io
    try:
        from PIL import Image, ImageDraw, ImageFont
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        print(f"[fop_vs_word_diff] missing dependency: {exc}")
        return None

    try:
        pdf_bytes = _b64.b64decode(fop_b64)
        fop_pages = convert_from_bytes(pdf_bytes, dpi=150)
        if not fop_pages:
            return None
        fop_img = fop_pages[0].convert("RGB")
    except Exception as exc:
        print(f"[fop_vs_word_diff] FOP page render failed: {exc}")
        return None

    if not word_img_path or not os.path.exists(word_img_path):
        # No Word page available — save FOP page alone with unique timestamp name
        import time as _time
        ts = int(_time.time() * 1000)
        out_path = out_path or f"/tmp/fop_vs_word_{ts}.png"
        fop_img.save(out_path, "PNG")
        return out_path

    try:
        word_img = Image.open(word_img_path).convert("RGB")
    except Exception as exc:
        print(f"[fop_vs_word_diff] Word image load failed: {exc}")
        return None

    # Scale both images to the same height for comparison
    TARGET_H = max(fop_img.height, word_img.height)
    def _scale_to_height(img, h):
        if img.height == h:
            return img
        ratio = h / img.height
        return img.resize((int(img.width * ratio), h), Image.LANCZOS)

    fop_img  = _scale_to_height(fop_img,  TARGET_H)
    word_img = _scale_to_height(word_img, TARGET_H)

    DIVIDER   = 6
    LABEL_H   = 28
    FONT_SIZE = 14
    total_w = fop_img.width + DIVIDER + word_img.width
    total_h = TARGET_H + LABEL_H

    composite = Image.new("RGB", (total_w, total_h), (240, 240, 240))

    # Labels
    draw = ImageDraw.Draw(composite)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", FONT_SIZE)
    except Exception:
        font = ImageFont.load_default()

    draw.rectangle([0, 0, fop_img.width, LABEL_H - 1], fill=(220, 235, 255))
    draw.text((8, 6), "FOP Render  (green=confirmed  red=unconfirmed)", fill=(0, 60, 160), font=font)

    word_x = fop_img.width + DIVIDER
    draw.rectangle([word_x, 0, total_w - 1, LABEL_H - 1], fill=(255, 235, 210))
    draw.text((word_x + 8, 6), "Original Word → PDF  (source of truth)", fill=(160, 60, 0), font=font)

    # Divider
    draw.rectangle([fop_img.width, 0, fop_img.width + DIVIDER - 1, total_h], fill=(100, 100, 100))

    # Paste images
    composite.paste(fop_img,  (0, LABEL_H))
    composite.paste(word_img, (word_x, LABEL_H))

    if out_path is None:
        import time as _time
        ts = int(_time.time() * 1000)  # millisecond timestamp → always a unique filename
        base_dir = os.path.dirname(word_img_path) if word_img_path else "/tmp"
        out_path = os.path.join(base_dir, f"fop_vs_word_{ts}.png")

    composite.save(out_path, "PNG")
    print(f"[fop_vs_word_diff] saved → {out_path}")
    return out_path


def _fop_pdf_to_html(cals_xml: str) -> str:
    """Return an HTML string that embeds a FOP-rendered PDF as an <iframe>.

    On error, returns an error message panel instead.
    """
    b64, err = _cals_to_fop_pdf(cals_xml)
    if err:
        return (
            "<div style='font-family:monospace;font-size:11px;color:#c00;"
            "background:#fff8f8;padding:10px;border:1px solid #f00;"
            "border-radius:4px;white-space:pre-wrap;'>"
            f"<b>FOP error:</b>\n{err}</div>"
        )
    return (
        f'<iframe src="data:application/pdf;base64,{b64}" '
        'style="width:100%;height:700px;border:1px solid #ccc;border-radius:4px;" '
        'title="FOP-rendered PDF preview"></iframe>'
    )


def _recheck_xml(cals_xml: str, table_meta: dict) -> tuple:
    """Re-run verify_table for a single edited table and return updated artefacts.

    Parameters
    ----------
    cals_xml   : edited CALS XML string (plain, without verify= attributes)
    table_meta : dict with keys pdf_path, page_idx, image_path, title

    Returns
    -------
    (annotated_xml, new_diff_img_path)
      annotated_xml    : CALS XML with verify="ok"/"unconfirmed" attributes baked in
                         (falls back to plain cals_xml on error)
      new_diff_img_path: path to re-annotated page image, or None
    """
    pdf_path  = table_meta.get("pdf_path", "")
    page_idx  = table_meta.get("page_idx", -1)
    img_path  = table_meta.get("image_path", "")
    title     = table_meta.get("title", "")

    if not pdf_path or not os.path.exists(pdf_path) or page_idx < 0:
        return cals_xml, None

    # Strip existing verify= attributes so re-verify starts clean
    import re as _re, xml.etree.ElementTree as ET
    try:
        root = ET.fromstring(cals_xml)
        for entry in root.iter("entry"):
            entry.attrib.pop("verify", None)
        _indent_xml(root)
        clean_xml = ET.tostring(root, encoding="unicode")
    except ET.ParseError:
        clean_xml = cals_xml

    result = verify_table(clean_xml, pdf_path, page_idx, title, img_path)
    annotated_xml      = result.get("annotated_xml") or clean_xml
    new_diff_img_path  = result.get("annotated_image_path") or None
    return annotated_xml, new_diff_img_path


def save_table_xml(all_tables: list, selected_idx: int, new_xml: str) -> tuple:
    """Persist an edited CALS XML back to every storage location.

    Updates:
      - all_tables[selected_idx]["xml"]   (in-memory catalog)
      - data/table_catalog.json           (on-disk catalog)
      - data/doc_segments.json            (doc sequence, if table_index matches)

    Returns (updated_all_tables, status_message).
    """
    import json as _json

    if not all_tables or selected_idx is None or selected_idx >= len(all_tables):
        return all_tables, "❌ No table selected."
    if not new_xml or not new_xml.strip():
        return all_tables, "❌ XML is empty — not saved."

    # Validate XML before saving
    try:
        import xml.etree.ElementTree as ET
        ET.fromstring(new_xml)
    except ET.ParseError as exc:
        return all_tables, f"❌ XML parse error: {exc}"

    # 1. Update in-memory list
    import copy
    updated = copy.deepcopy(all_tables)
    updated[selected_idx]["xml"] = new_xml
    table_index = updated[selected_idx].get("table_index")

    # 2. Overwrite table_catalog.json
    catalog_path = os.path.join(DATA_DIR, "table_catalog.json")
    try:
        if os.path.exists(catalog_path):
            with open(catalog_path, "r", encoding="utf-8") as _f:
                catalog = _json.load(_f)
            for entry in catalog:
                if entry.get("table_index") == table_index:
                    entry["xml"] = new_xml
            with open(catalog_path, "w", encoding="utf-8") as _f:
                _json.dump(catalog, _f, ensure_ascii=False, indent=2)
    except Exception as exc:
        return updated, f"⚠️ Saved in memory but table_catalog.json failed: {exc}"

    # 3. Overwrite doc_segments.json
    segs_path = os.path.join(DATA_DIR, "doc_segments.json")
    try:
        if os.path.exists(segs_path):
            with open(segs_path, "r", encoding="utf-8") as _f:
                segs = _json.load(_f)
            changed = False
            for seg in segs:
                if seg.get("type") == "table" and seg.get("table_index") == table_index:
                    seg["xml"] = new_xml
                    changed = True
            if changed:
                with open(segs_path, "w", encoding="utf-8") as _f:
                    _json.dump(segs, _f, ensure_ascii=False, indent=2)
    except Exception as exc:
        return updated, f"⚠️ Saved table_catalog.json but doc_segments.json failed: {exc}"

    title = updated[selected_idx].get("title", f"table {table_index}")
    return updated, f"✅ Saved — '{title[:50]}'"


def update_entry_text(all_tables: list, selected_idx: int, eid: str, new_text: str) -> tuple:
    """Update a single <entry> cell text in the clean XML and persist to disk.

    Finds the entry by structural position (same algorithm as _assign_ids so
    the eid from the annotation span maps 1-to-1).  Strips verify= attributes
    before saving — they are display-only and regenerated on next verification.

    Returns (updated_all_tables, new_clean_xml, status_message).
    """
    import copy
    import xml.etree.ElementTree as ET

    if not all_tables or selected_idx is None or selected_idx >= len(all_tables):
        return all_tables, "", "❌ No table selected."

    clean_xml = all_tables[selected_idx].get("xml", "")
    if not clean_xml:
        return all_tables, "", "❌ No XML for this table."

    try:
        root = ET.fromstring(clean_xml)
    except ET.ParseError as exc:
        return all_tables, clean_xml, f"❌ XML parse error: {exc}"

    # Find entry by replicating _assign_ids traversal order
    target = None
    row_idx = 0
    outer_break = False
    for tg in root.findall("tgroup"):
        if outer_break:
            break
        for section in (tg.find("thead"), tg.find("tbody")):
            if section is None:
                continue
            for row in section.findall("row"):
                for entry in row.findall("entry"):
                    col_key = (
                        entry.get("colname")
                        or entry.get("namest")
                        or f"x{row_idx}"
                    )
                    if f"entry-r{row_idx}-{col_key}" == eid:
                        target = entry
                        outer_break = True
                        break
                if outer_break:
                    break
                row_idx += 1
            if outer_break:
                break

    if target is None:
        return all_tables, clean_xml, f"❌ Entry '{eid}' not found in XML."

    # Replace content — remove child elements, set plain text
    for child in list(target):
        target.remove(child)
    target.text = new_text
    # Strip display-only verify attrs (regenerated on next verification)
    target.attrib.pop("verify", None)
    target.attrib.pop("verify-reason", None)

    _indent_xml(root)
    new_clean_xml = ET.tostring(root, encoding="unicode")

    updated, msg = save_table_xml(all_tables, selected_idx, new_clean_xml)
    return updated, new_clean_xml, msg
    """Level 2: Convert a .docx to per-page PNG images via LibreOffice + pdf2image.

    Returns a list of absolute image paths, one per page (0-indexed).
    """
    import subprocess
    from pdf2image import convert_from_path

    os.makedirs(img_dir, exist_ok=True)

    # Convert docx → PDF using LibreOffice headless
    pdf_path = os.path.join(img_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", img_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            print(f"[render_pages] LibreOffice failed: {result.stderr}")
            return [], pdf_path
    except Exception as e:
        print(f"[render_pages] LibreOffice error: {e}")
        return [], pdf_path

    if not os.path.exists(pdf_path):
        print(f"[render_pages] PDF not found after conversion: {pdf_path}")
        return [], pdf_path

    # Convert PDF pages → PNG images
    try:
        pages = convert_from_path(pdf_path, dpi=dpi)
    except Exception as e:
        print(f"[render_pages] pdf2image error: {e}")
        return [], pdf_path

    page_paths = []
    for i, page in enumerate(pages):
        page_path = os.path.join(img_dir, f"page_{i:03d}.png")
        page.save(page_path, "PNG")
        page_paths.append(page_path)

    print(f"[render_pages] Rendered {len(page_paths)} pages to {img_dir}")
    return page_paths, pdf_path


def _build_pdf_page_cache(pdf_path: str) -> dict:
    """Open *pdf_path* once and pre-compute per-page word lookup structures for
    `_annotate_entry_styles_with_pdfplumber`.

    Returns a dict keyed by page index (int):
        {page_idx: (line_lookup, word_lookup, x0_min, prefix_index)}

    line_lookup    : normalised line text -> (x0_first, first_bold, any_bold)
    word_lookup    : single normalised word -> is_bold
    x0_min         : leftmost x0 on page (flush-left margin baseline)
    prefix_index   : maps every 1-6 token prefix of each line to that line's
                     tuple — allows O(1) lookup instead of O(n) scan.
    """
    try:
        import pdfplumber
    except ImportError:
        return {}

    from collections import defaultdict

    cache: dict = {}
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                words = page.extract_words(extra_attrs=["fontname", "size"])
                if not words:
                    cache[page_idx] = ({}, {}, 0.0, {})
                    continue

                y_to_line: dict = defaultdict(list)
                for w in words:
                    y_to_line[round(w["top"])].append(w)

                line_lookup: dict = {}
                prefix_index: dict = {}
                for line_words in y_to_line.values():
                    line_words.sort(key=lambda w: w["x0"])
                    key = " ".join(w["text"] for w in line_words).strip().lower()
                    key = " ".join(key.split())
                    val = (
                        line_words[0]["x0"],
                        "Bold" in line_words[0]["fontname"],
                        any("Bold" in w["fontname"] for w in line_words),
                    )
                    line_lookup[key] = val
                    # Index every 1-6 token prefix
                    tokens = key.split()
                    for length in range(1, min(7, len(tokens) + 1)):
                        p = " ".join(tokens[:length])
                        if p not in prefix_index:
                            prefix_index[p] = val  # first/shortest match wins

                word_lookup: dict = {}
                for w in words:
                    word_lookup[" ".join(w["text"].split()).lower()] = "Bold" in w["fontname"]

                x0_min = min(w["x0"] for w in words)
                cache[page_idx] = (line_lookup, word_lookup, x0_min, prefix_index)
    except Exception as exc:
        print(f"[pdf-cache] Failed to build cache for {pdf_path}: {exc}")

    return cache


def _annotate_entry_styles_with_pdfplumber(
    cals_xml: str,
    pdf_path: str,
    page_idx: int,
    _page_cache: dict = None,
) -> str:
    """Annotate CALS <entry> bold/indent attributes using font metadata extracted
    directly from the LibreOffice-generated PDF via pdfplumber.

    Bold detection  — an entry whose text starts with a word rendered in a font
                      containing 'Bold' (e.g. 'LiberationSans-Bold') gets bold="true".
    Indent detection — label-column entries (first colspec column) get an indent
                      level derived from their x0 position relative to the minimum
                      x0 on the page.  Each ~14 pt step maps to one indent level (0–3).
    Value columns    — bold is set by font; indent is never written on non-label
                      columns (numeric columns must not be left-padded).

    Returns the annotated XML string, or *cals_xml* unchanged on any error.
    """
    import xml.etree.ElementTree as ET

    if not pdf_path or not os.path.exists(pdf_path):
        print("[pdfplumber-annotate] PDF not found, skipping")
        return cals_xml

    try:
        import pdfplumber
    except ImportError:
        print("[pdfplumber-annotate] pdfplumber not installed, skipping")
        return cals_xml

    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError as exc:
        print(f"[pdfplumber-annotate] XML parse error: {exc}")
        return cals_xml

    # Use pre-built cache if supplied; otherwise build it on the fly (one-table path)
    if _page_cache is not None:
        page_data = _page_cache.get(page_idx)
    else:
        page_data = None

    if page_data is None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                if page_idx >= len(pdf.pages):
                    print(f"[pdfplumber-annotate] page_idx {page_idx} out of range")
                    return cals_xml
                page = pdf.pages[page_idx]
                words = page.extract_words(extra_attrs=["fontname", "size"])
        except Exception as exc:
            print(f"[pdfplumber-annotate] pdfplumber error: {exc}")
            return cals_xml

        if not words:
            return cals_xml

        from collections import defaultdict
        y_to_line: dict = defaultdict(list)
        for w in words:
            y_to_line[round(w["top"])].append(w)

        line_lookup: dict = {}
        prefix_index: dict = {}
        for line_words in y_to_line.values():
            line_words.sort(key=lambda w: w["x0"])
            key = " ".join(" ".join(w["text"] for w in line_words).strip().split()).lower()
            val = (
                line_words[0]["x0"],
                "Bold" in line_words[0]["fontname"],
                any("Bold" in w["fontname"] for w in line_words),
            )
            line_lookup[key] = val
            tokens = key.split()
            for length in range(1, min(7, len(tokens) + 1)):
                p = " ".join(tokens[:length])
                if p not in prefix_index:
                    prefix_index[p] = val

        word_lookup: dict = {}
        for w in words:
            word_lookup[" ".join(w["text"].split()).lower()] = "Bold" in w["fontname"]

        x0_min = min(w["x0"] for w in words)
    else:
        line_lookup, word_lookup, x0_min, prefix_index = page_data
    # One indent step ≈ 14 pt (LiberationSans 10 pt rendered at 150 dpi)
    INDENT_STEP = 14.0

    def _indent_level(x0: float) -> int:
        return min(3, max(0, round((x0 - x0_min) / INDENT_STEP)))

    # Identify label column names (first colspec per tgroup)
    label_colnames: set = set()
    for tg in root.findall("tgroup"):
        cols = tg.findall("colspec")
        if cols:
            label_colnames.add(cols[0].get("colname", ""))

    applied = 0
    for tg in root.findall("tgroup"):
        for section in ("thead", "tbody"):
            for tbody in tg.findall(section):
                for row in tbody.findall("row"):
                    for entry in row.findall("entry"):
                        text = " ".join("".join(entry.itertext()).split())
                        if not text:
                            continue

                        colname = entry.get("colname", entry.get("namest", ""))
                        is_label_col = colname in label_colnames

                        # Match entry text against PDF lines using progressively
                        # shorter prefixes (O(1) via prefix_index; avoids O(n) scan).
                        tokens = text.lower().split()
                        matched_x0 = None
                        matched_first_bold = False
                        matched_any_bold = False
                        for length in range(min(6, len(tokens)), 0, -1):
                            prefix = " ".join(tokens[:length])
                            if prefix in prefix_index:
                                matched_x0, matched_first_bold, matched_any_bold = prefix_index[prefix]
                                break

                        # Bold: label column uses first-word bold only (avoids value
                        # column contamination); value columns use word_lookup on the
                        # exact cell text (e.g. bare number like "77,673").
                        if is_label_col:
                            is_bold = matched_first_bold
                        else:
                            # Prefer exact word match, fall back to any_bold on line
                            is_bold = word_lookup.get(text.lower(), matched_any_bold)

                        if is_bold:
                            entry.set("bold", "true")
                        else:
                            entry.attrib.pop("bold", None)

                        if is_label_col and matched_x0 is not None:
                            lvl = _indent_level(matched_x0)
                            if lvl > 0:
                                entry.set("indent", str(lvl))
                            else:
                                entry.attrib.pop("indent", None)
                        else:
                            # Value columns: never indent
                            entry.attrib.pop("indent", None)

                        applied += 1

    print(f"[pdfplumber-annotate] Annotated {applied} entries on page {page_idx}")
    return ET.tostring(root, encoding="unicode")


def _annotate_entry_styles_with_vlm(
    cals_xml: str,
    img_path: str,
    lmstudio_url: str = "http://localhost:1234/v1",
    batch_size: int = 25,
    max_img_px: int = 512,
) -> str:
    """Query a vision LLM (Qwen2.5-VL via LM Studio) to annotate bold/indent
    styles on CALS <entry> elements in *cals_xml*.  Only body (non-header) rows
    are annotated.

    The page image is resized to at most *max_img_px* on its longest side and
    re-encoded as JPEG to keep token count low.  Cell lists longer than
    *batch_size* are split across multiple requests (each re-uses the same
    image) so the context window is never exceeded.

    Returns the input *cals_xml* unchanged on any error.
    """
    import base64, io, json, re
    import xml.etree.ElementTree as ET

    if not img_path or not os.path.exists(img_path):
        return cals_xml

    try:
        from openai import OpenAI
    except ImportError:
        print("[VLM] openai package not installed — skipping style annotation")
        return cals_xml

    # --- resize image and encode as base64 data URI (cached per path+size) ---
    _cache_key = (img_path, max_img_px)
    if _cache_key in _vlm_img_b64_cache:
        img_data_url = _vlm_img_b64_cache[_cache_key]
        print(f"[VLM] Image data URI served from cache for {os.path.basename(img_path)}")
    else:
        try:
            from PIL import Image as _PILImage
            with _PILImage.open(img_path) as _im:
                _im = _im.convert("RGB")
                w, h = _im.size
                scale = min(1.0, max_img_px / max(w, h))
                nw, nh = int(w * scale), int(h * scale)
                if scale < 1.0:
                    _im = _im.resize((nw, nh), _PILImage.LANCZOS)
                buf = io.BytesIO()
                _im.save(buf, format="PNG")
                img_b64 = base64.b64encode(buf.getvalue()).decode()
            img_data_url = f"data:image/png;base64,{img_b64}"
            print(f"[VLM] Image resized to {nw}x{nh} PNG ({len(buf.getvalue())//1024} KB)")
        except Exception as exc:
            print(f"[VLM] Image resize failed ({exc}), falling back to raw file")
            mime = mimetypes.guess_type(img_path)[0] or "image/png"
            with open(img_path, "rb") as fh:
                img_b64 = base64.b64encode(fh.read()).decode()
            img_data_url = f"data:{mime};base64,{img_b64}"
        _vlm_img_b64_cache[_cache_key] = img_data_url

    # --- detect which model is currently loaded in LM Studio (cached per URL) ---
    client = OpenAI(base_url=lmstudio_url, api_key="lm-studio")
    if lmstudio_url in _vlm_model_id_cache:
        model_id = _vlm_model_id_cache[lmstudio_url]
    else:
        try:
            models = client.models.list()
            model_id = models.data[0].id if models.data else "local-model"
            print(f"[VLM] Using model: {model_id}")
        except Exception:
            model_id = "local-model"
        _vlm_model_id_cache[lmstudio_url] = model_id

    # --- parse CALS XML and collect body entry (element, text) pairs ---
    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError as exc:
        print(f"[VLM] XML parse error: {exc}")
        return cals_xml

    body_entries = []   # list of (entry_element, stripped_text)
    for tg in root.findall("tgroup"):
        for tbody in tg.findall("tbody"):
            for row in tbody.findall("row"):
                for entry in row.findall("entry"):
                    text = " ".join("".join(entry.itertext()).split())
                    if text:
                        body_entries.append((entry, text))

    if not body_entries:
        return cals_xml

    # --- helper: call VLM for one batch, return annotation list or [] ---
    def _call_vlm(batch):
        entries_json = json.dumps(
            [{"id": item["id"], "text": item["text"]} for item in batch],
            ensure_ascii=False,
        )
        prompt = (
            "You are analyzing a financial table image.\n"
            "Below is a JSON list of body cell text values from the table.\n\n"
            "For each cell determine:\n"
            "  bold   — true if the text is visually bold/heavy, false otherwise\n"
            "  indent — indent level: 0=none, 1=slight, 2=more, 3=deep\n\n"
            "Return ONLY a compact JSON array (same order as input).\n"
            "Format: [{\"id\":<int>,\"bold\":<bool>,\"indent\":<0-3>},...]\n"
            "No explanation, no markdown fences.\n\n"
            f"Cells:\n{entries_json}"
        )
        try:
            resp = client.chat.completions.create(
                model=model_id,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url",
                         "image_url": {"url": img_data_url}},
                        {"type": "text", "text": prompt},
                    ],
                }],
                max_tokens=1024,
                temperature=0.0,
            )
            raw = resp.choices[0].message.content.strip()
            cleaned = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE).rstrip("`").strip()
            result = json.loads(cleaned)
            if not isinstance(result, list):
                raise ValueError("Expected JSON array")
            return result
        except Exception as exc:
            print(f"[VLM] LM Studio call failed: {exc}")
            return []

    # --- chunk entries into batches and collect all annotations ---
    all_items = [{"id": i, "text": e[1]} for i, e in enumerate(body_entries)]
    ann_map = {}
    for start in range(0, len(all_items), batch_size):
        chunk = all_items[start: start + batch_size]
        results = _call_vlm(chunk)
        for a in results:
            if isinstance(a, dict) and "id" in a:
                ann_map[int(a["id"])] = a

    # --- apply annotations back to the XML elements ---
    applied = 0
    for i, (entry_el, _text) in enumerate(body_entries):
        ann = ann_map.get(i)
        if not ann:
            continue
        if ann.get("bold"):
            entry_el.set("bold", "true")
        indent = ann.get("indent", 0)
        if isinstance(indent, int) and indent > 0:
            entry_el.set("indent", str(indent))
        applied += 1

    print(f"[VLM] Style annotation: {applied}/{len(body_entries)} entries annotated")
    return ET.tostring(root, encoding="unicode")


# ---------------------------------------------------------------------------
# Annotation Agent helpers
# ---------------------------------------------------------------------------

def _compare_annotation_sets(xml_a: str, xml_b: str) -> list:
    """Compare two annotated CALS XMLs and return a list of per-entry conflicts.

    Each conflict dict:
        {"text": str,
         "pdfplumber": {"bold": bool, "indent": int},
         "vlm":        {"bold": bool, "indent": int}}
    Only entries that exist in both XMLs and differ on bold or indent are returned.
    """
    import xml.etree.ElementTree as ET

    def _entry_map(xml_str: str) -> dict:
        try:
            root = ET.fromstring(xml_str)
        except ET.ParseError:
            return {}
        result = {}
        for e in root.iter("entry"):
            text = " ".join("".join(e.itertext()).split())
            if text:
                result[text] = {
                    "bold":   e.get("bold") == "true",
                    "indent": int(e.get("indent", 0)),
                }
        return result

    a = _entry_map(xml_a)
    b = _entry_map(xml_b)

    conflicts = []
    for text, aa in a.items():
        bb = b.get(text)
        if bb and (aa["bold"] != bb["bold"] or aa["indent"] != bb["indent"]):
            conflicts.append({"text": text, "pdfplumber": aa, "vlm": bb})
    return conflicts


def _reconcile_with_llm(table_title: str, conflicts: list, llm_url: str) -> list:
    """Ask a local LLM (via OpenAI-compatible API) to reconcile annotation conflicts.

    The LLM receives the conflict list in text form and returns a JSON array of
    final decisions: [{"text": str, "bold": bool, "indent": int}, ...]

    Guidelines injected into the prompt:
    - Trust pdfplumber for bold (reads actual font name from PDF).
    - Trust pdfplumber for indent (reads actual x0 coordinate).
    - Only prefer VLM when there is a clear semantic reason (e.g. a total row
      that should be bold but the PDF lost the style).

    Returns [] on any error; caller falls back to pdfplumber result.
    """
    import json as _json, re

    try:
        from openai import OpenAI
    except ImportError:
        return []

    client = OpenAI(base_url=llm_url, api_key="lm-studio")
    if llm_url in _vlm_model_id_cache:
        model_id = _vlm_model_id_cache[llm_url]
    else:
        try:
            models = client.models.list()
            model_id = models.data[0].id if models.data else "local-model"
        except Exception:
            model_id = "local-model"
        _vlm_model_id_cache[llm_url] = model_id

    # Cap at 30 conflicts to stay within typical context windows
    payload = _json.dumps(conflicts[:30], ensure_ascii=False)

    prompt = (
        f'You are a document annotation expert reviewing a financial table: "{table_title}".\n\n'
        "Two annotation methods disagree on the entries listed below.\n"
        "  pdfplumber — reads the actual PDF font name (contains 'Bold') and x0 position. "
        "This is ground truth for bold and indent.\n"
        "  vlm — reads the visual page image. May capture nuance but can hallucinate.\n\n"
        "Rules:\n"
        "1. Prefer pdfplumber bold unless the cell is clearly a financial total/header that "
        "should be bold but was lost in the PDF conversion.\n"
        "2. Prefer pdfplumber indent (x0-derived) in all cases.\n"
        "3. Never indent non-label (numeric/date) columns.\n\n"
        f"Conflicts:\n{payload}\n\n"
        'Output ONLY a compact JSON array (no markdown, no explanation):\n'
        '[{"text":"<cell text>","bold":<true/false>,"indent":<0-3>},...]'
    )

    try:
        resp = client.chat.completions.create(
            model=model_id,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1024,
            temperature=0.0,
        )
        raw = resp.choices[0].message.content.strip()
        cleaned = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE).rstrip("`").strip()
        result = _json.loads(cleaned)
        if isinstance(result, list):
            return result
    except Exception as exc:
        print(f"[reconcile] LLM call failed: {exc}")
    return []


def _apply_reconciliation(cals_xml: str, decisions: list) -> str:
    """Apply LLM reconciliation decisions to a CALS XML string.

    Each decision: {"text": str, "bold": bool, "indent": int}
    Only entries whose text matches a decision are modified; all others are
    left as-is (preserving the pdfplumber baseline).
    """
    import xml.etree.ElementTree as ET

    if not decisions:
        return cals_xml

    dec_map = {
        d["text"]: d
        for d in decisions
        if isinstance(d, dict) and "text" in d
    }

    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError:
        return cals_xml

    for entry in root.iter("entry"):
        text = " ".join("".join(entry.itertext()).split())
        d = dec_map.get(text)
        if d is None:
            continue
        if d.get("bold"):
            entry.set("bold", "true")
        else:
            entry.attrib.pop("bold", None)
        indent = int(d.get("indent", 0))
        if indent > 0:
            entry.set("indent", str(indent))
        else:
            entry.attrib.pop("indent", None)

    return ET.tostring(root, encoding="unicode")


def _draw_bbox_on_image(img_path: str, bbox, color: str = "blue") -> str:
    """Draw a semi-transparent bounding box on a copy of *img_path* and return
    the path to the annotated copy.  If anything fails the original path is
    returned unchanged.

    Parameters
    ----------
    img_path : str
        Absolute path to the source PNG.
    bbox : list | None
        Normalised [x0, y0, x1, y1] (0‥1 relative to image dimensions) as
        stored in doc_segments.json / table_catalog.json.  ``None`` → no-op.
    color : str
        ``"blue"`` for table segments (41, 128, 185),
        ``"red"``  for paragraph segments (255, 0, 0).
    """
    if not bbox or not img_path or not os.path.exists(img_path):
        return img_path
    try:
        from PIL import Image, ImageDraw
        # Blue matches table run colour; red matches paragraph run colour.
        _COLORS = {
            "blue": ((41,  128, 185,  80), (41,  128, 185, 230)),
            "red":  ((255,   0,   0,  60), (255,   0,   0, 220)),
        }
        fill, outline = _COLORS.get(color, _COLORS["blue"])
        img = Image.open(img_path).convert("RGBA")
        iw, ih = img.size
        x0 = int(bbox[0] * iw)
        y0 = int(bbox[1] * ih)
        x1 = int(bbox[2] * iw)
        y1 = int(bbox[3] * ih)
        overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)
        draw.rectangle([x0, y0, x1, y1], fill=fill, outline=outline, width=3)
        result = Image.alpha_composite(img, overlay).convert("RGB")
        out_path = img_path.replace(".png", "_bbox.png")
        result.save(out_path, "PNG")
        return out_path
    except Exception as _e:
        print(f"[_draw_bbox_on_image] {_e}")
        return img_path


def _build_annotation_review_html(table_entry: dict) -> str:
    """Build an HTML comparison table showing pdfplumber vs all available VLM
    annotations (Qwen, Gemma, or any other model that has been run).

    Detects all ``xml_vlm*`` keys in the catalog entry dynamically so new models
    appear automatically without code changes.  Each VLM source gets its own
    bold / indent / match column group.

    Returns a self-contained HTML string.
    """
    import xml.etree.ElementTree as ET

    def _entry_map(xml_str: str) -> dict:
        """text -> (bold_bool, indent_int)"""
        if not xml_str:
            return {}
        try:
            root = ET.fromstring(xml_str)
        except ET.ParseError:
            return {}
        result: dict = {}
        for e in root.iter("entry"):
            txt = " ".join("".join(e.itertext()).split())
            if txt:
                result[txt] = (
                    e.get("bold") == "true",
                    int(e.get("indent", 0)),
                )
        return result

    def _model_label(key: str) -> str:
        """xml_vlm_qwen25vl -> 'Qwen2.5-VL', xml_vlm_gemma3 -> 'Gemma3', etc."""
        suffix = key[len("xml_vlm"):].lstrip("_")
        if not suffix:
            return "VLM"  # bare legacy key (should not appear any more)
        # Pretty-print known patterns
        _pretty = {
            "qwen25vl": "Qwen2.5-VL",
            "qwen25":   "Qwen2.5",
            "qwen35":   "Qwen3.5",
            "qwen3":    "Qwen3",
            "gemma3":   "Gemma3",
            "gemma":    "Gemma",
            "phi35":    "Phi-3.5",
            "phi4":     "Phi-4",
            "llama32":  "Llama3.2",
            "llama3":   "Llama3",
            "mistral":  "Mistral",
        }
        return _pretty.get(suffix, suffix.capitalize())

    pdf_map   = _entry_map(table_entry.get("xml_pdfplumber", ""))
    final_map = _entry_map(table_entry.get("xml", ""))

    # Collect all VLM annotation sources present in this entry, in insertion order
    # (xml_vlm first for backward compat, then xml_vlm_gemma, xml_vlm_phi, etc.)
    vlm_sources: list = []  # list of (key, label, entry_map)
    for key in sorted(table_entry.keys(),
                      key=lambda k: (0 if k == "xml_vlm" else 1, k)):
        if key.startswith("xml_vlm") and table_entry.get(key):
            vlm_sources.append((key, _model_label(key), _entry_map(table_entry[key])))

    # Collect all unique cell texts from pdfplumber (or fall back to final)
    all_texts = list(dict.fromkeys(list(pdf_map.keys()) or list(final_map.keys())))

    if not all_texts:
        return ("<p style='color:#888;font-family:sans-serif;font-size:12px;'>"
                "No annotation data available. Run 🤖 Annotation Agent first.</p>")

    total = len(all_texts)

    # VLM column colours — cycle through a palette per model
    _vlm_colours = ["#f39c12", "#9b59b6", "#1abc9c", "#e74c3c"]

    def _bool_icon(b: bool) -> str:
        return "<b style='color:#2ecc71;'>B</b>" if b else "<span style='color:#666;'>—</span>"

    def _indent_badge(n: int) -> str:
        if n == 0:
            return "<span style='color:#555;'>0</span>"
        colours = ["", "#5dade2", "#f39c12", "#e74c3c"]
        return f"<span style='color:{colours[min(n,3)]};font-weight:bold;'>{n}</span>"

    # --- column headers ---
    vlm_header_html = ""
    for ci, (_key, label, _vmap) in enumerate(vlm_sources):
        col = _vlm_colours[ci % len(_vlm_colours)]
        vlm_header_html += (
            f"<th style='padding:4px 8px;color:{col};'>{label} bold</th>"
            f"<th style='padding:4px 8px;color:{col};'>{label} ind</th>"
            f"<th style='padding:4px 8px;color:#aaa;'>vs pdf</th>"
        )
    if not vlm_sources:
        vlm_header_html = (
            "<th style='padding:4px 8px;color:#555;'>VLM bold</th>"
            "<th style='padding:4px 8px;color:#555;'>VLM ind</th>"
            "<th style='padding:4px 8px;color:#555;'>Match</th>"
        )

    # --- table rows ---
    total_conflicts = 0
    rows_html = ""
    for txt in all_texts:
        pb, pi = pdf_map.get(txt, (False, 0))
        fb, fi = final_map.get(txt, (pb, pi))

        row_conflict = False
        vlm_cells = ""
        for ci, (_key, label, vmap) in enumerate(vlm_sources):
            vb, vi = vmap.get(txt, (None, None))
            if vb is not None:
                match = (pb == vb and pi == vi)
                if not match:
                    row_conflict = True
                status = (
                    "<span style='color:#2ecc71;font-size:11px;'>✓</span>" if match
                    else "<span style='color:#f39c12;font-size:11px;font-weight:bold;'>⚠</span>"
                )
                vlm_cells += (
                    f"<td style='text-align:center;padding:3px 6px;'>{_bool_icon(vb)}</td>"
                    f"<td style='text-align:center;padding:3px 6px;'>{_indent_badge(vi)}</td>"
                    f"<td style='text-align:center;padding:3px 6px;'>{status}</td>"
                )
            else:
                vlm_cells += (
                    "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
                    "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
                    "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
                )
        if not vlm_sources:
            vlm_cells = (
                "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
                "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
                "<td style='text-align:center;color:#555;padding:3px 6px;'>—</td>"
            )

        if row_conflict:
            total_conflicts += 1
        row_bg = "background:#1e1a0e;" if row_conflict else ""
        rows_html += (
            f"<tr style='{row_bg}'>"
            f"<td style='padding:3px 8px;max-width:220px;overflow:hidden;"
            f"text-overflow:ellipsis;white-space:nowrap;font-size:11px;color:#ddd;'"
            f" title='{txt}'>{txt}</td>"
            f"<td style='text-align:center;padding:3px 6px;'>{_bool_icon(pb)}</td>"
            f"<td style='text-align:center;padding:3px 6px;'>{_indent_badge(pi)}</td>"
            f"{vlm_cells}"
            f"<td style='text-align:center;padding:3px 6px;'>{_bool_icon(fb)}</td>"
            f"<td style='text-align:center;padding:3px 6px;'>{_indent_badge(fi)}</td>"
            "</tr>"
        )

    title      = table_entry.get("title", "")
    method_tag = table_entry.get("annotation_version", "not run")
    models_run = ", ".join(label for _, label, _ in vlm_sources) if vlm_sources else "none"
    conflict_badge = (
        f"<span style='color:#f39c12;'>{total_conflicts} conflicts</span>"
        if total_conflicts
        else (
            f"<span style='color:#2ecc71;'>no conflicts</span>"
            if vlm_sources else
            "<span style='color:#555;'>VLM not run</span>"
        )
    )

    return (
        f"<div style='background:#1a1a2e;border:1px solid #3a3a5c;border-radius:6px;"
        f"padding:10px 14px;font-family:monospace;font-size:12px;'>"
        f"<div style='color:#7ec8e3;margin-bottom:8px;'>"
        f"{title or 'Table'} &nbsp;·&nbsp; method: <b>{method_tag}</b>"
        f" &nbsp;·&nbsp; models: <b>{models_run}</b>"
        f" &nbsp;·&nbsp; {total} entries &nbsp;·&nbsp; {conflict_badge}</div>"
        f"<div style='overflow-x:auto;max-height:520px;overflow-y:auto;'>"
        f"<table style='border-collapse:collapse;width:100%;'>"
        f"<thead><tr style='border-bottom:1px solid #444;'>"
        f"<th style='padding:4px 8px;text-align:left;color:#aaa;'>Cell text</th>"
        f"<th style='padding:4px 8px;color:#5dade2;'>pdf bold</th>"
        f"<th style='padding:4px 8px;color:#5dade2;'>pdf ind</th>"
        f"{vlm_header_html}"
        f"<th style='padding:4px 8px;color:#2ecc71;'>Final bold</th>"
        f"<th style='padding:4px 8px;color:#2ecc71;'>Final ind</th>"
        f"</tr></thead>"
        f"<tbody>{rows_html}</tbody>"
        f"</table></div></div>"
    )


def _recolor_docx_for_annotation(src_docx: str, dst_docx: str) -> None:
    """Recolor every run in *src_docx* and write to *dst_docx*.

    Uses the same strategy as render_pages.py:
      - Non-table paragraph runs → GREY  (130, 130, 130)
      - Table cell runs           → BLUE  ( 41, 128, 185)
      - Header / footer runs      → GREY  (same as paragraphs)

    python-docx's iter over block items gives perfect paragraph/table
    separation directly from the OOXML structure — no text matching needed.
    """
    from docx import Document as _Document
    from docx.document import Document as _DocumentType
    from docx.oxml.ns import qn as _qn
    from docx.shared import RGBColor as _RGB
    from docx.table import Table as _Table
    from docx.text.paragraph import Paragraph as _Paragraph

    _BLUE = _RGB(0x29, 0x80, 0xB9)   # (41, 128, 185)
    _RED  = _RGB(0xFF, 0x00, 0x00)   # (255, 0, 0)

    def _iter_blocks(parent):
        if isinstance(parent, _DocumentType):
            elm = parent.element.body
        elif hasattr(parent, "_tc"):
            elm = parent._tc
        else:
            elm = parent._element
        for child in elm.iterchildren():
            if child.tag == _qn("w:p"):
                yield _Paragraph(child, parent)
            elif child.tag == _qn("w:tbl"):
                yield _Table(child, parent)

    def _color_para(para, color):
        for run in para.runs:
            run.font.color.rgb = color

    def _color_table(table, color):
        for row in table.rows:
            for cell in row.cells:
                for block in _iter_blocks(cell):
                    if isinstance(block, _Paragraph):
                        _color_para(block, color)
                    elif isinstance(block, _Table):
                        _color_table(block, color)

    doc = _Document(src_docx)
    for block in _iter_blocks(doc):
        if isinstance(block, _Paragraph):
            _color_para(block, _RED)
        elif isinstance(block, _Table):
            _color_table(block, _BLUE)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf is None:
                continue
            for block in _iter_blocks(hf):
                if isinstance(block, _Paragraph):
                    _color_para(block, _RED)
                elif isinstance(block, _Table):
                    _color_table(block, _BLUE)
    doc.save(dst_docx)


def _create_annotated_pages_from_docx(docx_path: str, img_dir: str,
                                       page_indices: list = None,
                                       dpi: int = 150) -> dict:
    """Render per-page annotated PNGs using the recolor approach.

    1. Recolor the .docx (paragraphs=grey, table cells=blue) into a temp file.
    2. Convert the recolored .docx to PDF via LibreOffice.
    3. Render each PDF page to a PNG via PyMuPDF (fitz).

    Returns dict: page_idx (0-based) → absolute path of annotated PNG.
    Only pages whose index is in *page_indices* are written; when
    *page_indices* is None all pages are written.
    """
    import subprocess, tempfile
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("[annot] PyMuPDF (fitz) not installed — falling back to pdf2image")
        fitz = None

    os.makedirs(img_dir, exist_ok=True)
    result_map: dict = {}

    with tempfile.TemporaryDirectory() as tmp:
        colored_docx = os.path.join(tmp, "colored_annot.docx")
        try:
            _recolor_docx_for_annotation(docx_path, colored_docx)
        except Exception as exc:
            print(f"[annot] recolor failed: {exc}")
            return result_map

        pdf_out = os.path.join(tmp, "colored_annot.pdf")
        try:
            r = subprocess.run(
                ["libreoffice", "--headless", "--norestore",
                 "--convert-to", "pdf", "--outdir", tmp, colored_docx],
                capture_output=True, text=True, timeout=120,
            )
            if r.returncode != 0 or not os.path.exists(pdf_out):
                print(f"[annot] LibreOffice failed: {r.stderr[:200]}")
                return result_map
        except Exception as exc:
            print(f"[annot] LibreOffice error: {exc}")
            return result_map

        # Render pages
        if fitz:
            zoom = dpi / 72.0
            mat  = fitz.Matrix(zoom, zoom)
            doc  = fitz.open(pdf_out)
            for pnum in range(len(doc)):
                if page_indices is not None and pnum not in page_indices:
                    continue
                pix  = doc[pnum].get_pixmap(matrix=mat, alpha=False)
                out  = os.path.join(img_dir, f"page_{pnum:03d}_annot.png")
                pix.save(out)
                result_map[pnum] = out
            doc.close()
        else:
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_out, dpi=dpi)
            for pnum, img in enumerate(images):
                if page_indices is not None and pnum not in page_indices:
                    continue
                out = os.path.join(img_dir, f"page_{pnum:03d}_annot.png")
                img.save(out, "PNG")
                result_map[pnum] = out

    print(f"[annot] annotated {len(result_map)} page(s) → {img_dir}")
    return result_map


def _create_annotated_page(pdf_path: str, page_idx: int,
                            page_img_path: str, output_path: str,
                            table_titles: list = None,
                            cell_tokens: set = None):
    """Compatibility shim — kept so existing call sites don't break.
    The real annotation is now done by _create_annotated_pages_from_docx()
    which is called directly from _load_docx_direct with the .docx path.
    This shim is a no-op; it returns None so callers fall through gracefully.
    """
    return None


def _load_docx_direct(fpath: str, on_progress=None):
    """Load a .docx file using python-docx directly, bypassing unstructured.

    Level 1: Uses tabulate to render each table as an aligned text grid.
    Level 2: Renders the full document to per-page PNGs via LibreOffice + pdf2image.
             Each table element gets metadata['image_path'] pointing to the page
             image it appears on (tracked by table order vs page count).

    Iterates the document body in order so that the paragraph immediately
    preceding each table (its title/caption) is attached to that table element.

    Document sequence:
    Every element (paragraph or table) receives a ``doc_seq`` integer that records
    its 0-based position in the original body order.  Tables also carry an
    ``insertion_key`` of the form ``{{TABLE:t<N>}}`` so that the full document
    text flow can be reconstructed with transformed tables substituted in.
    The interleaved sequence is also written to the ``_doc_sequence`` attribute
    on the returned list for use by ``upload_files``.

    Duplicate-elimination — Layer 1 (parse time)
    --------------------------------------------
    In a Word document the caption of a table is almost always the paragraph
    element that immediately precedes the ``<w:tbl>`` element in the body XML.
    A naïve implementation that commits every paragraph immediately would write
    the caption text twice: once as a standalone ``paragraph`` segment and again
    as the ``[Title: …]`` prefix inside the ``table`` segment.

    To prevent this the loop uses a one-slot deferred buffer (``_pending_entry``):

    * A paragraph is *not* appended to ``doc_sequence`` / ``elements`` straight
      away; instead it is held in ``_pending_entry``.
    * When the next body child is a ``<w:tbl>`` the buffer is discarded —
      ``_pending_entry = None`` — because the text already lives in
      ``seg["title"]`` (and therefore in the ``[Title: …]`` line of the table
      block). ``pending_para`` is still set so the table entry is built
      correctly.
    * When the next body child is another ``<w:p>`` the buffer is flushed first
      (the previous paragraph was a real standalone paragraph, not a caption).
    * After the loop a drain step flushes any remaining entry (documents that
      end with a paragraph rather than a table).

    Result: every caption appears exactly once in ``doc_segments.json``, inside
    its table segment. ``reconstruct_document`` then emits it exactly once via
    the ``[Title: …]`` prefix, with no separate paragraph line above.
    """
    import docx
    from langchain.schema import Document
    from tabulate import tabulate

    class _DocList(list):
        """list subclass that can carry extra attributes (e.g. _doc_sequence)."""
        pass

    doc = docx.Document(fpath)
    elements = _DocList()
    table_index = 0
    doc_seq = 0          # monotonic counter across all body children
    # Full ordered sequence of body segments for document reconstruction.
    # Each entry: {"seq": int, "type": "paragraph"|"table", "text": str,
    #              "insertion_key": str (tables only), "table_index": int (tables only)}
    doc_sequence = []

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}
    pending_para = ""
    _pending_entry = None  # (doc_sequence_dict, Document) buffered until we know it's not a table title

    img_dir = os.path.join(os.path.dirname(fpath), "table_images")

    if on_progress:
        on_progress(0, 0, "Converting document to PDF\u2026")

    # Level 2: convert to PDF and render all pages
    page_images, pdf_path = _render_docx_pages_to_images(fpath, img_dir)
    total_pages = len(page_images)

    # Extract per-page text from the PDF to detect which pages contain tables
    page_texts = []
    if pdf_path and os.path.exists(pdf_path):
        try:
            from pypdf import PdfReader
            reader = PdfReader(pdf_path)
            page_texts = [p.extract_text() or "" for p in reader.pages]
            print(f"[_load_docx_direct] Extracted text from {len(page_texts)} PDF pages")
        except Exception as exc:
            print(f"[_load_docx_direct] pypdf error: {exc}")

    # Normalise page texts once for case-insensitive matching
    page_texts_norm = [pt.lower() for pt in page_texts]

    # Pre-scan: collect (cell_rows, title) per table so the title paragraph can
    # join the fingerprint — it is usually the most distinctive text on that page.
    table_cells_by_idx: dict = {}  # tidx -> (rows, title_str)
    _scan_idx = 0
    _pre_title = ""
    for _child in doc.element.body.iterchildren():
        _tag = _child.tag.split("}")[-1] if "}" in _child.tag else _child.tag
        if _tag == "p":
            _para = para_map.get(_child)
            _txt = _para.text.strip() if _para else ""
            if _txt:
                _pre_title = _txt
        elif _tag == "tbl":
            _tbl = table_map.get(_child)
            if _tbl is not None:
                _rows = []
                for _row in _tbl.rows:
                    _seen_tc: set = set()
                    _rcells = []
                    for _c in _row.cells:
                        if id(_c._tc) not in _seen_tc:
                            _seen_tc.add(id(_c._tc))
                            _rcells.append(_c.text.strip())
                    _rows.append(_rcells)
                table_cells_by_idx[_scan_idx] = ([r for r in _rows if any(r)], _pre_title)
                _scan_idx += 1
            _pre_title = ""

    if on_progress:
        _n = len(table_cells_by_idx)
        on_progress(0, _n, f"Found {_n} table{'s' if _n != 1 else ''}\u2014verifying\u2026")

    def _build_tokens(cell_rows, title):
        """Collect lowercase fingerprint tokens from title + all cell values."""
        tokens = []
        # Title words are the most distinctive — add each word ≥4 chars individually
        for word in title.split():
            w = word.strip("(),.:;-").lower()
            if len(w) >= 4:
                tokens.append(w)
        # Whole-cell values from every row (not just the first few)
        for row in cell_rows:
            for cell in row:
                cell_norm = cell.strip().lower()
                if len(cell_norm) >= 4:
                    tokens.append(cell_norm)
                # Also split multi-word cells so e.g. "Total Assets" → "total", "assets"
                for word in cell_norm.split():
                    w = word.strip("(),.:;-")
                    if len(w) >= 4 and w != cell_norm:
                        tokens.append(w)
        return tokens

    def _best_page_for_table(cell_rows, title=""):
        """Return 0-based page index scored by length-weighted token hits, or None."""
        if not page_texts_norm:
            return None
        tokens = _build_tokens(cell_rows, title)
        if not tokens:
            return None
        # Weight each hit by token length — longer tokens are more distinctive
        scores = [
            sum(len(t) for t in tokens if t in pt)
            for pt in page_texts_norm
        ]
        best = max(range(len(scores)), key=lambda i: scores[i])
        return best if scores[best] > 0 else None

    def _best_page_for_para(text):
        """Return 0-based page index that best matches a paragraph, or None.

        Strategy:
        1. Exact substring match of the full lowercased text.
        2. Length-weighted token match (same metric as _best_page_for_table).
        """
        if not page_texts_norm or not text:
            return None
        text_norm = text.lower().strip()
        # Exact match first — most reliable
        for i, pt in enumerate(page_texts_norm):
            if text_norm in pt:
                return i
        # Token fallback
        tokens = [w.strip("(),.:;-").lower() for w in text.split() if len(w) >= 4]
        if not tokens:
            return None
        scores = [sum(len(t) for t in tokens if t in pt) for pt in page_texts_norm]
        best = max(range(len(scores)), key=lambda i: scores[i])
        return best if scores[best] > 0 else None

    # Build table_index → image_path mapping (None for tables with no page match)
    table_to_page_img: dict = {}
    for tidx, (rows, title) in table_cells_by_idx.items():
        pg = _best_page_for_table(rows, title)
        if pg is not None and pg < total_pages:
            table_to_page_img[tidx] = page_images[pg]
            print(f"[_load_docx_direct] table {tidx} '{title[:40]}' → page {pg}")
        elif page_images and total_pages > 0:
            # Fallback to proportional distribution if pypdf gave no match
            fallback = min(int(tidx * total_pages / max(len(table_cells_by_idx), 1)), total_pages - 1)
            table_to_page_img[tidx] = page_images[fallback]
            print(f"[_load_docx_direct] table {tidx} '{title[:40]}' → fallback page {fallback}")

    table_pages_matched = sum(1 for v in table_to_page_img.values() if v is not None)
    text_only_pages = total_pages - len({v for v in table_to_page_img.values() if v})
    print(f"[_load_docx_direct] {table_pages_matched}/{len(table_cells_by_idx)} tables matched to pages; "
          f"{text_only_pages} text-only pages (no image stored)")

    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Flush any buffered paragraph that was not consumed as a table title
            if _pending_entry is not None:
                doc_sequence.append(_pending_entry[0])
                elements.append(_pending_entry[1])
                _pending_entry = None
            para = para_map.get(child)
            text = para.text.strip() if para else ""
            if text:
                pending_para = text
                para_pg = _best_page_for_para(text)
                para_img  = page_images[para_pg] if (para_pg is not None and para_pg < total_pages) else ""
                para_pgidx = para_pg if para_pg is not None else -1
                _pending_entry = (
                    {"seq": doc_seq, "type": "paragraph", "text": text,
                     "image_path": para_img, "page_idx": para_pgidx},
                    Document(
                        page_content=text,
                        metadata={
                            "source": fpath, "type": "paragraph", "doc_seq": doc_seq,
                            "image_path": para_img, "page_idx": para_pgidx,
                        }
                    )
                )
            doc_seq += 1

        elif tag == "tbl":
            # The immediately preceding paragraph is this table's title — consume
            # it here so it doesn't also appear as a standalone paragraph in the
            # reconstructed document (which would duplicate the title text).
            _pending_entry = None
            table = table_map.get(child)
            if table is None:
                continue

            # Extract cell data, de-duplicating only truly merged cells.
            # Use id(cell._tc) (XML element identity) so two distinct cells with
            # the same text (e.g. two "0" columns) are both kept, while a merged
            # cell that python-docx returns multiple times for each column it spans
            # is included only once.
            def _cell_text(cell):
                """Return full text of a cell using robust OOXML extraction."""
                from docx.oxml.ns import qn as _qn
                return _extract_tc_text(cell._tc, _qn)

            cell_rows = []
            for row in table.rows:
                seen_tc: set = set()
                row_cells = []
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id not in seen_tc:
                        seen_tc.add(tc_id)
                        row_cells.append(_cell_text(cell))
                cell_rows.append(row_cells)

            cell_rows = [r for r in cell_rows if any(c for c in r)]
            print(f"[table {table_index}] '{pending_para[:40]}' — {len(cell_rows)} rows, "
                  f"sample: {cell_rows[0] if cell_rows else []}")
            if not cell_rows:
                pending_para = ""
                continue

            # Level 1: tabulate text rendering — use "simple" (not "grid") to keep
            # char count low enough that large tables (e.g. full balance sheets) fit
            # within the LLM context budget without truncation.
            table_text = tabulate(cell_rows, tablefmt="simple")
            title_line = f"[Title: {pending_para}]\n" if pending_para else ""

            # Level 2: use fingerprint-matched page image
            img_path = table_to_page_img.get(table_index)

            # Build CALS XML first — it is the reference for the inspection step below
            try:
                cals_xml = _build_cals_xml(table, title=pending_para, table_id=f"t{table_index}")
            except Exception as _xe:
                print(f"[_build_cals_xml] table {table_index}: {_xe}")
                cals_xml = None

            # Phase 1 inspection: compare CALS XML against pdfplumber PDF cross-check
            page_idx_for_table = next(
                (pi for pi, pg_img in enumerate(page_images) if pg_img == img_path), None
            )
            inspection = None
            if pdf_path and os.path.exists(pdf_path) and page_idx_for_table is not None and cals_xml:
                inspection = verify_table(
                    cals_xml, pdf_path, page_idx_for_table, pending_para,
                    img_path=img_path or ""
                )

            # Phase 2: style annotation — adds bold/indent attributes to <entry>
            #
            # Primary (always): read <w:b> and <w:ind> directly from OOXML.
            # This is fully reliable because it uses Word's own markup — text
            # wrapping, run-splitting, and \xa0 in the PDF are irrelevant.
            #
            # Supplementary (optional, controlled by STYLE_ANNOTATE_METHOD):
            #   pdfplumber — font name + x0 from the LibreOffice PDF; can refine
            #                indent levels with pixel-accurate x0 measurements.
            #   vlm        — Qwen2.5-VL via LM Studio (requires local model server)
            if cals_xml:
                cals_xml = _annotate_entry_styles_from_ooxml(cals_xml, table)
            if cals_xml and STYLE_ANNOTATE_METHOD == "pdfplumber" and pdf_path and page_idx_for_table is not None:
                cals_xml = _annotate_entry_styles_with_pdfplumber(
                    cals_xml, pdf_path, page_idx_for_table
                )
            elif cals_xml and STYLE_ANNOTATE_METHOD == "vlm" and img_path:
                cals_xml = _annotate_entry_styles_with_vlm(
                    cals_xml, img_path, lmstudio_url=VLM_LMSTUDIO_URL
                )

            import json as _json
            insertion_key = f"{{{{TABLE:t{table_index}}}}}"
            doc_sequence.append({
                "seq": doc_seq,
                "type": "table",
                "table_index": table_index,
                "insertion_key": insertion_key,
                "title": pending_para,
                "text": f"{title_line}{table_text}",
                "page_idx": page_idx_for_table if page_idx_for_table is not None else -1,
                "image_path": img_path or "",
            })
            elements.append(Document(
                page_content=f"{title_line}{table_text}",
                metadata={
                    "source": fpath,
                    "type": "table",
                    "table_index": table_index,
                    "title": pending_para,
                    "image_path": img_path or "",
                    "pdf_path": pdf_path or "",
                    "page_idx": page_idx_for_table if page_idx_for_table is not None else -1,
                    "cell_rows": _json.dumps(cell_rows),
                    "inspection": _json.dumps(inspection) if inspection is not None else "",
                    "xml": cals_xml or "",
                    "doc_seq": doc_seq,
                    "insertion_key": insertion_key,
                }
            ))
            doc_seq += 1
            if on_progress:
                _tot = len(table_cells_by_idx)
                _lbl = (pending_para or f"Table {table_index + 1}")[:45]
                on_progress(table_index + 1, _tot, f"Verified {table_index + 1}/{_tot}: {_lbl}")
            table_index += 1
            pending_para = ""

    # Flush any trailing paragraph (document ends with paragraph, not a table)
    if _pending_entry is not None:
        doc_sequence.append(_pending_entry[0])
        elements.append(_pending_entry[1])
        _pending_entry = None

    print(f"[_load_docx_direct] Loaded {table_index} tables, {len(page_images)} page images")

    # ── Merge consecutive paragraph segments ──────────────────────────────────
    # All consecutive paragraph segments on the same page are joined with a
    # newline into a single segment.  Table segments are never merged — each
    # table always remains its own segment.
    _merged_ds: list = []
    _merged_el: list = []
    _pbuf_ds:   list = []   # buffer of consecutive paragraph dicts
    _pbuf_el:   list = []   # parallel buffer of Document objects

    def _flush_pbuf():
        if not _pbuf_ds:
            return
        if len(_pbuf_ds) == 1:
            _merged_ds.append(_pbuf_ds[0])
            _merged_el.append(_pbuf_el[0])
        else:
            _mtxt = "\n".join(s["text"] for s in _pbuf_ds)
            _mseq = _pbuf_ds[0]["seq"]
            _merged_ds.append({
                "seq":        _mseq,
                "type":       "paragraph",
                "text":       _mtxt,
                "image_path": _pbuf_ds[0].get("image_path", ""),
                "page_idx":   _pbuf_ds[0].get("page_idx", -1),
            })
            _merged_el.append(Document(
                page_content=_mtxt,
                metadata={**_pbuf_el[0].metadata, "doc_seq": _mseq},
            ))
        _pbuf_ds.clear()
        _pbuf_el.clear()

    for _seg, _elem in zip(doc_sequence, elements):
        if _seg["type"] == "paragraph":
            _same_page = (not _pbuf_ds or
                          _pbuf_ds[0].get("page_idx") == _seg.get("page_idx"))
            if _same_page:
                _pbuf_ds.append(_seg)
                _pbuf_el.append(_elem)
            else:
                _flush_pbuf()
                _pbuf_ds.append(_seg)
                _pbuf_el.append(_elem)
        else:                                # table — always commit directly
            _flush_pbuf()
            _merged_ds.append(_seg)
            _merged_el.append(_elem)

    _flush_pbuf()
    doc_sequence = _merged_ds
    elements     = _DocList(_merged_el)

    # Re-index: assign clean seq values (0, 1, 2, …) after merging and keep
    # element metadata.doc_seq in sync.  table_index values on table segments
    # are already correct and are left unchanged.
    for _new_seq, (_seg, _elem) in enumerate(zip(doc_sequence, elements)):
        _seg["seq"] = _new_seq
        _elem.metadata["doc_seq"] = _new_seq

    print(f"[_load_docx_direct] After paragraph merge: {len(doc_sequence)} segments")

    # ── Approximate bounding boxes ─────────────────────────────────────────────
    # For each segment, store a normalised bbox [x0, y0, x1, y1] (values 0‥1
    # relative to the page) that indicates the segment's rough position on its
    # page.  "Conceptual" accuracy is fine — we probe a few key phrases/tokens
    # with pdfplumber word coordinates and take the union.
    #
    # Paragraph segments: probe each original paragraph line (split on "\n")
    #   with its first 4 words to get boxes spread across the full vertical span.
    # Table segments: probe every cell token (len ≥ 4 chars) found on the page.
    try:
        import pdfplumber as _pdfp
        _pdfp_ok = True
    except ImportError:
        _pdfp_ok = False
        print("[_load_docx_direct] pdfplumber not available — bounding boxes skipped")

    _bbox_strip = str.maketrans('', '', '(),.:;%\'"' + '\xa0')

    def _bnorm(w):
        return w.lower().replace('\xa0', ' ').translate(_bbox_strip).strip()

    def _first_match_boxes(pg_words, pg_wnorms, probe_norms):
        """Return list of (x0,top,x1,bot) for the first consecutive phrase match."""
        plen = len(probe_norms)
        if plen == 0:
            return []
        for i in range(len(pg_wnorms) - plen + 1):
            if pg_wnorms[i:i + plen] == probe_norms:
                return [(pg_words[j]["x0"], pg_words[j]["top"],
                         pg_words[j]["x1"], pg_words[j]["bottom"])
                        for j in range(i, i + plen)]
        return []

    def _union_norm_bbox(boxes, pw, ph):
        """Normalised [x0, y0, x1, y1] over all boxes, or None if empty."""
        if not boxes:
            return None
        return [
            round(min(b[0] for b in boxes) / pw, 4),
            round(min(b[1] for b in boxes) / ph, 4),
            round(max(b[2] for b in boxes) / pw, 4),
            round(max(b[3] for b in boxes) / ph, 4),
        ]

    if _pdfp_ok and pdf_path and os.path.exists(pdf_path):
        import json as _bbox_json
        _page_word_cache: dict = {}   # page_idx → (words, wnorms, width, height)

        def _pg_words(pi, _pdf_obj):
            if pi not in _page_word_cache:
                _pg  = _pdf_obj.pages[pi]
                _ws  = _pg.extract_words(keep_blank_chars=False, use_text_flow=False)
                _wn  = [_bnorm(w["text"]) for w in _ws]
                _page_word_cache[pi] = (_ws, _wn, float(_pg.width), float(_pg.height))
            return _page_word_cache[pi]

        with _pdfp.open(pdf_path) as _pdf_bbox:
            for _seg, _elem in zip(doc_sequence, elements):
                _pi = _seg.get("page_idx", -1)
                if _pi < 0 or _pi >= len(_pdf_bbox.pages):
                    _seg["bbox"] = None
                    _elem.metadata["bbox"] = None
                    continue

                _ws, _wn, _pw, _ph = _pg_words(_pi, _pdf_bbox)
                _all_boxes = []

                if _seg["type"] == "paragraph":
                    # Probe each original paragraph line (joined by "\n") with
                    # its first 4 normalised words so boxes span the full height.
                    for _line in _seg["text"].split("\n"):
                        _tw = [_bnorm(w) for w in _line.split() if _bnorm(w)]
                        _probe = _tw[:4]
                        if len(_probe) >= 2:
                            _all_boxes += _first_match_boxes(_ws, _wn, _probe)

                else:  # table — probe cell tokens as individual word hits
                    try:
                        _cr = _bbox_json.loads(_elem.metadata.get("cell_rows", "[]"))
                    except Exception:
                        _cr = []
                    _cell_toks = {_bnorm(str(c)) for row in _cr for c in row
                                  if len(_bnorm(str(c))) >= 4}
                    for _wobj, _wn_tok in zip(_ws, _wn):
                        if _wn_tok in _cell_toks:
                            _all_boxes.append((_wobj["x0"], _wobj["top"],
                                               _wobj["x1"], _wobj["bottom"]))

                _bbox = _union_norm_bbox(_all_boxes, _pw, _ph)
                _seg["bbox"] = _bbox
                _elem.metadata["bbox"] = _bbox

        print(f"[_load_docx_direct] Bounding boxes computed for {len(doc_sequence)} segments")
    else:
        for _seg, _elem in zip(doc_sequence, elements):
            _seg["bbox"] = None
            _elem.metadata["bbox"] = None

    # ── Page annotation ────────────────────────────────────────────────────────
    # Recolor the original .docx (paragraphs=grey, table cells=blue), convert
    # to PDF via LibreOffice, and render per-page annotated PNGs.  This is the
    # same approach as render_pages.py — structure comes directly from the OOXML
    # so no text matching or boundary detection is needed.
    _content_pages = sorted({s["page_idx"] for s in doc_sequence
                              if s.get("page_idx", -1) >= 0})
    _annot_img_map: dict = {}
    if _content_pages:
        _annot_img_map = _create_annotated_pages_from_docx(
            fpath, img_dir, page_indices=set(_content_pages)
        )
    if _annot_img_map:
        for s in doc_sequence:
            _pi = s.get("page_idx", -1)
            if _pi in _annot_img_map:
                s["annotated_image_path"] = _annot_img_map[_pi]
        _seq_to_annot = {
            s["seq"]: _annot_img_map[s["page_idx"]]
            for s in doc_sequence if s.get("page_idx", -1) in _annot_img_map
        }
        for elem in elements:
            _ds = elem.metadata.get("doc_seq", -1)
            if _ds in _seq_to_annot:
                elem.metadata["annotated_image_path"] = _seq_to_annot[_ds]
        print(f"[_load_docx_direct] Pages annotated: {sorted(_annot_img_map)}")

    # Attach sequence to the list so upload_files can persist it without a second pass
    elements._doc_sequence = doc_sequence
    return elements


def _annotate_cals_xml(cals_xml: str, unconfirmed_values: set, reason_map: dict = None) -> str:
    """Return a copy of the CALS XML with verify="unconfirmed" on entries
    whose text is not confirmed by the pdfplumber cross-check.
    Entries that are confirmed get verify="ok".
    When reason_map is supplied, unconfirmed entries also get a verify-reason attribute
    explaining why the match failed (e.g. what pdfplumber found instead).
    """
    import re
    import xml.etree.ElementTree as ET

    def _norm(v):
        v = str(v or "").strip()
        return v.lower() if v else None

    if not cals_xml:
        return cals_xml
    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError:
        return cals_xml

    for entry in root.iter("entry"):
        # Join all inner text (covers plain text AND <para> children)
        text = " ".join("".join(entry.itertext()).split())
        if not text:
            continue
        norm = _norm(text)
        if norm in unconfirmed_values:
            entry.set("verify", "unconfirmed")
            if reason_map:
                entry.set("verify-reason",
                          reason_map.get(norm, "Not found in PDF text extraction"))
        else:
            entry.set("verify", "ok")

    _indent_xml(root)
    return ET.tostring(root, encoding="unicode")


def _annotate_page_image(img_path: str, pdf_path: str, page_idx: int,
                         unconfirmed_values: set, confirmed_values: set = None,
                         dpi: int = 150) -> str:
    """Draw colour-coded boxes on every table cell in the page image:
      - GREEN  (semi-transparent) : cell value confirmed by pdfplumber cross-check
      - RED    (semi-transparent) : cell value in CALS XML but NOT confirmed by pdfplumber
      - YELLOW (semi-transparent) : cell has no text / empty

    Uses pdfplumber's per-cell bboxes for precise placement.
    Saves the annotated image as <original>_diff.png and returns its path.
    Falls back to returning img_path unchanged on any error.
    """
    import re
    try:
        from PIL import Image, ImageDraw, ImageFont
        import pdfplumber
    except ImportError as exc:
        print(f"[annotate_page_image] missing dependency: {exc}")
        return img_path

    if not img_path or not os.path.exists(img_path):
        return img_path
    if not os.path.exists(pdf_path):
        return img_path

    confirmed_values = confirmed_values or set()

    def _norm(v):
        v = str(v or "").strip()
        return v.lower() if v else None

    # Colour palette  (R, G, B, A)
    COLOR_OK       = (40,  180,  60, 70)    # green fill
    COLOR_OK_BORD  = (20,  140,  30, 210)   # green border
    COLOR_BAD      = (220,  50,  50, 90)    # red fill
    COLOR_BAD_BORD = (200,   0,   0, 220)   # red border
    COLOR_EMPTY    = (200, 180,   0, 50)    # yellow fill (no text)
    COLOR_EMPTY_B  = (160, 140,   0, 160)   # yellow border

    try:
        img = Image.open(img_path).convert("RGBA")
        overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
        draw = ImageDraw.Draw(overlay)
        scale = dpi / 72.0

        # Try to load a small font for labels; fall back to default if unavailable
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 11)
        except Exception:
            font = ImageFont.load_default()

        cells_drawn = 0
        with pdfplumber.open(pdf_path) as pdf:
            if page_idx >= len(pdf.pages):
                return img_path
            page = pdf.pages[page_idx]
            # Lattice (border-based) first; fall back to stream (text-alignment) for
            # borderless tables produced by LibreOffice from Word .docx files.
            tables = page.find_tables()
            if not tables:
                tables = page.find_tables({
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                })
            for tbl in tables:
                extracted = tbl.extract()   # [[text, ...], ...]
                for row_cells, row_texts in zip(tbl.rows, extracted or []):
                    for cell_bbox, cell_text in zip(row_cells.cells, row_texts or []):
                        if cell_bbox is None:
                            continue
                        x0, top, x1, bottom = cell_bbox
                        px0 = int(x0     * scale)
                        py0 = int(top    * scale)
                        px1 = int(x1     * scale)
                        py1 = int(bottom * scale)

                        norm = _norm(cell_text or "")
                        if not norm:
                            fill, border, label = COLOR_EMPTY, COLOR_EMPTY_B, ""
                        elif norm in unconfirmed_values:
                            fill, border, label = COLOR_BAD, COLOR_BAD_BORD, "\u2717"
                        else:
                            # confirmed (either in confirmed_values set or simply not flagged)
                            fill, border, label = COLOR_OK, COLOR_OK_BORD, "\u2713"

                        draw.rectangle([px0, py0, px1, py1],
                                       fill=fill, outline=border, width=2)
                        # Small label in top-left corner of cell when there's room
                        if label and (px1 - px0) > 14 and (py1 - py0) > 10:
                            draw.text((px0 + 3, py0 + 2), label, fill=border, font=font)
                        # For red cells: overlay the raw text pdfplumber found so the user
                        # can see exactly what the PDF extractor saw in that cell position
                        if fill == COLOR_BAD and cell_text and (px1 - px0) > 40 and (py1 - py0) > 20:
                            raw_label = (cell_text.strip()).replace("\n", " ")[:24]
                            draw.text((px0 + 3, py0 + 13), raw_label,
                                      fill=(160, 0, 0), font=font)
                        cells_drawn += 1

        annotated = Image.alpha_composite(img, overlay).convert("RGB")
        out_path = img_path.replace(".png", "_diff.png")
        annotated.save(out_path, "PNG")
        print(f"[annotate_page_image] {cells_drawn} cells annotated \u2192 {out_path}")
        return out_path
    except Exception as exc:
        print(f"[annotate_page_image] {exc}")
        return img_path


def verify_table(cals_xml: str, pdf_path: str, page_idx: int,
                 table_title: str = "", img_path: str = "") -> dict:
    """Inspect extraction quality by comparing CALS XML entries (canonical Word source)
    against pdfplumber's independent PDF extraction.

    Using CALS XML as the reference (rather than raw cell_rows) gives a span-aware
    comparison: merged cells appear exactly once, column widths are preserved, and
    the extracted value set is not polluted by python-docx's virtual grid fill.

    Returns a dict with: table, cals_shape, pdf_shape, in_cals_not_pdf,
    in_pdf_not_cals, coverage_pct, verdict.
    """
    import re
    import xml.etree.ElementTree as ET

    try:
        import pdfplumber
    except ImportError:
        return {"table": table_title, "verdict": "SKIP", "reason": "pdfplumber not installed"}

    def _normalize(v):
        v = str(v or "").strip()
        return v.lower() if v else None

    # ------------------------------------------------------------------
    # 1. Extract value set from CALS XML (span-aware canonical source)
    # ------------------------------------------------------------------
    if not cals_xml:
        return {"table": table_title, "verdict": "SKIP", "reason": "no CALS XML available"}
    try:
        root = ET.fromstring(cals_xml)
    except ET.ParseError as exc:
        return {"table": table_title, "verdict": "SKIP", "reason": f"CALS parse error: {exc}"}

    cals_entries = []
    cals_multipara_values: set = set()  # entries with <para> children — eligible for multi-line match
    cals_norm_to_raw: dict = {}          # normalized value → first raw entry text (for word search)
    for entry in root.iter("entry"):
        # Join all inner text (covers plain text AND <para> children)
        text = " ".join("".join(entry.itertext()).split())
        if text:
            cals_entries.append(text)
            _n_raw = _normalize(text)
            if _n_raw and _n_raw not in cals_norm_to_raw:
                cals_norm_to_raw[_n_raw] = text
            # Track which entries use <para> children (explicit line-break structure)
            if entry.find("para") is not None:
                _n = _normalize(text)
                if _n:
                    cals_multipara_values.add(_n)

    tgroup = root.find("tgroup")
    cals_cols = int(tgroup.get("cols", 0)) if tgroup is not None else 0
    cals_rows = sum(1 for _ in root.iter("row"))
    cals_values = {_normalize(v) for v in cals_entries if _normalize(v) is not None}

    # ------------------------------------------------------------------
    # 2. Extract value set from pdfplumber (independent PDF render check)
    # ------------------------------------------------------------------
    if not os.path.exists(pdf_path):
        return {"table": table_title, "verdict": "SKIP", "reason": f"PDF not found: {pdf_path}"}

    pdf_ref = None
    _page_words: list = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if page_idx >= len(pdf.pages):
                return {"table": table_title, "verdict": "SKIP",
                        "reason": f"page_idx {page_idx} out of range ({len(pdf.pages)} pages)"}
            page = pdf.pages[page_idx]
            pdf_ref = page.extract_table()
            if pdf_ref is None:
                pdf_ref = page.extract_table(table_settings={
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                })
            # Also capture raw words for the column-grouping fallback (section 2c)
            _page_words = page.extract_words(keep_blank_chars=False,
                                             x_tolerance=3, y_tolerance=3)
    except Exception as exc:
        return {"table": table_title, "verdict": "SKIP", "reason": str(exc)}

    if pdf_ref is None:
        return {"table": table_title, "verdict": "SKIP",
                "reason": "pdfplumber found no table on this page"}

    pdf_rows_count = len(pdf_ref)
    pdf_cols_count = max((len(r or []) for r in pdf_ref), default=0)
    pdf_values = {_normalize(c) for row in pdf_ref
                  for c in (row or []) if _normalize(c) is not None}

    # Build norm→raw map so we can explain mismatches (what pdfplumber actually found)
    pdf_raw_map: dict = {}  # norm_val → first raw text that produced it
    for _row in (pdf_ref or []):
        for _c in (_row or []):
            _n = _normalize(_c)
            if _n is not None and _n not in pdf_raw_map:
                pdf_raw_map[_n] = str(_c or "").strip()

    # ------------------------------------------------------------------
    # 2b. Supplement with Camelot (handles merged/border cells better)
    # ------------------------------------------------------------------
    camelot_values: set = set()
    try:
        import camelot  # optional; silently skipped if not installed
        _ctables = camelot.read_pdf(
            pdf_path,
            pages=str(page_idx + 1),   # camelot uses 1-based page numbers
            flavor="lattice",
            suppress_stdout=True,
        )
        if len(_ctables) == 0:
            # lattice found nothing — retry with stream flavor
            _ctables = camelot.read_pdf(
                pdf_path,
                pages=str(page_idx + 1),
                flavor="stream",
                suppress_stdout=True,
            )
        for _ct in _ctables:
            for _row in _ct.data:
                for _cell in _row:
                    _n = _normalize(_cell)
                    if _n is not None:
                        camelot_values.add(_n)
                        if _n not in pdf_raw_map:  # keep first raw text seen per norm value
                            pdf_raw_map[_n] = str(_cell or "").strip()
        if camelot_values:
            print(f"[verify_table] Camelot added {len(camelot_values)} values "
                  f"({len(camelot_values - pdf_values)} new)")
            pdf_values = pdf_values | camelot_values
    except ImportError:
        pass   # camelot not installed — pdfplumber-only mode
    except Exception as _ce:
        print(f"[verify_table] Camelot fallback failed: {_ce}")

    # ------------------------------------------------------------------
    # 2c. Multi-line header cell reconstruction
    #     pdfplumber's text strategy can bisect a word at a column boundary
    #     (e.g. "September" → "Sept" | "ember 30,"), and extract_text() fuses
    #     adjacent column headers that appear on the same visual line.
    #     Fix: group words by y-line, split each line into horizontal x-clusters
    #     (gap > 20pt), then pair consecutive-line clusters whose x-ranges overlap.
    #     This reconstructs multi-line cell text like "September 30,\n2025".
    # ------------------------------------------------------------------
    try:
        if _page_words:
            from collections import defaultdict as _ddict

            def _x_clusters(word_list):
                """Split one line's words into horizontally-separated clusters."""
                ws = sorted(word_list, key=lambda w: w["x0"])
                clusters = []
                cur = [ws[0]]
                cur_right = ws[0]["x1"]
                for _ww in ws[1:]:
                    if _ww["x0"] - cur_right > 20:
                        clusters.append(cur)
                        cur = [_ww]
                    else:
                        cur.append(_ww)
                    cur_right = max(cur_right, _ww["x1"])
                if cur:
                    clusters.append(cur)
                return clusters

            # Group words into y-lines (bucket to 3pt)
            _lines: dict = _ddict(list)
            for _ww in _page_words:
                _lines[round(_ww["top"] / 3) * 3].append(_ww)
            _sorted_lines = sorted(_lines.items())

            _cross_line_vals: set = set()
            for _i in range(len(_sorted_lines) - 1):
                _y1, _words1 = _sorted_lines[_i]
                _y2, _words2 = _sorted_lines[_i + 1]
                if _y2 - _y1 > 20:   # too far apart — not the same cell
                    continue
                for _c1 in _x_clusters(_words1):
                    _c1_x0 = min(w["x0"] for w in _c1)
                    _c1_x1 = max(w["x1"] for w in _c1)
                    for _c2 in _x_clusters(_words2):
                        _c2_x0 = min(w["x0"] for w in _c2)
                        _c2_x1 = max(w["x1"] for w in _c2)
                        # Only pair if the x-ranges overlap (same column)
                        if _c2_x0 > _c1_x1 + 10 or _c1_x0 > _c2_x1 + 10:
                            continue
                        _l1 = " ".join(w["text"] for w in sorted(_c1, key=lambda w: w["x0"]))
                        _l2 = " ".join(w["text"] for w in sorted(_c2, key=lambda w: w["x0"]))
                        _phrase = _l1 + " " + _l2
                        _n = _normalize(_phrase)
                        if _n is not None and _n not in pdf_values:
                            _cross_line_vals.add(_n)
                            if _n not in pdf_raw_map:
                                pdf_raw_map[_n] = _phrase
            # *** Do NOT merge into pdf_values directly ***
            # These phrases only confirm entries that explicitly model a line break
            # via <para> children (cals_multipara_values). A plain-text entry like
            # "September 30, 2025" must be found by pdfplumber/camelot directly.
            pdf_multiline_values = _cross_line_vals
            _new_grp = pdf_multiline_values & cals_multipara_values
            if _new_grp:
                print(f"[verify_table] Multi-line header reconstruction confirmed "
                      f"{len(_new_grp)} <para>-structured values: {sorted(_new_grp)[:5]}")
    except Exception as _cge:
        print(f"[verify_table] Multi-line header reconstruction failed: {_cge}")

    # ------------------------------------------------------------------
    # 3. Compare: CALS is the reference, PDF is the cross-check
    #
    # effective_pdf_values = direct extractions  UNION
    #   (multi-line reconstructions INTERSECT <para>-structured CALS entries)
    #
    # The intersection gate means a plain-text "September 30, 2025" entry can
    # only be confirmed by direct extraction, not by the multi-line fallback.
    # A <para>-structured entry CAN be confirmed by multi-line reconstruction
    # because the author explicitly modelled the line break.
    # ------------------------------------------------------------------
    try:
        effective_pdf_values = pdf_values | (pdf_multiline_values & cals_multipara_values)
    except NameError:
        effective_pdf_values = pdf_values  # pdf_multiline_values not set (old path)

    # ------------------------------------------------------------------
    # 3b. OOXML-sourced page-word confirmation
    #     pdfplumber's extract_table() often fails to detect borderless label
    #     rows (e.g. "Commitments and contingencies") even though the words
    #     ARE present on the page.  For every value that table-extraction
    #     missed, look it up as free-form words via extract_words() instead.
    #     Only non-numeric values use this path: numbers must be confirmed by
    #     structured table extraction (pdfplumber / Camelot) to be meaningful.
    # ------------------------------------------------------------------
    if _page_words:
        _pw_tokens = {w["text"].lower() for w in _page_words}
        _word_confirmed: set = set()
        for _nv in (cals_values - effective_pdf_values):
            # Look up the original OOXML/CALS raw text for this normalised value
            _raw = cals_norm_to_raw.get(_nv, "")
            # Skip numeric values — table-structure extraction is the authority
            try:
                float(re.sub(r"[\$,\(\)\u2014\u2013\s]", "", _raw).replace("-", ""))
                continue
            except (ValueError, TypeError):
                pass
            if not _raw:
                continue
            # All significant words (>=3 chars) must be found anywhere on the page
            _toks = [t.lower() for t in _raw.split() if len(t) >= 3]
            if _toks and all(t in _pw_tokens for t in _toks):
                _word_confirmed.add(_nv)
        if _word_confirmed:
            print(
                f"[verify_table] OOXML word-search confirmed {len(_word_confirmed)} text-label "
                f"value(s) missed by pdfplumber table detection: "
                f"{sorted(cals_norm_to_raw.get(v, v) for v in _word_confirmed)[:5]}"
            )
            effective_pdf_values = effective_pdf_values | _word_confirmed

    in_cals_not_pdf = sorted(str(v) for v in cals_values - effective_pdf_values)
    in_pdf_not_cals = sorted(str(v) for v in effective_pdf_values - cals_values)

    # Coverage: what % of CALS values are confirmed by pdfplumber
    confirmed = len(cals_values & effective_pdf_values)
    coverage = confirmed / max(len(cals_values), 1)

    verdict = "PASS" if coverage >= 0.85 and not in_cals_not_pdf else "WARN" if coverage >= 0.6 else "FAIL"

    unconfirmed_set = set(in_cals_not_pdf)  # already normalized strings
    confirmed_set   = set(str(v) for v in cals_values & effective_pdf_values)

    # Build per-value reason strings that explain each mismatch to the user
    reason_map: dict = {}
    for _cals_raw in cals_entries:
        _norm = _normalize(_cals_raw)
        if _norm not in unconfirmed_set:
            continue
        # Look for PDF raw texts whose norm shares a common prefix with this value
        _candidates = []
        for _pdf_norm, _pdf_raw in pdf_raw_map.items():
            if not _pdf_norm or not _norm:
                continue
            _overlap = min(len(_pdf_norm), len(_norm), 6)
            if _overlap >= 4 and _pdf_norm[:_overlap] == _norm[:_overlap]:
                _candidates.append(f'"{_pdf_raw}"')
        if _candidates:
            reason_map[_norm] = (
                f"PDF extracted {', '.join(_candidates[:2])} — "
                f"normalised differently (likely split cell or line-break)"
            )
        else:
            reason_map[_norm] = (
                "Not found in PDF text extraction "
                "(merged cell, borderless table, or layout issue)"
            )

    # ------------------------------------------------------------------
    # 4. Build annotations — CALS XML with verify attributes, page image with colour boxes
    # ------------------------------------------------------------------
    annotated_xml = None
    annotated_image_path = None
    if in_cals_not_pdf:
        try:
            annotated_xml = _annotate_cals_xml(cals_xml, unconfirmed_set, reason_map=reason_map)
        except Exception as _ae:
            print(f"[verify_table] annotate_cals_xml failed: {_ae}")
    try:
        annotated_image_path = _annotate_page_image(
            img_path, pdf_path, page_idx, unconfirmed_set, confirmed_set
        )
    except Exception as _ae:
        print(f"[verify_table] annotate_page_image failed: {_ae}")

    report = {
        "table": table_title,
        "cals_shape": {"rows": cals_rows, "cols": cals_cols, "unique_values": len(cals_values)},
        "pdf_shape":  {"rows": pdf_rows_count, "cols": pdf_cols_count, "unique_values": len(pdf_values)},
        "confirmed_by_pdf": confirmed,
        "coverage_pct": f"{coverage:.0%}",
        "in_cals_not_pdf": in_cals_not_pdf[:20],
        "in_pdf_not_cals": in_pdf_not_cals[:20],
        "verdict": verdict,
        "annotated_xml": annotated_xml,
        "annotated_image_path": annotated_image_path,
    }
    print(f"[verify_table] '{table_title[:40]}' → {verdict} | "
          f"coverage={coverage:.0%} in_cals_not_pdf={len(in_cals_not_pdf)} "
          f"in_pdf_not_cals={len(in_pdf_not_cals)}")
    return report


# ---------------------------------------------------------------------------
# TEDS-based snapshot comparison
# ---------------------------------------------------------------------------

class _TableNode:
    """Lightweight node for the APTED tree: one cell or one structural tag."""

    def __init__(self, tag: str, content: str = "",
                 colspan: int = 1, rowspan: int = 1,
                 bold: bool = False, indent: int = 0):
        self.tag = tag          # "td", "tr", "table"
        self.content = content
        self.colspan = colspan
        self.rowspan = rowspan
        self.bold = bold
        self.indent = indent
        self.children: list = []

    # APTED requires a 'children' attribute and a string label.
    def __repr__(self):
        return (f"<{self.tag} cs={self.colspan} rs={self.rowspan} "
                f"b={int(self.bold)} i={self.indent}>")


def _cals_to_tree(cals_xml: str) -> "_TableNode":
    """Parse CALS XML into a _TableNode tree suitable for APTED comparison.

    Tree shape:
        table
          └─ tr  (one per <row>)
               └─ td  (one per <entry> in DOM order)

    Span information is taken from namest/nameend (colspan) and morerows (rowspan).
    Bold and indent come from the VLM-annotated attributes.
    """
    import xml.etree.ElementTree as ET

    root = ET.fromstring(cals_xml)

    # Build colspec index {colname: col_number} for colspan calculation
    colspec_order: list = []
    for cs in root.iter("colspec"):
        name = cs.get("colname", "")
        if name:
            colspec_order.append(name)
    col_index = {name: i for i, name in enumerate(colspec_order)}

    table_node = _TableNode("table")

    for row_el in root.iter("row"):
        tr_node = _TableNode("tr")
        for entry_el in row_el:
            if entry_el.tag != "entry":
                continue
            text = " ".join("".join(entry_el.itertext()).split())
            bold = entry_el.get("bold", "") == "true"
            try:
                indent = int(entry_el.get("indent", "0"))
            except ValueError:
                indent = 0

            # Colspan from namest/nameend + colspec_order
            namest = entry_el.get("namest", "")
            nameend = entry_el.get("nameend", "")
            if namest and nameend and namest in col_index and nameend in col_index:
                colspan = col_index[nameend] - col_index[namest] + 1
            else:
                colspan = 1

            # Rowspan from morerows attribute
            try:
                rowspan = int(entry_el.get("morerows", "0")) + 1
            except ValueError:
                rowspan = 1

            td_node = _TableNode(
                "td", content=text,
                colspan=max(colspan, 1), rowspan=max(rowspan, 1),
                bold=bold, indent=indent,
            )
            tr_node.children.append(td_node)
        if tr_node.children:
            table_node.children.append(tr_node)

    return table_node


def _teds_rename_cost(node_a: "_TableNode", node_b: "_TableNode",
                      w_content: float = 0.70,
                      w_bold: float = 0.15,
                      w_indent: float = 0.15) -> float:
    """Cost of renaming node_a into node_b for APTED.

    Structural mismatch (different tag or span) → cost 1.0 (maximum).
    For td nodes, partial costs for content, bold, and indent divergence.
    """
    if node_a.tag != node_b.tag:
        return 1.0
    if node_a.tag in ("table", "tr"):
        return 0.0   # structural containers — no rename cost beyond children

    # td node
    if node_a.colspan != node_b.colspan or node_a.rowspan != node_b.rowspan:
        return 1.0   # span changed — structural mismatch

    cost = 0.0

    # Content divergence via normalised Levenshtein
    ca, cb = node_a.content, node_b.content
    if ca or cb:
        try:
            from distance import levenshtein as _lev
            max_len = max(len(ca), len(cb), 1)
            cost += w_content * (_lev(ca, cb) / max_len)
        except ImportError:
            # Fallback: exact match
            cost += w_content * (0.0 if ca == cb else 1.0)

    if node_a.bold != node_b.bold:
        cost += w_bold

    if node_a.indent != node_b.indent:
        cost += w_indent * min(abs(node_a.indent - node_b.indent) / 3.0, 1.0)

    return min(cost, 1.0)


def _apted_distance(tree_a: "_TableNode", tree_b: "_TableNode",
                    rename_cost_fn) -> float:
    """Compute tree edit distance between tree_a and tree_b using APTED.

    Requires the `apted` package (`pip install apted`).
    Falls back to a naive size-difference estimate if unavailable.
    """
    try:
        from apted import APTED, Config

        class _APTEDConfig(Config):
            def rename(self, n1, n2):
                return rename_cost_fn(n1, n2)
            def children(self, node):
                return node.children

        apted_inst = APTED(tree_a, tree_b, _APTEDConfig())
        return apted_inst.compute_edit_distance()
    except ImportError:
        # Rough fallback: |size_A - size_B| normalised by max size
        def _size(n):
            return 1 + sum(_size(c) for c in n.children)
        sa, sb = _size(tree_a), _size(tree_b)
        return float(abs(sa - sb))


def _tree_size(node: "_TableNode") -> int:
    return 1 + sum(_tree_size(c) for c in node.children)


def compare_snapshots(xml_a: str, xml_b: str) -> dict:
    """Compare two CALS XML table snapshots using TEDS (Tree Edit Distance Similarity).

    Computes:
      - teds_full:    combined structure + content + bold/indent score (0–1)
      - teds_struct:  structure only — span grid and row order (0–1)
      - verdict:      "PASS" / "WARN" / "FAIL"

    Plus per-dimension diff lists:
      - lost_bold, gained_bold: cell texts where bold changed
      - indent_changed:         [(text, old_indent, new_indent), ...]
      - missing_values:         normalised values in A but not B
      - extra_values:           normalised values in B but not A
      - span_changed:           [(text, (cs_a,rs_a), (cs_b,rs_b)), ...]

    Args:
        xml_a: CALS XML of the reference snapshot (original document).
        xml_b: CALS XML of the re-transformed output.

    Returns a dict with all fields above, or a dict with "error" if parsing fails.
    """
    import re
    import xml.etree.ElementTree as ET

    def _norm(v: str) -> str | None:
        v = re.sub(r"[\$,\(\)\s\u2014\u2013]", "", str(v or ""))
        v = v.replace("-", "").strip()
        if not v:
            return None
        try:
            return str(float(v))
        except ValueError:
            return v.lower()

    # --- Parse both trees ---
    try:
        tree_a = _cals_to_tree(xml_a)
    except Exception as exc:
        return {"error": f"Failed to parse xml_a: {exc}"}
    try:
        tree_b = _cals_to_tree(xml_b)
    except Exception as exc:
        return {"error": f"Failed to parse xml_b: {exc}"}

    n_a = _tree_size(tree_a)
    n_b = _tree_size(tree_b)
    norm_factor = max(n_a, n_b, 1)

    # --- Full TEDS (structure + content + style) ---
    dist_full = _apted_distance(tree_a, tree_b, _teds_rename_cost)
    teds_full = max(0.0, 1.0 - dist_full / norm_factor)

    # --- Structural TEDS (spans + row order only) ---
    def _struct_cost(na, nb):
        if na.tag != nb.tag:
            return 1.0
        if na.tag == "td" and (na.colspan != nb.colspan or na.rowspan != nb.rowspan):
            return 1.0
        return 0.0

    dist_struct = _apted_distance(tree_a, tree_b, _struct_cost)
    teds_struct = max(0.0, 1.0 - dist_struct / norm_factor)

    # --- Per-cell diff (row-by-row, position-matched) ---
    lost_bold: list = []
    gained_bold: list = []
    indent_changed: list = []
    span_changed: list = []

    rows_a = tree_a.children   # list of tr nodes
    rows_b = tree_b.children

    for tr_a, tr_b in zip(rows_a, rows_b):
        for td_a, td_b in zip(tr_a.children, tr_b.children):
            label = td_a.content or td_b.content
            if td_a.bold and not td_b.bold:
                lost_bold.append(label)
            elif not td_a.bold and td_b.bold:
                gained_bold.append(label)
            if td_a.indent != td_b.indent:
                indent_changed.append((label, td_a.indent, td_b.indent))
            if td_a.colspan != td_b.colspan or td_a.rowspan != td_b.rowspan:
                span_changed.append((
                    label,
                    (td_a.colspan, td_a.rowspan),
                    (td_b.colspan, td_b.rowspan),
                ))

    # --- Value-set diff ---
    def _entry_values(cals_xml: str) -> set:
        root = ET.fromstring(cals_xml)
        vals = set()
        for entry in root.iter("entry"):
            n = _norm(" ".join("".join(entry.itertext()).split()))
            if n:
                vals.add(n)
        return vals

    try:
        vals_a = _entry_values(xml_a)
        vals_b = _entry_values(xml_b)
        missing_values = sorted(vals_a - vals_b)
        extra_values = sorted(vals_b - vals_a)
    except Exception:
        missing_values = []
        extra_values = []

    # --- Verdict ---
    if teds_full >= 0.95 and teds_struct >= 1.0:
        verdict = "PASS"
    elif teds_full >= 0.80 or teds_struct >= 0.95:
        verdict = "WARN"
    else:
        verdict = "FAIL"

    print(f"[compare_snapshots] teds_full={teds_full:.3f} teds_struct={teds_struct:.3f} "
          f"verdict={verdict} lost_bold={len(lost_bold)} indent_changed={len(indent_changed)}")

    return {
        "teds_full":      round(teds_full, 4),
        "teds_struct":    round(teds_struct, 4),
        "verdict":        verdict,
        "lost_bold":      lost_bold,
        "gained_bold":    gained_bold,
        "indent_changed": indent_changed,
        "missing_values": missing_values,
        "extra_values":   extra_values,
        "span_changed":   span_changed,
        "tree_sizes":     {"a": n_a, "b": n_b},
    }


def fop_teds_verify(cals_xml: str,
                   page_width_mm: float = None,
                   page_height_mm: float = None,
                   theme: str = "verify") -> dict:
    """Render *cals_xml* to PDF via Apache FOP, re-extract the table with
    pdfplumber, then compute a TEDS comparison between the original CALS XML
    and the re-extracted version.

    This is the canonical end-to-end regression check:
      original CALS  ─►  FOP PDF  ─►  pdfplumber  ─►  plain CALS  ─►  TEDS

    Returns the ``compare_snapshots`` dict extended with:
      ``fop_pdf_pages``   – int: number of pages in the rendered PDF
      ``fop_rows``        – int: rows pdfplumber recovered
      ``fop_cols``        – int: columns pdfplumber recovered
      ``error``           – str (only present on failure)
    """
    import base64
    import io
    import re
    import xml.etree.ElementTree as ET

    # --- 1. Render to FOP PDF ---
    try:
        pdf_b64, fop_err = _cals_to_fop_pdf(
            cals_xml, page_width_mm, page_height_mm, theme=theme
        )
    except Exception as exc:
        return {"error": f"FOP render failed: {exc}"}
    if fop_err:
        return {"error": f"FOP error: {fop_err[:400]}"}

    pdf_bytes = base64.b64decode(pdf_b64)

    # --- 2. Re-extract table from FOP PDF ---
    try:
        import pdfplumber
    except ImportError:
        return {"error": "pdfplumber not installed"}

    all_rows: list = []
    n_pages = 0
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            n_pages = len(pdf.pages)
            for pg in pdf.pages:
                tbl = pg.extract_table()
                if tbl:
                    all_rows.extend(tbl)
    except Exception as exc:
        return {"error": f"pdfplumber extraction failed: {exc}"}

    if not all_rows:
        return {
            "error": "pdfplumber found no table in the FOP PDF",
            "fop_pdf_pages": n_pages,
            "fop_rows": 0,
            "fop_cols": 0,
        }

    ncols = max(len(r) for r in all_rows)

    # --- 3. Build plain CALS from re-extracted rows ---
    def _esc(t: str) -> str:
        return (
            str(t or "")
            .strip()
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    lines = [f'<table><tgroup cols="{ncols}">']
    lines += [f'  <colspec colname="c{i + 1}"/>' for i in range(ncols)]
    lines += ["  <tbody>"]
    for row in all_rows:
        lines.append("    <row>")
        for cell in row:
            lines.append(f"      <entry>{_esc(cell)}</entry>")
        lines.append("    </row>")
    lines += ["  </tbody>", "</tgroup></table>"]
    xml_fop_reparsed = "\n".join(lines)

    # --- 4. TEDS comparison ---
    result = compare_snapshots(cals_xml, xml_fop_reparsed)
    result["fop_pdf_pages"] = n_pages
    result["fop_rows"]      = len(all_rows)
    result["fop_cols"]      = ncols

    print(
        f"[fop_teds_verify] pages={n_pages} rows={len(all_rows)} cols={ncols} "
        f"teds_full={result['teds_full']} teds_struct={result['teds_struct']} "
        f"verdict={result['verdict']}"
    )
    return result


def _extract_via_pdfplumber_lattice(pdf_path: str, page_idx: int):
    """Extract table rows using pdfplumber lattice (border-based) mode."""
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            if page_idx >= len(pdf.pages):
                return None
            tbl = pdf.pages[page_idx].extract_table()
            return [[str(c or "").strip() for c in row] for row in tbl] if tbl else None
    except Exception as exc:
        print(f"[pdfplumber-lattice] {exc}")
        return None


def _extract_via_pdfplumber_stream(pdf_path: str, page_idx: int):
    """Extract table rows using pdfplumber stream (whitespace) mode."""
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            if page_idx >= len(pdf.pages):
                return None
            tbl = pdf.pages[page_idx].extract_table(table_settings={
                "vertical_strategy": "text",
                "horizontal_strategy": "text",
                "intersection_x_tolerance": 15,
                "intersection_y_tolerance": 15,
            })
            return [[str(c or "").strip() for c in row] for row in tbl] if tbl else None
    except Exception as exc:
        print(f"[pdfplumber-stream] {exc}")
        return None


def _extract_via_pypdf_layout(pdf_path: str, page_idx: int):
    """Extract table rows using pypdf layout-preserving text extraction."""
    try:
        import re
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        if page_idx >= len(reader.pages):
            return None
        text = reader.pages[page_idx].extract_text(extraction_mode="layout")
        if not text:
            return None
        rows = []
        for line in text.splitlines():
            if not line.strip():
                continue
            cells = [c.strip() for c in re.split(r'  +', line) if c.strip()]
            if cells:
                rows.append(cells)
        return rows if rows else None
    except Exception as exc:
        print(f"[pypdf-layout] {exc}")
        return None


def optimize_extraction_agent(title: str, original_cell_rows: list,
                              pdf_path: str, page_idx: int) -> dict:
    """Iteratively tries 4 extraction strategies for a single table.

    Scoring: builds the union of all values found across all strategies, then
    scores each strategy by what percentage of that union it covers.  Avoids
    unreliable PASS/FAIL verdicts from pdfplumber (which truncates cell text).

    Returns a dict with 'title', 'recommended_strategy', and 'rounds'.
    """
    import re

    def _norm(v):
        return re.sub(r'[\s$,.()\-\u2014]', '', str(v or '')).lower()

    def _vset(rows):
        return {_norm(c) for row in rows for c in row if len(_norm(c)) >= 3}

    strategies = [("python-docx (current)", original_cell_rows)]
    if os.path.exists(pdf_path):
        for name, rows in [
            ("pdfplumber-lattice", _extract_via_pdfplumber_lattice(pdf_path, page_idx)),
            ("pdfplumber-stream",  _extract_via_pdfplumber_stream(pdf_path, page_idx)),
            ("pypdf-layout",       _extract_via_pypdf_layout(pdf_path, page_idx)),
        ]:
            if rows:
                strategies.append((name, rows))

    # Build union of all distinct values found across every strategy
    strategy_vsets = {name: _vset(rows) for name, rows in strategies}
    union = set().union(*strategy_vsets.values())

    rounds = []
    for name, rows in strategies:
        vs = strategy_vsets[name]
        coverage = len(vs & union) / max(len(union), 1)
        missing_from_this = sorted(union - vs)[:10]
        rounds.append({
            "strategy": name,
            "shape": [len(rows), max((len(r) for r in rows), default=0)],
            "unique_values_found": len(vs),
            "union_coverage": f"{coverage:.0%}",
            "sample_values_missing_from_this_strategy": missing_from_this,
        })
        print(f"[optimize] '{name}': {len(vs)} values, {coverage:.0%} coverage")

    best_idx = max(range(len(rounds)), key=lambda i: rounds[i]["unique_values_found"])
    rounds[best_idx]["recommended"] = True

    return {
        "title": title,
        "total_unique_values_across_strategies": len(union),
        "recommended_strategy": rounds[best_idx]["strategy"],
        "rounds": rounds,
    }

def load_documents_from_files(file_paths: List[str], on_progress=None) -> List[Any]:
    """Load and return documents from supported file types.

    The returned list may carry a ``_doc_sequence`` attribute when a .docx file
    was processed — callers can use it to reconstruct the document text order.
    """
    docs = []
    doc_sequence = None  # captured from the first .docx that provides it

    for fpath in file_paths:
        ext = os.path.splitext(fpath)[-1].lower()

        try:
            if ext == ".pdf":
                loader = UnstructuredPDFLoader(fpath, mode="elements")
                loaded = loader.load()
            elif ext in [".docx", ".doc"]:
                # Use python-docx directly — unstructured is incompatible with this docx version
                loaded = _load_docx_direct(fpath, on_progress=on_progress)
                # Capture the document sequence (paragraphs + table placeholders in order)
                if doc_sequence is None:
                    doc_sequence = getattr(loaded, "_doc_sequence", None)
            elif ext in [".txt", ".md"]:
                loader = TextLoader(fpath)
                loaded = loader.load()
            elif ext == ".csv":
                loader = CSVLoader(fpath)
                loaded = loader.load()
            else:
                print(f"[load_documents] Skipping unsupported file type: {fpath}")
                continue

            docs.append(loaded)
            print(f"[load_documents] Successfully loaded {fpath} ({len(loaded)} elements)")
        except Exception as e:
            print(f"[load_documents] Failed to load {fpath}: {e}")

    class _DocList(list):
        """list subclass that can carry extra attributes (e.g. _doc_sequence)."""
        pass

    flat = _DocList(item for sublist in docs for item in sublist)
    if doc_sequence is not None:
        flat._doc_sequence = doc_sequence
    return flat


def split_documents(docs: List[Any]):
    """Split documents into smaller chunks using recursive splitter."""
    print(f"[split_documents] Splitting {len(docs)} docs with chunk size {CHUNK_SIZE}, overlap {CHUNK_OVERLAP}")

    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )
    return splitter.split_documents(docs)

def embed_documents(doc_splits: List[Any]):
    """Embed and store the split documents into Chroma vectorstore."""
    try:
        print(f"[embed_documents] Embedding {len(doc_splits)} chunks using model: {EMBEDDINGS_MODEL}")

        vectorstore = Chroma.from_documents(
            documents=doc_splits,
            collection_name="rag-chroma",
            embedding=HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL),
            persist_directory=DATA_DIR,
        )
        return vectorstore

    except Exception as e:
        print(f"[embed_documents] Vectorstore creation failed: {e}")
        return None


def reconstruct_document(
    table_catalog: list = None,
    segments_path: str = None,
    output_format: str = "text",
) -> str:
    """Reconstruct the full document text by merging transformed CALS XML back into
    the original paragraph sequence.

    Duplicate-elimination — Layer 2 (reconstruction time)
    ------------------------------------------------------
    ``doc_segments.json`` is the single source of truth produced by
    ``_load_docx_direct`` (Layer 1 already removed caption-paragraph duplicates).
    Reconstruction is a single ordered traversal of that file:

    * ``paragraph`` segments → emitted verbatim. Because Layer 1 never writes a
      caption as a standalone paragraph segment, no title text appears here.
    * ``table`` segments → emitted as ``[Title: …]\n<CALS XML>``.  The title
      comes from ``seg["title"]`` (the caption consumed at parse time), and the
      CALS XML is looked up from the catalog by ``insertion_key``.

    There is no secondary scan of the paragraph list or the catalog, so no
    segment can appear more than once.  Correctness is guaranteed by the
    uniqueness of ``seq`` values and the one-pass loop structure.

    Parameters
    ----------
    table_catalog : list, optional
        List of catalog dicts (as stored in table_catalog.json).  When provided,
        each table's ``insertion_key`` is substituted with its current ``xml``
        (CALS format).  When ``None`` the catalog is loaded from disk.
    segments_path : str, optional
        Path to doc_segments.json.  Defaults to ``data/doc_segments.json``.
    output_format : str
        ``"text"``  — plain text with CALS XML blocks inline (default)
        ``"xml"``   — wraps everything in a ``<document>`` root; paragraphs become
                      ``<para>`` elements and tables keep their CALS ``<table>`` root

    Returns
    -------
    str
        The reconstructed document as a string, tables at their original positions,
        original paragraph text preserved in sequence order.
    """
    import json as _json

    if segments_path is None:
        segments_path = os.path.join(DATA_DIR, "doc_segments.json")

    if not os.path.exists(segments_path):
        return ""

    with open(segments_path, encoding="utf-8") as _f:
        segments = _json.load(_f)

    if table_catalog is None:
        catalog_path = os.path.join(DATA_DIR, "table_catalog.json")
        if os.path.exists(catalog_path):
            with open(catalog_path, encoding="utf-8") as _f:
                table_catalog = _json.load(_f)
        else:
            table_catalog = []

    # Build a lookup: insertion_key → current xml (may have been updated by agent/VLM)
    key_to_xml: dict = {}
    for entry in table_catalog:
        ik = entry.get("insertion_key", "")
        if ik:
            key_to_xml[ik] = entry.get("xml") or entry.get("xml_pdfplumber") or ""

    # Sort segments by original body order
    segments_sorted = sorted(segments, key=lambda s: s.get("seq", 0))

    if output_format == "xml":
        parts = ["<document>"]
        for seg in segments_sorted:
            if seg["type"] == "paragraph":
                import html as _html
                parts.append(f"  <para>{_html.escape(seg['text'])}</para>")
            elif seg["type"] == "table":
                ik = seg.get("insertion_key", "")
                xml_block = key_to_xml.get(ik, "")
                if xml_block:
                    # Indent each line of the CALS block for readability
                    indented = "\n".join("  " + ln for ln in xml_block.splitlines())
                    parts.append(indented)
                else:
                    parts.append(f"  <!-- TABLE PLACEHOLDER: {ik} -->")
        parts.append("</document>")
        return "\n".join(parts)
    else:
        parts = []
        for seg in segments_sorted:
            if seg["type"] == "paragraph":
                parts.append(seg["text"])
            elif seg["type"] == "table":
                ik = seg.get("insertion_key", "")
                xml_block = key_to_xml.get(ik, "")
                title = seg.get("title", "")
                header = f"\n--- TABLE: {title} ---\n" if title else "\n--- TABLE ---\n"
                if xml_block:
                    parts.append(f"{header}{xml_block}\n--- END TABLE ---")
                else:
                    parts.append(f"{header}[No XML available for {ik}]\n--- END TABLE ---")
        return "\n\n".join(parts)


def _fo_esc(text: str) -> str:
    """Escape a string for use as XML / XSL-FO text content."""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


def _entry_to_fo_cell(entry, fo: dict, is_head: bool = False) -> str:
    """Convert a CALS <entry> element to an fo:table-cell XML string."""
    import re as _re

    # ── column / row spans ────────────────────────────────────────────────────
    span_attrs = ""
    namest  = entry.get("namest", "")
    nameend = entry.get("nameend", "")
    if namest and nameend:
        try:
            s = int(_re.sub(r"^[^0-9]*", "", namest) or "1")
            e = int(_re.sub(r"^[^0-9]*", "", nameend) or "1")
            if e > s:
                span_attrs += f' number-columns-spanned="{e - s + 1}"'
        except ValueError:
            pass
    morerows = entry.get("morerows", "")
    if morerows:
        try:
            mr = int(morerows)
            if mr > 0:
                span_attrs += f' number-rows-spanned="{mr + 1}"'
        except ValueError:
            pass

    # ── vertical alignment ────────────────────────────────────────────────────
    valign = {"top": "before", "bottom": "after"}.get(entry.get("valign", ""), "center")

    if is_head:
        bg, border = fo["head_bg"], fo["head_cell_border"]
    else:
        bg, border = fo["bg_default"], fo["body_cell_border"]

    cell_attrs = (
        f'border="{border}" background-color="{bg}" padding="3pt"'
        f' display-align="{valign}"{span_attrs}'
    )

    align     = entry.get("align", "")
    align_attr = f' text-align="{align}"' if align else ""

    if is_head:
        block_attrs = (
            f'font-family="{fo["font_family"]}" font-size="{fo["font_size"]}"'
            f' font-weight="{fo["head_font_weight"]}" color="{fo["head_color"]}"'
            + align_attr
        )
    else:
        fw         = "bold" if entry.get("bold") == "true" else "normal"
        indent_mm  = {"1": "6", "2": "12", "3": "18"}.get(entry.get("indent", ""), "")
        indent_attr = f' start-indent="{indent_mm}mm"' if indent_mm else ""
        block_attrs = (
            f'font-family="{fo["font_family"]}" font-size="{fo["font_size"]}"'
            f' font-weight="{fw}" color="{fo["body_color"]}"'
            + align_attr + indent_attr
        )

    # ── cell content ──────────────────────────────────────────────────────────
    paras = entry.findall("para")
    if paras:
        pieces = []
        for i, p in enumerate(paras):
            pieces.append(_fo_esc(p.text or ""))
            if i < len(paras) - 1:
                pieces.append("<fo:block/>")
        content = "".join(pieces)
    else:
        content = _fo_esc(entry.text or "")

    return (
        f'<fo:table-cell {cell_attrs}>'
        f'<fo:block {block_attrs}>{content}</fo:block>'
        f'</fo:table-cell>'
    )


def _cals_xml_to_fo_block(cals_xml: str, fo: dict) -> str:
    """Convert a CALS <table> XML string to XSL-FO markup (fo:block wrapping fo:table)."""
    import xml.etree.ElementTree as _ET
    if not cals_xml or not cals_xml.strip():
        return ""
    try:
        xml_str = cals_xml.strip()
        if not xml_str.startswith("<table"):
            xml_str = f"<table>{xml_str}</table>"
        root = _ET.fromstring(xml_str)
    except _ET.ParseError as exc:
        return (
            f'<fo:block color="#c00" font-family="monospace" space-before="4pt">'
            f'[Table XML error: {_fo_esc(str(exc))}]</fo:block>'
        )

    for entry in root.iter("entry"):
        entry.attrib.pop("verify", None)
        entry.attrib.pop("verify-reason", None)
    _sanitize_rowspans(root)

    parts = ['<fo:block space-before="8pt" space-after="8pt">']

    title_el = root.find("title")
    if title_el is not None and title_el.text:
        parts.append(
            f'<fo:block font-family="{fo["caption_font"]}" font-size="{fo["caption_size"]}"'
            f' font-weight="bold" space-after="4pt">{_fo_esc(title_el.text)}</fo:block>'
        )

    for tgroup in root.findall("tgroup"):
        try:
            cols = int(tgroup.get("cols", "1") or "1")
        except (ValueError, TypeError):
            cols = 1

        parts.append(
            f'<fo:table table-layout="fixed" width="100%"'
            f' border-collapse="collapse" border="{fo["table_border"]}">'
        )

        colspecs = tgroup.findall("colspec")
        if colspecs:
            for cs in colspecs:
                cw = cs.get("colwidth", "") or ""
                if cw.endswith("*"):
                    cw = f"proportional-column-width({cw[:-1]})"
                if not cw:
                    cw = "proportional-column-width(1)"
                parts.append(f'<fo:table-column column-width="{cw}"/>')
        else:
            for _ in range(max(cols, 1)):
                parts.append('<fo:table-column column-width="proportional-column-width(1)"/>')

        thead = tgroup.find("thead")
        if thead is not None:
            parts.append("<fo:table-header>")
            for row in thead.findall("row"):
                parts.append("<fo:table-row>")
                for entry in row.findall("entry"):
                    parts.append(_entry_to_fo_cell(entry, fo, is_head=True))
                parts.append("</fo:table-row>")
            parts.append("</fo:table-header>")

        tbody = tgroup.find("tbody")
        parts.append("<fo:table-body>")
        rows = tbody.findall("row") if tbody is not None else []
        for row in rows:
            parts.append("<fo:table-row>")
            for entry in row.findall("entry"):
                parts.append(_entry_to_fo_cell(entry, fo, is_head=False))
            parts.append("</fo:table-row>")
        if not rows:
            # FOP requires at least one row in fo:table-body
            parts.append(
                '<fo:table-row><fo:table-cell><fo:block/></fo:table-cell></fo:table-row>'
            )
        parts.append("</fo:table-body>")
        parts.append("</fo:table>")

    parts.append("</fo:block>")
    return "\n".join(parts)


def _docx_to_fop_pdf(docx_path: str, key_to_xml: dict, fo: dict, env: dict) -> tuple:
    """Read a .docx and render a full-document A4 PDF via Apache FOP.

    Walks the docx body in order; for each paragraph emits an fo:block preserving
    the Word style (heading level, alignment, bold/italic runs, space-before/after,
    left-indent); for each table emits the CALS fo:table via _cals_xml_to_fo_block.

    Returns (b64_pdf, error_msg).
    """
    import subprocess, tempfile, base64 as _b64, os as _os

    try:
        from docx import Document as _DocxDoc
    except ImportError:
        return None, "python-docx not available."
    try:
        docx = _DocxDoc(docx_path)
    except Exception as exc:
        return None, f"Failed to open {docx_path}: {exc}"

    # WD_ALIGN_PARAGRAPH → fo text-align value
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_AL
        _ALIGN = {
            _WD_AL.LEFT:    "left",
            _WD_AL.CENTER:  "center",
            _WD_AL.RIGHT:   "right",
            _WD_AL.JUSTIFY: "justify",
        }
    except Exception:
        _ALIGN = {}

    # Word style name → extra FO block properties (override theme defaults)
    _PARA_STYLES = {
        "Title":     {"font-size": "18pt", "font-weight": "bold",
                      "text-align": "center", "space-before": "14pt", "space-after": "10pt"},
        "Subtitle":  {"font-size": "14pt", "font-style": "italic",
                      "text-align": "center", "space-before": "6pt",  "space-after": "8pt"},
        "Heading 1": {"font-size": "14pt", "font-weight": "bold",
                      "space-before": "14pt", "space-after": "6pt",  "keep-with-next": "always"},
        "Heading 2": {"font-size": "12pt", "font-weight": "bold",
                      "space-before": "12pt", "space-after": "4pt",  "keep-with-next": "always"},
        "Heading 3": {"font-size": "11pt", "font-weight": "bold",
                      "space-before": "10pt", "space-after": "4pt",  "keep-with-next": "always"},
        "Heading 4": {"font-size": "10pt", "font-weight": "bold",
                      "space-before": "8pt",  "space-after": "2pt",  "keep-with-next": "always"},
        "Heading 5": {"font-size": "10pt", "font-weight": "bold",   "font-style": "italic",
                      "space-before": "6pt"},
        "Heading 6": {"font-size": "10pt", "font-style": "italic",  "space-before": "4pt"},
    }

    # Index body children
    para_map  = {p._element: p for p in docx.paragraphs}
    table_map = {t._element: t for t in docx.tables}  # noqa: F841

    fo_blocks    = []
    table_counter = 0

    for child in docx.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            para = para_map.get(child)
            if para is None:
                continue
            text = (para.text or "").strip()
            if not text:
                fo_blocks.append('<fo:block space-after="2pt"/>')
                continue

            style_name = (para.style.name if para.style else None) or "Normal"

            props = {
                "font-family": fo["font_family"],
                "font-size":   fo["font_size"],
                "color":       fo["body_color"],
                "line-height": "1.4",
                "space-after": "4pt",
            }
            if style_name in _PARA_STYLES:
                props.update(_PARA_STYLES[style_name])

            # Alignment from paragraph object
            try:
                al = _ALIGN.get(para.alignment)
                if al:
                    props["text-align"] = al
            except Exception:
                pass

            # Space before / after from paragraph_format (beats style default)
            try:
                pf = para.paragraph_format
                if pf.space_before is not None and pf.space_before.pt:
                    props["space-before"] = f"{pf.space_before.pt:.0f}pt"
                if pf.space_after is not None and pf.space_after.pt:
                    props["space-after"]  = f"{pf.space_after.pt:.0f}pt"
            except Exception:
                pass

            # Left indent
            try:
                pf = para.paragraph_format
                if pf.left_indent and pf.left_indent.pt > 0:
                    props["start-indent"] = f"{pf.left_indent.pt * 25.4 / 72:.1f}mm"
            except Exception:
                pass

            block_attrs = " ".join(f'{k}="{v}"' for k, v in props.items())

            # Build content from runs (per-run bold / italic / size)
            run_parts = []
            try:
                for run in para.runs:
                    if not run.text:
                        continue
                    rt = _fo_esc(run.text)
                    inl = []
                    if run.bold:
                        inl.append('font-weight="bold"')
                    elif run.bold is False and "font-weight" not in props:
                        inl.append('font-weight="normal"')
                    if run.italic:
                        inl.append('font-style="italic"')
                    elif run.italic is False and "font-style" not in props:
                        inl.append('font-style="normal"')
                    if run.font.size:
                        inl.append(f'font-size="{run.font.size.pt:.0f}pt"')
                    if inl:
                        run_parts.append(f'<fo:inline {" ".join(inl)}>{rt}</fo:inline>')
                    else:
                        run_parts.append(rt)
            except Exception:
                run_parts = [_fo_esc(text)]

            content = "".join(run_parts) or _fo_esc(text)
            fo_blocks.append(f"<fo:block {block_attrs}>{content}</fo:block>")

        elif tag == "tbl":
            ik = "{{TABLE:t%d}}" % table_counter
            table_counter += 1
            cals_xml = key_to_xml.get(ik, "")
            if cals_xml:
                fo_blocks.append(_cals_xml_to_fo_block(cals_xml, fo))
            else:
                fo_blocks.append(
                    f'<fo:block color="#888" font-family="monospace" font-size="9pt"'
                    f' space-before="8pt" space-after="8pt">'
                    f'[Table {table_counter - 1}: XML not available]</fo:block>'
                )

    # ── Assemble fo:root ──────────────────────────────────────────────────────
    fo_doc = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<fo:root xmlns:fo="http://www.w3.org/1999/XSL/Format">\n'
        '  <fo:layout-master-set>\n'
        '    <fo:simple-page-master master-name="a4"\n'
        '        page-width="210mm" page-height="297mm"\n'
        '        margin-top="20mm" margin-bottom="20mm"\n'
        '        margin-left="20mm" margin-right="20mm">\n'
        '      <fo:region-body/>\n'
        '    </fo:simple-page-master>\n'
        '  </fo:layout-master-set>\n'
        '  <fo:page-sequence master-reference="a4">\n'
        '    <fo:flow flow-name="xsl-region-body">\n'
        + "\n".join(fo_blocks) + "\n"
        '    </fo:flow>\n'
        '  </fo:page-sequence>\n'
        '</fo:root>\n'
    )

    with tempfile.TemporaryDirectory() as tmp:
        fo_path  = _os.path.join(tmp, "document.fo")
        pdf_path = _os.path.join(tmp, "output.pdf")

        with open(fo_path, "w", encoding="utf-8") as f:
            f.write(fo_doc)

        fop_cmd = ["fop", fo_path, pdf_path]
        try:
            r = subprocess.run(fop_cmd, capture_output=True, text=True, env=env, timeout=180)
        except FileNotFoundError:
            return None, "fop not found on PATH. Install with: apt install fop"
        except subprocess.TimeoutExpired:
            return None, "FOP timed out after 180 s."

        stderr = r.stderr or ""
        errors = [l for l in stderr.splitlines()
                  if "ERROR" in l or "Exception" in l or "FATAL" in l]

        if not _os.path.exists(pdf_path) or _os.path.getsize(pdf_path) < 100:
            return None, "FOP failed:\n" + ("\n".join(errors) or stderr[:500])

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return _b64.b64encode(pdf_bytes).decode(), ""


def reconstruct_to_pdf(
    table_catalog: list = None,
    segments_path: str = None,
    theme: str = "finance",
) -> tuple:
    """Render the full reconstructed document as an A4 PDF via Apache FOP.

    Primary path: re-reads the original .docx (path stored in the catalog's
    ``source`` field) with python-docx to recover full paragraph styling — heading
    levels, alignment, bold/italic runs, space-before/after, left-indent — and
    generates XSL-FO directly.  CALS tables are rendered inline via
    ``_cals_xml_to_fo_block`` at the exact position they appear in the docx body.

    Fallback (docx not on disk): builds a plain ``<document>`` XML via
    ``reconstruct_document`` and applies ``_DOCUMENT_TO_FO_XSLT``.

    Parameters
    ----------
    table_catalog : list, optional
        Catalog list (as in table_catalog.json).  Loaded from disk when None.
    segments_path : str, optional
        Path to doc_segments.json.  Defaults to ``data/doc_segments.json``.
    theme : str
        One of the keys in ``_CALS_FO_THEMES``.  Defaults to ``"finance"``.

    Returns
    -------
    (b64_pdf : str | None, error_msg : str)
    """
    import subprocess, tempfile, base64, xml.etree.ElementTree as ET
    import os as _os

    # ── Load catalog ──────────────────────────────────────────────────────────
    if table_catalog is None:
        catalog_path = os.path.join(DATA_DIR, "table_catalog.json")
        if os.path.exists(catalog_path):
            import json as _json
            with open(catalog_path, encoding="utf-8") as _f:
                table_catalog = _json.load(_f)
        else:
            table_catalog = []

    # Build insertion_key → CALS XML lookup
    key_to_xml: dict = {}
    for entry in (table_catalog or []):
        ik = entry.get("insertion_key", "")
        if ik:
            key_to_xml[ik] = entry.get("xml") or entry.get("xml_pdfplumber") or ""

    if not key_to_xml:
        return None, "No document content found. Upload a Word file first."

    # ── Java 11 detection (FOP 2.8 requirement) ───────────────────────────────
    env = _os.environ.copy()
    for _jh in ("/usr/lib/jvm/java-11-openjdk-arm64",
                "/usr/lib/jvm/java-11-openjdk-amd64"):
        if _os.path.isdir(_jh):
            env["JAVA_HOME"] = _jh
            break

    fo = _CALS_FO_THEMES.get(theme, _CALS_FO_THEMES["finance"])

    # ── Primary path: re-read original .docx for full style fidelity ─────────
    docx_path = None
    for entry in (table_catalog or []):
        src = entry.get("source", "")
        if src and src.lower().endswith(".docx") and _os.path.exists(src):
            docx_path = src
            break

    if docx_path:
        return _docx_to_fop_pdf(docx_path, key_to_xml, fo, env)

    # ── Fallback: XSLT approach from plain <document> XML ────────────────────
    doc_xml = reconstruct_document(
        table_catalog=table_catalog,
        segments_path=segments_path,
        output_format="xml",
    )
    if not doc_xml.strip():
        return None, "No document content found. Upload a Word file first."

    try:
        root = ET.fromstring(doc_xml)
        for entry in root.iter("entry"):
            entry.attrib.pop("verify", None)
            entry.attrib.pop("verify-reason", None)
        _sanitize_rowspans(root)
        doc_xml_clean = ET.tostring(root, encoding="unicode")
    except ET.ParseError as exc:
        return None, f"Failed to parse reconstructed XML: {exc}"

    xslt_str = (
        _DOCUMENT_TO_FO_XSLT
        .replace("__FO_TABLE_BORDER__",     fo["table_border"])
        .replace("__FO_HEAD_CELL_BORDER__",  fo["head_cell_border"])
        .replace("__FO_HEAD_BG__",           fo["head_bg"])
        .replace("__FO_FONT_FAMILY__",       fo["font_family"])
        .replace("__FO_FONT_SIZE__",         fo["font_size"])
        .replace("__FO_HEAD_FONT_WEIGHT__",  fo["head_font_weight"])
        .replace("__FO_HEAD_COLOR__",        fo["head_color"])
        .replace("__FO_BODY_CELL_BORDER__",  fo["body_cell_border"])
        .replace("__FO_BODY_COLOR__",        fo["body_color"])
        .replace("__FO_BG_DEFAULT__",        fo["bg_default"])
        .replace("__FO_CAPTION_FONT__",      fo["caption_font"])
        .replace("__FO_CAPTION_SIZE__",      fo["caption_size"])
    )

    with tempfile.TemporaryDirectory() as tmp:
        xslt_path = _os.path.join(tmp, "doc2fo.xsl")
        xml_path  = _os.path.join(tmp, "document.xml")
        pdf_path  = _os.path.join(tmp, "output.pdf")

        with open(xslt_path, "w", encoding="utf-8") as f:
            f.write(xslt_str)
        with open(xml_path, "w", encoding="utf-8") as f:
            f.write('<?xml version="1.0" encoding="UTF-8"?>\n' + doc_xml_clean)

        fop_cmd = ["fop", "-xsl", xslt_path, "-xml", xml_path, "-pdf", pdf_path]
        try:
            r = subprocess.run(fop_cmd, capture_output=True, text=True, env=env, timeout=120)
        except FileNotFoundError:
            return None, "fop not found on PATH. Install with: apt install fop"
        except subprocess.TimeoutExpired:
            return None, "FOP timed out after 120 s."

        stderr = r.stderr or ""
        errors = [l for l in stderr.splitlines()
                  if "ERROR" in l or "Exception" in l or "FATAL" in l]

        if not _os.path.exists(pdf_path) or _os.path.getsize(pdf_path) < 100:
            return None, "FOP failed:\n" + ("\n".join(errors) or stderr[:500])

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return base64.b64encode(pdf_bytes).decode(), ""


def _set_tc_shading(tc, fill_hex: str) -> None:
    """Set background fill colour on a table cell's oxml element (CT_Tc).

    Parameters
    ----------
    tc : docx.oxml.table.CT_Tc
        The ``cell._tc`` element to shade.
    fill_hex : str
        Six-hex-digit colour without ``#``, e.g. ``"EBF3FB"``.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlEl
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(_qn("w:shd")):
        tcPr.remove(existing)
    shd = _OxmlEl("w:shd")
    shd.set(_qn("w:val"),   "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"),  fill_hex.lstrip("#").upper())
    tcPr.append(shd)


def write_back_to_docx(
    table_catalog: list = None,
    docx_src_path: str = None,
    output_path: str = None,
) -> tuple:
    """Write edited CALS table XML back into the original .docx file.

    For each table in the catalog, finds the matching table in the docx
    (by ``table_index`` / ``doc_seq`` order — docx table N corresponds to
    ``{{TABLE:tN}}``) and overwrites cell text with the current CALS XML.

    Only cell *text* is updated.  Existing cell formatting (font, shading,
    borders) is preserved where possible; per-run text is replaced cleanly.

    Limitations
    -----------
    * Structural changes (adding / removing rows or columns) require the CALS
      row/column count to match the original Word table exactly.  If they
      differ, the affected table is skipped and noted in the status message.
    * Merged cells: writing to a merged cell updates the master cell; the
      other spanning cells in the merge are left empty (Word behavior).

    Parameters
    ----------
    table_catalog : list, optional
        Loaded from ``data/table_catalog.json`` when None.
    docx_src_path : str, optional
        Path to the source ``.docx``.  Auto-detected from catalog when None.
    output_path : str, optional
        Where to save the result.  Defaults to a temp file.

    Returns
    -------
    (output_path : str | None, status_msg : str)
    """
    import json as _json, copy as _copy, tempfile as _tmp, os as _os

    # ── load catalog ──────────────────────────────────────────────────────────
    if table_catalog is None:
        cat_path = os.path.join(DATA_DIR, "table_catalog.json")
        if not os.path.exists(cat_path):
            return None, "No table catalog found."
        with open(cat_path, encoding="utf-8") as _f:
            table_catalog = _json.load(_f)

    if not table_catalog:
        return None, "Table catalog is empty."

    # ── find source .docx ─────────────────────────────────────────────────────
    if docx_src_path is None:
        for entry in table_catalog:
            src = entry.get("source", "")
            if src and src.lower().endswith(".docx") and _os.path.exists(src):
                docx_src_path = src
                break

    if not docx_src_path or not _os.path.exists(docx_src_path):
        return None, "Original .docx file not found on disk."

    try:
        from docx import Document as _DocxDoc
        import xml.etree.ElementTree as _ET
    except ImportError as exc:
        return None, f"Missing dependency: {exc}"

    try:
        docx = _DocxDoc(docx_src_path)
    except Exception as exc:
        return None, f"Failed to open {docx_src_path}: {exc}"

    # ── build table_index → catalog entry map ─────────────────────────────────
    idx_to_entry: dict = {}
    for entry in table_catalog:
        ti = entry.get("table_index")
        if ti is not None:
            idx_to_entry[int(ti)] = entry

    skipped   = []
    updated   = []

    docx_tables = docx.tables

    for ti, entry in sorted(idx_to_entry.items()):
        if ti >= len(docx_tables):
            skipped.append(f"t{ti}: index out of range (docx has {len(docx_tables)} tables)")
            continue

        cals_xml = entry.get("xml") or entry.get("xml_pdfplumber") or ""
        if not cals_xml.strip():
            skipped.append(f"t{ti}: no CALS XML available")
            continue

        # Parse CALS XML → cell text grid
        try:
            xml_str = cals_xml.strip()
            if not xml_str.startswith("<table"):
                xml_str = f"<table>{xml_str}</table>"
            root_el = _ET.fromstring(xml_str)
        except _ET.ParseError as exc:
            skipped.append(f"t{ti}: XML parse error — {exc}")
            continue

        # Collect row/cell texts from all tgroups in order
        cals_rows = []
        for tgroup in root_el.findall("tgroup"):
            for section in (tgroup.find("thead"), tgroup.find("tbody")):
                if section is None:
                    continue
                for row_el in section.findall("row"):
                    row_texts = []
                    for entry_el in row_el.findall("entry"):
                        # Collect all text inside <entry> (including <para> children)
                        cell_text = "".join(entry_el.itertext()).strip()
                        row_texts.append(cell_text)
                    cals_rows.append(row_texts)

        docx_table = docx_tables[ti]

        # Deduplicate merged cells in docx rows (same _tc element appears multiple times)
        def _dedup_row(row):
            seen, cells = set(), []
            for c in row.cells:
                cid = id(c._tc)
                if cid not in seen:
                    seen.add(cid)
                    cells.append(c)
            return cells

        docx_rows = [_dedup_row(r) for r in docx_table.rows]

        if len(docx_rows) != len(cals_rows):
            skipped.append(
                f"t{ti} '{entry.get('title','')}': "
                f"row count mismatch (docx={len(docx_rows)}, cals={len(cals_rows)}) — skipped"
            )
            continue

        col_mismatch = False
        for r_idx, (d_row, c_row) in enumerate(zip(docx_rows, cals_rows)):
            if len(d_row) != len(c_row):
                skipped.append(
                    f"t{ti} row {r_idx}: col count mismatch "
                    f"(docx={len(d_row)}, cals={len(c_row)}) — table skipped"
                )
                col_mismatch = True
                break

        if col_mismatch:
            continue

        # Write cell text back, preserving run formatting
        for d_row, c_row in zip(docx_rows, cals_rows):
            for cell, new_text in zip(d_row, c_row):
                # Strategy: clear all paragraphs except the first, then update
                # the first paragraph's runs while keeping their character style.
                paras = cell.paragraphs
                # Remove extra paragraphs (beyond first)
                for extra_para in paras[1:]:
                    p_el = extra_para._element
                    p_el.getparent().remove(p_el)

                first_para = cell.paragraphs[0]
                runs = first_para.runs
                if runs:
                    # Update first run, clear the rest
                    runs[0].text = new_text
                    for extra_run in runs[1:]:
                        extra_run.text = ""
                else:
                    # No runs — add one
                    first_para.add_run(new_text)

        # ── apply per-table style transform (e.g. alternating row stripe) ─────
        transform = entry.get("transform") or {}
        if transform.get("type") == "stripe":
            stripe_hex = (transform.get("color") or "EBF3FB").lstrip("#").upper()
            even_hex   = "FFFFFF"
            for r_idx, docx_row in enumerate(docx_table.rows):
                fill = stripe_hex if r_idx % 2 == 1 else even_hex
                seen_tcs: set = set()
                for cell in docx_row.cells:
                    cid = id(cell._tc)
                    if cid not in seen_tcs:
                        seen_tcs.add(cid)
                        _set_tc_shading(cell._tc, fill)

        updated.append(f"t{ti} '{entry.get('title',''[:30])}'")

    if not updated:
        return None, "No tables were updated." + (
            (" Issues: " + "; ".join(skipped)) if skipped else ""
        )

    # ── save ──────────────────────────────────────────────────────────────────
    if output_path is None:
        tf = _tmp.NamedTemporaryFile(
            delete=False, suffix=".docx", prefix="document_edited_"
        )
        tf.close()
        output_path = tf.name

    try:
        docx.save(output_path)
    except Exception as exc:
        return None, f"Failed to save docx: {exc}"

    status = f"✅ Updated {len(updated)} table(s): {', '.join(updated)}."
    if skipped:
        status += f" Skipped: {'; '.join(skipped)}."
    return output_path, status


def compute_export_preview(table_catalog: list = None) -> dict:
    """Build a summary dict for the export review panel.

    Returns a dict with:
      ``text_stats``  — char/word/whitespace/other counts from doc_segments.json
      ``table_rows``  — per-table list with TEDS vs pdfplumber baseline + transform info
      ``total_tables``, ``edited_tables``, ``transformed_tables`` — aggregate counts
    """
    import json as _json

    if table_catalog is None:
        cat_path = os.path.join(DATA_DIR, "table_catalog.json")
        if not os.path.exists(cat_path):
            return {"error": "No catalog found"}
        with open(cat_path, encoding="utf-8") as _f:
            table_catalog = _json.load(_f)

    if not table_catalog:
        return {"error": "Catalog is empty"}

    # ── text segments ─────────────────────────────────────────────────────────
    text_stats = {
        "seg_count": 0, "char_count": 0,
        "word_count": 0, "ws_count": 0, "other_count": 0,
    }
    seg_path = os.path.join(DATA_DIR, "doc_segments.json")
    try:
        if os.path.exists(seg_path):
            with open(seg_path, encoding="utf-8") as _f:
                segments = _json.load(_f)
            for seg in segments:
                if seg.get("type") != "paragraph":
                    continue
                txt = seg.get("text") or ""
                text_stats["seg_count"]   += 1
                text_stats["char_count"]  += len(txt)
                text_stats["word_count"]  += len(txt.split())
                text_stats["ws_count"]    += sum(1 for c in txt if c.isspace())
                text_stats["other_count"] += sum(1 for c in txt
                                                  if not c.isalnum() and not c.isspace())
    except Exception as _e:
        print(f"[export_preview] text stats error: {_e}")

    # ── table segments ────────────────────────────────────────────────────────
    table_rows = []
    for i, entry in enumerate(table_catalog):
        xml_current    = entry.get("xml") or ""
        xml_pdfplumber = entry.get("xml_pdfplumber") or ""
        # Fall back to the original extracted XML so TEDS can always be
        # computed even when the annotation agent has not been run yet.
        xml_baseline      = xml_pdfplumber or xml_current
        baseline_is_self  = bool(xml_baseline and not xml_pdfplumber)
        transform    = entry.get("transform") or {}
        edited       = bool(xml_pdfplumber and xml_current
                            and xml_current.strip() != xml_pdfplumber.strip())
        has_transform = bool(transform.get("type") and transform["type"] != "none")

        teds_full = teds_struct = None
        if xml_baseline and xml_current:
            try:
                result = compare_snapshots(xml_baseline, xml_current)
                if "error" not in result:
                    teds_full   = result["teds_full"]
                    teds_struct = result["teds_struct"]
            except Exception as _te:
                print(f"[export_preview] teds table {i}: {_te}")

        table_rows.append({
            "index":            i,
            "title":            (entry.get("title") or f"Table {i + 1}")[:55],
            "edited":           edited,
            "has_transform":    has_transform,
            "transform_type":   transform.get("type", ""),
            "transform_color":  transform.get("color", ""),
            "teds_full":        teds_full,
            "teds_struct":      teds_struct,
            "baseline_is_self": baseline_is_self,
        })

    n_edited      = sum(1 for r in table_rows if r["edited"])
    n_transformed = sum(1 for r in table_rows if r["has_transform"])
    return {
        "text_stats":         text_stats,
        "table_rows":         table_rows,
        "total_tables":       len(table_rows),
        "edited_tables":      n_edited,
        "transformed_tables": n_transformed,
    }


## Main function that use helper functions 

# Sidecar file that records the paths of the most recently uploaded files
# so the direct_generate graph node can read them without going through the vectorstore.
_UPLOADED_FILES_RECORD = os.path.join(DATA_DIR, "_uploaded_files.json")

def record_uploaded_files(file_paths: List[str]) -> None:
    """Persist the list of uploaded file paths to disk."""
    import json
    with open(_UPLOADED_FILES_RECORD, "w") as f:
        json.dump(file_paths, f)

def get_uploaded_files() -> List[str]:
    """Return the list of previously uploaded file paths.
    Falls back to scanning DATA_DIR for document files if the record is missing.
    """
    import json
    if os.path.exists(_UPLOADED_FILES_RECORD):
        try:
            with open(_UPLOADED_FILES_RECORD) as f:
                paths = json.load(f)
            # Only return paths that still exist on disk
            return [p for p in paths if os.path.exists(p)]
        except Exception:
            pass
    # Fallback: scan DATA_DIR for any supported document files
    supported = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"}
    found = [
        os.path.join(DATA_DIR, f)
        for f in os.listdir(DATA_DIR)
        if os.path.splitext(f)[1].lower() in supported
    ]
    if found:
        print(f"[get_uploaded_files] No record file found, discovered from data dir: {found}")
        record_uploaded_files(found)  # write the record for next time
    return found

def clear_uploaded_files() -> None:
    """Clear the uploaded files record."""
    if os.path.exists(_UPLOADED_FILES_RECORD):
        os.remove(_UPLOADED_FILES_RECORD)


def load_persisted_state():
    """Reload the vector store and table catalog that were persisted by a previous upload.

    Returns (vectorstore, table_catalog) where either may be None/[] if nothing
    was previously saved.  Safe to call at startup — returns (None, []) if
    DATA_DIR contains no prior upload artefacts.
    """
    import json as _json

    # ── 1. Table catalog ──────────────────────────────────────────────────
    catalog_path = os.path.join(DATA_DIR, "table_catalog.json")
    table_catalog = []
    if os.path.exists(catalog_path):
        try:
            with open(catalog_path, encoding="utf-8") as _f:
                table_catalog = _json.load(_f)
            print(f"[startup] Loaded table catalog: {len(table_catalog)} tables from {catalog_path}")
        except Exception as _e:
            print(f"[startup] Could not read table catalog: {_e}")

    # ── 2. Chroma vector store ────────────────────────────────────────────
    # Only reconnect if the SQLite shard file is present (collection exists).
    chroma_db = os.path.join(DATA_DIR, "chroma.sqlite3")
    vectorstore = None
    if os.path.exists(chroma_db):
        try:
            vectorstore = Chroma(
                collection_name="rag-chroma",
                embedding_function=HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL),
                persist_directory=DATA_DIR,
            )
            n = vectorstore._collection.count()
            print(f"[startup] Reconnected to Chroma: {n} vectors in 'rag-chroma'")
            if n == 0:
                # Empty collection — treat as no prior state
                vectorstore = None
        except Exception as _e:
            print(f"[startup] Could not reconnect to Chroma: {_e}")
            vectorstore = None

    return vectorstore, table_catalog


def upload_files(file_paths, on_progress=None):
    """Upload files into the vector store pipeline.

    Returns:
        (vectorstore, table_catalog) — vectorstore may be None on failure;
        table_catalog is a list of dicts (one per table) with all display fields.
    """
    import json as _json

    if not file_paths:
        print("[upload_files] No documents provided.")
        return None, []

    try:
        # Copy files to DATA_DIR so they persist beyond Gradio's /tmp lifetime
        if on_progress:
            on_progress(0, 0, "Copying files\u2026")
        persistent_paths = []
        for fpath in file_paths:
            dest = os.path.join(DATA_DIR, os.path.basename(fpath))
            if os.path.abspath(fpath) != os.path.abspath(dest):
                shutil.copy2(fpath, dest)
            persistent_paths.append(dest)

        record_uploaded_files(persistent_paths)  # save persistent paths before any processing

        if on_progress:
            on_progress(0, 0, "Extracting tables\u2026")
        docs_list = load_documents_from_files(persistent_paths, on_progress=on_progress)

        if not docs_list:
            print("[upload_files] No documents successfully loaded.")
            return None, []

        # Collect doc_sequence from the first .docx loader result that has it.
        # _load_docx_direct attaches ._doc_sequence to its returned list; that list
        # is wrapped inside load_documents_from_files → we recover it via a scan of
        # each loaded sub-list stored on the flattened docs_list's _doc_sequence.
        doc_sequence = getattr(docs_list, "_doc_sequence", None)

        # Build table catalog from the pre-split documents so metadata is intact
        table_catalog = []
        for doc in docs_list:
            if doc.metadata.get("type") != "table":
                continue
            m = doc.metadata
            try:
                cell_rows = _json.loads(m["cell_rows"]) if m.get("cell_rows") else []
            except Exception:
                cell_rows = []
            try:
                inspection = _json.loads(m["inspection"]) if m.get("inspection") else {}
            except Exception:
                inspection = {}
            table_catalog.append({
                "title":         m.get("title", ""),
                "page_idx":      m.get("page_idx", -1),
                "image_path":    m.get("image_path", ""),
                "pdf_path":      m.get("pdf_path", ""),
                "table_index":   m.get("table_index", 0),
                "xml":           m.get("xml", ""),
                "inspection":    inspection,
                "cell_rows":     cell_rows,
                "source":        m.get("source", ""),
                "doc_seq":       m.get("doc_seq", -1),
                "insertion_key": m.get("insertion_key", ""),
                "annotated_image_path": m.get("annotated_image_path", ""),
                "bbox":          m.get("bbox", None),
            })

        # Persist catalog alongside uploaded data for potential reload
        catalog_path = os.path.join(DATA_DIR, "table_catalog.json")
        try:
            with open(catalog_path, "w", encoding="utf-8") as _f:
                _json.dump(table_catalog, _f, ensure_ascii=False, indent=2)
            print(f"[upload_files] Saved table catalog ({len(table_catalog)} tables) → {catalog_path}")
        except Exception as _ce:
            print(f"[upload_files] Could not write table catalog: {_ce}")

        # Persist document sequence for reconstruction (text + table placeholders)
        if doc_sequence:
            segments_path = os.path.join(DATA_DIR, "doc_segments.json")
            try:
                with open(segments_path, "w", encoding="utf-8") as _f:
                    _json.dump(doc_sequence, _f, ensure_ascii=False, indent=2)
                print(f"[upload_files] Saved doc sequence ({len(doc_sequence)} segments) → {segments_path}")
            except Exception as _se:
                print(f"[upload_files] Could not write doc_segments: {_se}")

        if on_progress:
            on_progress(0, 0, "Chunking text\u2026")
        doc_splits = split_documents(docs_list)
        if on_progress:
            on_progress(0, 0, "Building vector store\u2026")
        vs = embed_documents(doc_splits)
        return vs, table_catalog

    except Exception as e:
        print(f"[upload_files] Pipeline failed: {e}")
        return None, []




def _clear(
    persist_directory: str = None,
    collection_name: str = "rag-chroma",
    delete_all: bool = True
):
    """Clear the Chroma collection and optionally delete all shard folders (excluding hidden files)."""
    if persist_directory is None:
        persist_directory = DATA_DIR
    try:
        # Clear the collection via Chroma client
        vectorstore = Chroma(
            collection_name=collection_name,
            embedding_function=HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL),
            persist_directory=persist_directory,
        )
        vectorstore._client.delete_collection(name=collection_name)
        vectorstore._client.create_collection(name=collection_name)
        print(f"[clear] Collection '{collection_name}' cleared.")

        if delete_all:
            for item in os.listdir(persist_directory):
                if item.startswith(".") or item.startswith("chroma.") or item.startswith("readme-images"):
                    continue  # Skip hidden files like .gitkeep and the sqlite3 file

                path = os.path.join(persist_directory, item)
                try:
                    if os.path.isfile(path):
                        os.remove(path)
                        print(f"[clear] Removed file: {item}")
                    elif os.path.isdir(path):
                        shutil.rmtree(path)
                        print(f"[clear] Removed directory: {item}")
                except Exception as file_err:
                    print(f"[clear] Could not delete {item}: {file_err}")

    except Exception as e:
        print(f"[clear] Failed to clear vector store: {e}")


      
# def clear():
#     """ This is a helper function for emptying the collection the vector store. """
#     vectorstore = Chroma(
#         collection_name="rag-chroma",
#         embedding_function=NVIDIAEmbeddings(model=EMBEDDINGS_MODEL),
#         persist_directory="/project/data",
#     )
    
#     vectorstore._client.delete_collection(name="rag-chroma")
#     vectorstore._client.create_collection(name="rag-chroma")

def get_retriever(): 
    """ This is a helper function for returning the retriever object of the vector store. """
    vectorstore = Chroma(
        collection_name="rag-chroma",
        embedding_function=HuggingFaceEmbeddings(model_name=EMBEDDINGS_MODEL),
        persist_directory=DATA_DIR,
    )
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    return retriever
